import json
import os
import re
import sqlite3
import threading
import time
import traceback
from collections import Counter
from copy import deepcopy
from datetime import datetime
from functools import lru_cache
from hashlib import sha1
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pandas as pd
from flask import Flask, abort, g, has_request_context, jsonify, redirect, render_template, request, send_file, url_for
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

try:
    import psycopg
    from psycopg.rows import dict_row
except Exception:
    psycopg = None
    dict_row = None

try:
    from psycopg_pool import ConnectionPool
except Exception:
    ConnectionPool = None

app = Flask(__name__)

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")
EXCEL_PATH = (
    BASE_DIR / "static" / "files" / "Annual Returns Pilot Sheet as @27th August, 2024.xlsx"
)
WAIVER_SUMMARY_PATH = BASE_DIR / "static" / "files" / "Waiver summary, 2025.xlsx"
DEFAULT_WORKBOOKS_DIR = BASE_DIR / "static" / "files"
CSV_PATH = BASE_DIR / "static" / "files" / "annual_returns_import.csv"
MAPPING_PATH = BASE_DIR / "static" / "files" / "annual_returns_columns.json"
DB_PATH = BASE_DIR / "annual_returns.db"
DATABASE_URL = os.getenv("DATABASE_URL", "").strip()
USE_POSTGRES = (
    psycopg is not None
    and (
        str(DATABASE_URL).strip().lower().startswith("postgresql://")
        or str(DATABASE_URL).strip().lower().startswith("postgres://")
    )
)
SHEET_NAME = "Annual Returns Database"
UPLOAD_DIR = BASE_DIR / "static" / "uploads"

MAX_PER_PAGE = 200
DEFAULT_PER_PAGE = 25
SIDEBAR_LINKS_CACHE_TTL_SECONDS = 15.0


def env_flag(name: str, default: bool = False) -> bool:
    raw = str(os.getenv(name, "1" if default else "0")).strip().lower()
    return raw in {"1", "true", "yes", "y", "on"}


IS_FLASK_RUN_CLI = str(os.getenv("FLASK_RUN_FROM_CLI", "")).strip().lower() in {"1", "true", "yes", "y", "on"}
RUN_STARTUP_MAINTENANCE = env_flag("RUN_STARTUP_MAINTENANCE", default=not IS_FLASK_RUN_CLI)
AUTO_SYNC_DEFAULT_WORKBOOKS = env_flag("AUTO_SYNC_DEFAULT_WORKBOOKS", default=not IS_FLASK_RUN_CLI)
AUTO_ENSURE_TEXT_INDEXES = env_flag("AUTO_ENSURE_TEXT_INDEXES", default=not IS_FLASK_RUN_CLI)

PERF_LOG_HTTP = env_flag("PERF_LOG_HTTP", default=True)
PERF_LOG_SQL = env_flag("PERF_LOG_SQL", default=True)
PERF_SQL_EXPLAIN = env_flag("PERF_SQL_EXPLAIN", default=False)
PERF_SQL_EXPLAIN_ANALYZE = env_flag("PERF_SQL_EXPLAIN_ANALYZE", default=False)
try:
    perf_default = "120" if not IS_FLASK_RUN_CLI else "800"
    PERF_SQL_SLOW_MS = max(1.0, float(os.getenv("PERF_SQL_SLOW_MS", perf_default)))
except Exception:
    PERF_SQL_SLOW_MS = 120.0 if not IS_FLASK_RUN_CLI else 800.0

_sidebar_uploaded_links_cache: list[dict[str, object]] = []
_sidebar_uploaded_links_cache_expires_at = 0.0
_sidebar_uploaded_links_cache_lock = threading.Lock()
_default_workbooks_sync_lock = threading.Lock()
_default_workbooks_synced = False
SELECT_CONFIG_CACHE_TTL_SECONDS = 120.0
DASHBOARD_CACHE_TTL_SECONDS = 90.0
KEYSET_PAGINATION_MIN_PAGE = 8

_cache_lock = threading.Lock()
_data_version = 0
_select_config_cache: dict[tuple[object, ...], tuple[float, object]] = {}
_dashboard_payload_cache: dict[tuple[object, ...], tuple[float, object]] = {}
_pg_pool = None
_pg_pool_lock = threading.Lock()
_uploaded_registry_ready = False
_uploaded_registry_lock = threading.Lock()
UNNAMED_SANITIZE_STATE_KEY = "unnamed_sanitize_v2_complete"


def db_display_target() -> str:
    if USE_POSTGRES:
        try:
            return f"PostgreSQL ({DATABASE_URL.split('@')[-1]})"
        except Exception:
            return "PostgreSQL"
    return str(DB_PATH)


def bump_data_version() -> None:
    global _data_version
    with _cache_lock:
        _data_version += 1
        _select_config_cache.clear()
        _dashboard_payload_cache.clear()
    clear_sidebar_uploaded_links_cache()


def _mapping_cache_fingerprint(mapping: list[dict[str, str]]) -> str:
    compact = [{"db_key": str(x.get("db_key", "")), "label": str(x.get("label", ""))} for x in mapping]
    return sha1(json.dumps(compact, separators=(",", ":"), ensure_ascii=True).encode("utf-8")).hexdigest()[:16]


def _cache_get(
    store: dict[tuple[object, ...], tuple[float, object]],
    key: tuple[object, ...],
) -> object | None:
    now = time.time()
    with _cache_lock:
        item = store.get(key)
        if not item:
            return None
        expires_at, value = item
        if expires_at <= now:
            store.pop(key, None)
            return None
        return deepcopy(value)


def _cache_set(
    store: dict[tuple[object, ...], tuple[float, object]],
    key: tuple[object, ...],
    ttl_seconds: float,
    value: object,
) -> None:
    with _cache_lock:
        store[key] = (time.time() + ttl_seconds, deepcopy(value))


class DbResult:
    def __init__(self, cursor, lastrowid=None):
        self._cursor = cursor
        self.lastrowid = lastrowid if lastrowid is not None else getattr(cursor, "lastrowid", None)

    def fetchone(self):
        return self._cursor.fetchone()

    def fetchall(self):
        return self._cursor.fetchall()


class DbConnection:
    def __init__(self, backend: str, conn, release_callback=None):
        self.backend = backend
        self._conn = conn
        self._release_callback = release_callback

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        if exc_type is None:
            try:
                self._conn.commit()
            except Exception:
                pass
        else:
            try:
                self._conn.rollback()
            except Exception:
                pass
        self._finalize()

    def commit(self):
        self._conn.commit()

    def rollback(self):
        self._conn.rollback()

    def close(self):
        self._finalize()

    def _finalize(self):
        if self._release_callback is not None:
            self._release_callback(self._conn)
            return
        self._conn.close()

    def _pg_table_info(self, table_name: str) -> DbResult:
        cur = self._conn.execute(
            """
            SELECT
                column_name AS name
            FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = %s
            ORDER BY ordinal_position
            """,
            (table_name,),
        )
        return DbResult(cur)

    def _rewrite_sql_for_postgres(self, sql: str) -> str:
        text = str(sql)
        text = re.sub(
            r"INTEGER\s+PRIMARY\s+KEY\s+AUTOINCREMENT",
            "BIGSERIAL PRIMARY KEY",
            text,
            flags=re.IGNORECASE,
        )
        text = text.replace("?", "%s")
        return text

    def execute(self, sql: str, params=None) -> DbResult:
        start = time.perf_counter()
        explain_ctx: tuple[str, tuple] | None = None
        params = tuple(params) if params is not None else tuple()
        if self.backend == "postgres":
            pragma = re.match(r"^\s*PRAGMA\s+table_info\(([^)]+)\)\s*$", str(sql), flags=re.IGNORECASE)
            if pragma:
                table_name = pragma.group(1).strip().strip('"').strip("'")
                result = self._pg_table_info(table_name)
                self._log_sql_timing(sql, params, (time.perf_counter() - start) * 1000.0)
                return result

            q = self._rewrite_sql_for_postgres(sql)
            cur = self._conn.execute(q, params)
            lastrowid = None
            if q.lstrip().upper().startswith("INSERT INTO"):
                try:
                    id_cur = self._conn.execute("SELECT LASTVAL() AS id")
                    row = id_cur.fetchone()
                    if isinstance(row, dict):
                        lastrowid = row.get("id")
                    elif row is not None:
                        lastrowid = row[0]
                except Exception:
                    lastrowid = None
            elapsed_ms = (time.perf_counter() - start) * 1000.0
            if self._is_explain_candidate(q):
                explain_ctx = (q, params)
            self._log_sql_timing(q, params, elapsed_ms)
            if explain_ctx:
                self._maybe_log_explain(explain_ctx[0], explain_ctx[1], elapsed_ms)
            return DbResult(cur, lastrowid=lastrowid)

        cur = self._conn.execute(sql, params)
        elapsed_ms = (time.perf_counter() - start) * 1000.0
        self._log_sql_timing(sql, params, elapsed_ms)
        return DbResult(cur, lastrowid=getattr(cur, "lastrowid", None))

    def executemany(self, sql: str, seq_of_params) -> DbResult:
        start = time.perf_counter()
        if self.backend == "postgres":
            q = self._rewrite_sql_for_postgres(sql)
            cur = self._conn.cursor(row_factory=dict_row)
            cur.executemany(q, [tuple(p) for p in seq_of_params])
            self._log_sql_timing(q, (), (time.perf_counter() - start) * 1000.0, is_many=True)
            return DbResult(cur)
        cur = self._conn.executemany(sql, seq_of_params)
        self._log_sql_timing(sql, (), (time.perf_counter() - start) * 1000.0, is_many=True)
        return DbResult(cur)

    def _is_explain_candidate(self, sql: str) -> bool:
        text = str(sql).lstrip().upper()
        if text.startswith("EXPLAIN"):
            return False
        return text.startswith("SELECT") or text.startswith("WITH")

    def _log_sql_timing(self, sql: str, params: tuple, elapsed_ms: float, is_many: bool = False) -> None:
        if has_request_context():
            g.sql_count = int(getattr(g, "sql_count", 0)) + 1
            g.sql_time_ms = float(getattr(g, "sql_time_ms", 0.0)) + elapsed_ms
            if elapsed_ms >= PERF_SQL_SLOW_MS:
                g.sql_slow_count = int(getattr(g, "sql_slow_count", 0)) + 1
        if not PERF_LOG_SQL:
            return
        mode = "SQL_MANY" if is_many else "SQL"
        level = app.logger.warning if elapsed_ms >= PERF_SQL_SLOW_MS else app.logger.info
        first_line = " ".join(str(sql).strip().split())
        route = f"{request.method} {request.path}" if has_request_context() else "no-request"
        level(
            "%s %.2fms route=%s backend=%s slow=%s query=%s",
            mode,
            elapsed_ms,
            route,
            self.backend,
            elapsed_ms >= PERF_SQL_SLOW_MS,
            first_line[:350],
        )

    def _maybe_log_explain(self, sql: str, params: tuple, elapsed_ms: float) -> None:
        if self.backend != "postgres":
            return
        if elapsed_ms < PERF_SQL_SLOW_MS:
            return
        if not PERF_SQL_EXPLAIN:
            return
        explain_prefix = "EXPLAIN (ANALYZE, BUFFERS, VERBOSE, FORMAT JSON) " if PERF_SQL_EXPLAIN_ANALYZE else "EXPLAIN (FORMAT JSON) "
        try:
            row = self._conn.execute(explain_prefix + sql, params).fetchone()
            route = f"{request.method} {request.path}" if has_request_context() else "no-request"
            app.logger.warning("SQL_EXPLAIN route=%s elapsed_ms=%.2f plan=%s", route, elapsed_ms, row)
        except Exception as exc:
            app.logger.warning("SQL_EXPLAIN_FAILED elapsed_ms=%.2f error=%s", elapsed_ms, exc)


def get_db() -> DbConnection:
    if USE_POSTGRES:
        try:
            pool = get_pg_pool()
            if pool is not None:
                conn = pool.getconn()
                return DbConnection("postgres", conn, release_callback=pool.putconn)
            conn = psycopg.connect(DATABASE_URL, row_factory=dict_row)
            return DbConnection("postgres", conn)
        except Exception as exc:
            app.logger.warning("PostgreSQL connection failed, falling back to SQLite: %s", exc)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return DbConnection("sqlite", conn)


def get_pg_pool():
    if not USE_POSTGRES or ConnectionPool is None:
        return None
    global _pg_pool
    if _pg_pool is not None:
        return _pg_pool
    with _pg_pool_lock:
        if _pg_pool is not None:
            return _pg_pool
        try:
            max_size = max(4, int(os.getenv("PG_POOL_MAX_SIZE", "20")))
        except Exception:
            max_size = 20
        try:
            _pg_pool = ConnectionPool(
                conninfo=DATABASE_URL,
                min_size=1,
                max_size=max_size,
                kwargs={"row_factory": dict_row},
                open=True,
            )
        except Exception as exc:
            app.logger.warning("Could not initialize psycopg pool, using direct connections: %s", exc)
            _pg_pool = None
        return _pg_pool


def quote_identifier(name: str) -> str:
    if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", name):
        raise ValueError(f"Invalid SQL identifier: {name}")
    return f'"{name}"'


def safe_index_name(*parts: str) -> str:
    raw = "_".join(re.sub(r"[^a-zA-Z0-9_]+", "_", str(p).strip().lower()).strip("_") for p in parts if p)
    raw = re.sub(r"_+", "_", raw).strip("_") or "idx_generic"
    if len(raw) <= 55:
        return raw
    digest = sha1(raw.encode("utf-8")).hexdigest()[:7]
    return f"{raw[:47]}_{digest}"


def ensure_postgres_trgm_extension(conn: DbConnection) -> None:
    if conn.backend != "postgres":
        return
    try:
        conn.execute("CREATE EXTENSION IF NOT EXISTS pg_trgm")
    except Exception as exc:
        app.logger.warning("pg_trgm extension unavailable: %s", exc)


def ensure_postgres_text_search_indexes(conn: DbConnection, table_name: str, db_keys: list[str]) -> None:
    if conn.backend != "postgres" or not AUTO_ENSURE_TEXT_INDEXES:
        return
    table_ident = quote_identifier(table_name)
    try:
        conn.execute(
            f"CREATE INDEX IF NOT EXISTS {quote_identifier(safe_index_name('idx', table_name, 'excel_row'))} "
            f"ON {table_ident} (excel_row_number)"
        )
    except Exception as exc:
        app.logger.warning("Could not create excel_row index for %s: %s", table_name, exc)
    ensure_postgres_trgm_extension(conn)
    for key in db_keys:
        idx_name = safe_index_name("idx", table_name, key, "trgm")
        try:
            conn.execute(
                f"CREATE INDEX IF NOT EXISTS {quote_identifier(idx_name)} "
                f"ON {table_ident} USING GIN (LOWER({quote_identifier(key)}) gin_trgm_ops)"
            )
        except Exception as exc:
            app.logger.warning("Could not create trigram index %s on %s: %s", idx_name, table_name, exc)


def ensure_uploaded_registry() -> None:
    global _uploaded_registry_ready
    if _uploaded_registry_ready:
        return
    with _uploaded_registry_lock:
        if _uploaded_registry_ready:
            return
        with get_db() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    original_filename TEXT NOT NULL,
                    stored_filename TEXT NOT NULL UNIQUE,
                    uploaded_at TEXT NOT NULL,
                    import_status TEXT NOT NULL DEFAULT 'completed',
                    import_started_at TEXT,
                    import_completed_at TEXT,
                    import_error TEXT
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS uploaded_sheets (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id BIGINT NOT NULL,
                    sheet_name TEXT NOT NULL,
                    table_name TEXT NOT NULL UNIQUE,
                    mapping_json TEXT NOT NULL,
                    validation_json TEXT NOT NULL,
                    row_count INTEGER NOT NULL DEFAULT 0,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(file_id) REFERENCES uploaded_files(id) ON DELETE CASCADE
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS uploaded_refresh_backups (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id BIGINT NOT NULL,
                    source_sheet_id BIGINT NOT NULL,
                    source_sheet_name TEXT,
                    source_table_name TEXT NOT NULL,
                    backup_table_name TEXT NOT NULL UNIQUE,
                    mapping_json TEXT NOT NULL,
                    validation_json TEXT NOT NULL,
                    snapshot_key TEXT,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(file_id) REFERENCES uploaded_files(id) ON DELETE CASCADE
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS dashboard_thresholds (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    scope_type TEXT NOT NULL,
                    scope_key TEXT NOT NULL,
                    config_json TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    UNIQUE(scope_type, scope_key)
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS app_state (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                )
                """
            )
            cols = {row["name"] for row in conn.execute("PRAGMA table_info(uploaded_refresh_backups)").fetchall()}
            if "source_sheet_name" not in cols:
                conn.execute("ALTER TABLE uploaded_refresh_backups ADD COLUMN source_sheet_name TEXT")
            if "snapshot_key" not in cols:
                conn.execute("ALTER TABLE uploaded_refresh_backups ADD COLUMN snapshot_key TEXT")
            conn.execute(
                "UPDATE uploaded_refresh_backups SET snapshot_key = 'legacy_' || id WHERE snapshot_key IS NULL OR snapshot_key = ''"
            )
            file_cols = {row["name"] for row in conn.execute("PRAGMA table_info(uploaded_files)").fetchall()}
            sheet_cols = {row["name"] for row in conn.execute("PRAGMA table_info(uploaded_sheets)").fetchall()}
            row_count_added = False
            if "row_count" not in sheet_cols:
                conn.execute("ALTER TABLE uploaded_sheets ADD COLUMN row_count INTEGER NOT NULL DEFAULT 0")
                row_count_added = True
            if "import_status" not in file_cols:
                conn.execute("ALTER TABLE uploaded_files ADD COLUMN import_status TEXT NOT NULL DEFAULT 'completed'")
            if "import_started_at" not in file_cols:
                conn.execute("ALTER TABLE uploaded_files ADD COLUMN import_started_at TEXT")
            if "import_completed_at" not in file_cols:
                conn.execute("ALTER TABLE uploaded_files ADD COLUMN import_completed_at TEXT")
            if "import_error" not in file_cols:
                conn.execute("ALTER TABLE uploaded_files ADD COLUMN import_error TEXT")
            conn.execute("UPDATE uploaded_files SET import_status = 'completed' WHERE import_status IS NULL OR import_status = ''")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_uploaded_sheets_file_id ON uploaded_sheets(file_id)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_uploaded_sheets_file_id_id ON uploaded_sheets(file_id, id)")
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_uploaded_refresh_backups_file_snapshot ON uploaded_refresh_backups(file_id, snapshot_key)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_uploaded_refresh_backups_source_sheet ON uploaded_refresh_backups(source_sheet_id)"
            )
            if row_count_added:
                stale_counts = conn.execute("SELECT id, table_name FROM uploaded_sheets").fetchall()
            else:
                stale_counts = conn.execute(
                    "SELECT id, table_name FROM uploaded_sheets WHERE row_count IS NULL OR row_count < 0"
                ).fetchall()
            for row in stale_counts:
                table_name = str(row["table_name"])
                if not table_exists(conn, table_name):
                    continue
                total = conn.execute(
                    f"SELECT COUNT(*) AS total FROM {quote_identifier(table_name)}"
                ).fetchone()["total"]
                conn.execute("UPDATE uploaded_sheets SET row_count = ? WHERE id = ?", (int(total), int(row["id"])))
            conn.commit()
        _uploaded_registry_ready = True


def get_app_state_value(key: str) -> str:
    ensure_uploaded_registry()
    with get_db() as conn:
        row = conn.execute("SELECT value FROM app_state WHERE key = ?", (key,)).fetchone()
    if row is None:
        return ""
    return str(row["value"] or "")


def set_app_state_value(key: str, value: str) -> None:
    ensure_uploaded_registry()
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO app_state (key, value, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(key) DO UPDATE SET value = excluded.value, updated_at = excluded.updated_at
            """,
            (key, str(value), now),
        )
        conn.commit()


DEFAULT_DASHBOARD_THRESHOLDS: dict[str, float] = {
    "late_filing_target_pct": 20.0,
    "recovery_target_pct": 70.0,
    "outstanding_target": 50000.0,
    "avg_processing_target_days": 14.0,
    "late_filing_alert_pct": 40.0,
    "recovery_alert_pct_min": 40.0,
    "outstanding_alert_max": 100000.0,
    "quality_alert_min": 70.0,
}


def normalize_thresholds(raw: dict[str, object] | None) -> dict[str, float]:
    merged = dict(DEFAULT_DASHBOARD_THRESHOLDS)
    if not isinstance(raw, dict):
        return merged
    for key, default in DEFAULT_DASHBOARD_THRESHOLDS.items():
        val = raw.get(key, default)
        try:
            merged[key] = float(val)
        except Exception:
            merged[key] = float(default)
    return merged


def get_dashboard_thresholds(scope_type: str, scope_key: str) -> dict[str, float]:
    ensure_uploaded_registry()
    with get_db() as conn:
        row = conn.execute(
            """
            SELECT config_json
            FROM dashboard_thresholds
            WHERE scope_type = ? AND scope_key = ?
            """,
            (scope_type, scope_key),
        ).fetchone()
    if row is None:
        return dict(DEFAULT_DASHBOARD_THRESHOLDS)
    try:
        parsed = json.loads(row["config_json"])
    except Exception:
        parsed = {}
    return normalize_thresholds(parsed)


def save_dashboard_thresholds(scope_type: str, scope_key: str, config: dict[str, object]) -> dict[str, float]:
    ensure_uploaded_registry()
    clean = normalize_thresholds(config)
    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO dashboard_thresholds (scope_type, scope_key, config_json, updated_at)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(scope_type, scope_key)
            DO UPDATE SET config_json = excluded.config_json, updated_at = excluded.updated_at
            """,
            (scope_type, scope_key, json.dumps(clean), datetime.utcnow().isoformat()),
        )
        conn.commit()
    return clean


def column_mapping_from_headers(headers: list[str]) -> list[dict[str, str]]:
    used: dict[str, int] = {}
    mapping = []
    for idx, label in enumerate(headers, start=1):
        db_key = sanitize_column(label, idx, used)
        mapping.append({"label": label, "db_key": db_key})
    return mapping


def is_unnamed_label(label: str) -> bool:
    text = str(label or "").strip().lower()
    if not text:
        return False
    return re.fullmatch(r"unnamed(?:[\s_:]*[0-9]+)?", text) is not None


def normalize_compact_label(text: str) -> str:
    return " ".join(re.sub(r"[^a-z0-9]+", " ", str(text).lower()).split())


WAIVER_RULE_FILE_NAME = "Waiver summary, 2025.xlsx"
WAIVER_RULE_SHEETS = {
    "WSapril,25",
    "wsmay25",
    "june, 2025",
    "july, 2025",
    "August, 2025",
    "sept_oct_2025",
    "Nov_Dec_2025",
    "jan,2026",
}
WAIVER_RULE_FILE_NORM = normalize_compact_label(WAIVER_RULE_FILE_NAME)
WAIVER_RULE_SHEET_NORMS = {normalize_compact_label(name) for name in WAIVER_RULE_SHEETS}
WAIVER_COMMENT_AI_ENABLED = env_flag("WAIVER_COMMENT_AI_ENABLED", default=True)
WAIVER_COMMENT_AI_MODEL = str(os.getenv("WAIVER_COMMENT_AI_MODEL", "gpt-5-nano")).strip() or "gpt-5-nano"
_waiver_ai_cache_lock = threading.Lock()
_waiver_ai_cache: dict[tuple[str, float], dict[str, object]] = {}


def is_waiver_rule_target_sheet(file_name: str, sheet_name: str) -> bool:
    return (
        normalize_compact_label(file_name) == WAIVER_RULE_FILE_NORM
        and normalize_compact_label(sheet_name) in WAIVER_RULE_SHEET_NORMS
    )


def get_requested_waiver_and_balance_keys(mapping: list[dict[str, str]]) -> tuple[str | None, str | None]:
    requested_key = None
    balance_key = None
    balance_best_score = -10_000
    for item in mapping:
        if not isinstance(item, dict):
            continue
        key = str(item.get("db_key", "")).strip()
        if not key:
            continue
        label_norm = normalize_compact_label(item.get("label", ""))
        key_norm = normalize_compact_label(key)
        if requested_key is None and (
            ("requested" in label_norm and "waiver" in label_norm)
            or ("requested" in key_norm and "waiver" in key_norm)
        ):
            requested_key = key
        if "balance" in label_norm or "balance" in key_norm:
            score = 0
            if label_norm == "balance" or key_norm == "balance":
                score += 100
            if "pending" in label_norm or "pending" in key_norm:
                score -= 60
            if "balance" in label_norm:
                score += 10
            if "balance" in key_norm:
                score += 10
            if score > balance_best_score:
                balance_best_score = score
                balance_key = key
    return requested_key, balance_key


def get_total_penalty_and_penalty_paid_keys(mapping: list[dict[str, str]]) -> tuple[str | None, str | None]:
    total_penalty_key = None
    penalty_paid_key = None
    for item in mapping:
        if not isinstance(item, dict):
            continue
        key = str(item.get("db_key", "")).strip()
        if not key:
            continue
        label_norm = normalize_compact_label(item.get("label", ""))
        key_norm = normalize_compact_label(key)
        text_norm = f"{label_norm} {key_norm}".strip()
        if total_penalty_key is None and "total" in text_norm and "penalty" in text_norm:
            total_penalty_key = key
        if penalty_paid_key is None and "penalty" in text_norm and "paid" in text_norm:
            penalty_paid_key = key
    return total_penalty_key, penalty_paid_key


def get_waiver_rule_keys(mapping: list[dict[str, str]]) -> dict[str, object]:
    requested_key, balance_key = get_requested_waiver_and_balance_keys(mapping)
    total_penalty_key, penalty_paid_key = get_total_penalty_and_penalty_paid_keys(mapping)
    comment_candidates: list[tuple[int, int, str]] = []
    committee_comment_candidates: list[tuple[int, int, str]] = []
    status_key = None
    granted_waiver_key = None
    for idx, item in enumerate(mapping):
        if not isinstance(item, dict):
            continue
        key = str(item.get("db_key", "")).strip()
        if not key:
            continue
        label_norm = normalize_compact_label(item.get("label", ""))
        key_norm = normalize_compact_label(key)
        # Only target the dedicated Status column; do not use FILES STATUS.
        if status_key is None:
            if key_norm == "status" or label_norm == "status":
                status_key = key
            elif (
                "status" in label_norm
                and "file" not in label_norm
                and "files" not in label_norm
                and "file" not in key_norm
                and "files" not in key_norm
            ):
                status_key = key
        if granted_waiver_key is None:
            text_norm_g = f"{label_norm} {key_norm}".strip()
            if "waiver" in text_norm_g and ("granted" in text_norm_g or "approved" in text_norm_g):
                granted_waiver_key = key
        text_norm = f"{label_norm} {key_norm}".strip()
        has_comment_signal = any(
            term in text_norm for term in ("comment", "remarks", "remark", "recommendation", "review note")
        )
        has_committee_signal = "committee" in text_norm
        has_payment_signal = "payment" in text_norm or "payable" in text_norm or "to pay" in text_norm
        if has_comment_signal or (has_committee_signal and has_payment_signal):
            score = 0
            if has_committee_signal:
                score += 100
            if has_comment_signal:
                score += 30
            if has_payment_signal:
                score += 20
            if "waiver" in text_norm:
                score += 5
            comment_candidates.append((score, -idx, key))
            if has_committee_signal:
                committee_comment_candidates.append((score, -idx, key))

    committee_comment_candidates.sort(reverse=True)
    comment_candidates.sort(reverse=True)
    comment_keys = [item[2] for item in committee_comment_candidates]
    if not comment_keys:
        # Strict waiver rule: if no committee comment column, do not infer from other fields.
        comment_keys = []
    comment_key = comment_keys[0] if comment_keys else None
    return {
        "requested": requested_key,
        "balance": balance_key,
        "total_penalty": total_penalty_key,
        "penalty_paid": penalty_paid_key,
        "granted_waiver": granted_waiver_key,
        "comment": comment_key,
        "comment_keys": comment_keys,
        "status": status_key,
    }


def parse_amount_value(value: object) -> float:
    text = str(value or "").strip()
    if not text:
        return 0.0
    compact = text.lower().replace(",", "")
    short = re.search(r"(-?\d+(?:\.\d+)?)\s*([km])\b", compact)
    if short:
        base = float(short.group(1))
        unit = short.group(2)
        mult = 1000.0 if unit == "k" else 1_000_000.0
        return max(0.0, base * mult)
    text = re.sub(r"[^\d.\-]+", "", compact)
    try:
        return max(0.0, float(text))
    except Exception:
        return 0.0


def format_amount_value(value: float) -> str:
    if abs(value - round(value)) < 1e-9:
        return str(int(round(value)))
    return f"{value:.2f}"


def _waiver_comment_heuristic(comment_text: str, requested_amount: float) -> dict[str, object]:
    comment = str(comment_text or "").strip()
    norm = normalize_compact_label(comment)
    if not norm:
        return {"decision": "unknown", "status": "", "subtract_amount": 0.0}

    reject_terms = [
        "reject",
        "rejected",
        "not approved",
        "declined",
        "denied",
        "no waiver",
    ]
    if any(term in norm for term in reject_terms):
        return {"decision": "rejected", "status": "Rejected", "subtract_amount": 0.0}

    full_terms = [
        "full waiver",
        "waived in full",
        "fully waived",
        "waiver granted in full",
    ]
    if any(term in norm for term in full_terms):
        return {"decision": "full_waiver", "status": "Full Waiver", "subtract_amount": max(0.0, requested_amount)}

    pay_patterns = [
        r"(?:to\s*pay|pay)\s*(?:ksh|kes|usd|ugx|tzs|amount|amount of)?\s*([0-9][0-9,]*(?:\.[0-9]+)?)",
        r"(?:payable|pay)\s*(?:is|=)?\s*(?:ksh|kes|usd|ugx|tzs)?\s*([0-9][0-9,]*(?:\.[0-9]+)?)",
    ]
    for pattern in pay_patterns:
        match = re.search(pattern, comment, flags=re.IGNORECASE)
        if match:
            pay_amount = parse_amount_value(match.group(1))
            if pay_amount > 0:
                subtract_amount = min(pay_amount, max(0.0, requested_amount))
                return {"decision": "to_pay_amount", "status": "Partially Approved", "subtract_amount": subtract_amount}

    payment_field_match = re.search(
        r"(?:payment|payments|amount)\s*(?:after committee comments)?\s*(?:is|=|:)?\s*(?:ksh|kes|usd|ugx|tzs)?\s*([0-9][0-9,]*(?:\.[0-9]+)?)",
        comment,
        flags=re.IGNORECASE,
    )
    if payment_field_match:
        pay_amount = parse_amount_value(payment_field_match.group(1))
        if pay_amount > 0:
            return {
                "decision": "to_pay_amount",
                "status": "Partially Approved",
                "subtract_amount": min(pay_amount, max(0.0, requested_amount)),
            }

    if re.fullmatch(r"\s*[0-9][0-9,]*(?:\.[0-9]+)?\s*", comment):
        pay_amount = parse_amount_value(comment)
        if pay_amount > 0:
            return {
                "decision": "to_pay_amount",
                "status": "Partially Approved",
                "subtract_amount": min(pay_amount, max(0.0, requested_amount)),
            }

    if "approved" in norm and "waiver" in norm:
        return {"decision": "full_waiver", "status": "Approved", "subtract_amount": max(0.0, requested_amount)}

    return {"decision": "unknown", "status": "", "subtract_amount": 0.0}


def _extract_amounts_from_comment(comment_text: str) -> dict[str, object]:
    text = str(comment_text or "")
    text_stripped = text.strip()
    amounts = {
        "to_pay": [],
        "waived": [],
        "percent_waiver": None,
        "full_waiver": False,
        "rejected": False,
    }
    norm = normalize_compact_label(text)
    if not norm:
        return amounts

    if any(term in norm for term in ["reject", "rejected", "declined", "denied", "no waiver"]):
        amounts["rejected"] = True
    if any(term in norm for term in ["full waiver", "fully waived", "waived in full", "waiver granted in full"]):
        amounts["full_waiver"] = True

    pct = re.search(r"(\d{1,3}(?:\.\d+)?)\s*%\s*waiver", text, flags=re.IGNORECASE)
    if pct:
        try:
            pct_v = float(pct.group(1))
            if 0.0 <= pct_v <= 100.0:
                amounts["percent_waiver"] = pct_v
        except Exception:
            pass

    # If the committee comment value is a plain amount (e.g. 10000, 20,000, 25k),
    # treat it as granted waiver amount unless it's a duration token.
    numeric_token = r"\s*[0-9][0-9,]*(?:\.[0-9]+)?(?:\s*[kKmM])?\s*"
    duration_re = r"\b(day|days|week|weeks|month|months|year|years)\b"
    if re.fullmatch(numeric_token, text_stripped):
        if not re.search(duration_re, text_stripped, flags=re.IGNORECASE):
            bare_amt = parse_amount_value(text_stripped)
            if bare_amt > 0:
                amounts["waived"].append(bare_amt)
                return amounts

    # Also support prefixed forms such as "committee_s_comments: 50000".
    for line in text.splitlines():
        candidate = str(line or "").strip()
        if not candidate:
            continue
        if ":" in candidate:
            candidate = candidate.split(":", 1)[1].strip()
        if re.fullmatch(numeric_token, candidate) and not re.search(duration_re, candidate, flags=re.IGNORECASE):
            bare_amt = parse_amount_value(candidate)
            if bare_amt > 0:
                amounts["waived"].append(bare_amt)

    amt_re = re.compile(r"(?<!\d)(\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)(?:\s*([kKmM]))?(?!\d)")
    for match in amt_re.finditer(text):
        raw_num = match.group(1)
        suffix = (match.group(2) or "").lower()
        amount = parse_amount_value(f"{raw_num}{suffix}")
        if amount <= 0:
            continue
        # Ignore duration counters such as "21 days", "24 months", etc.
        after_text = text[match.end() : match.end() + 24].lower()
        if re.match(r"^\s*(day|days|week|weeks|month|months|year|years)\b", after_text):
            continue

        ctx_start = max(0, match.start() - 48)
        ctx_end = min(len(text), match.end() + 48)
        context = normalize_compact_label(text[ctx_start:ctx_end])
        if "reinstatement" in context or ("filing" in context and "fee" in context):
            continue
        if "to pay" in context or "payable" in context or "pay " in context:
            amounts["to_pay"].append(amount)
        elif "waiv" in context or "grant" in context or "approve" in context:
            amounts["waived"].append(amount)
    return amounts


def compute_waiver_financials(
    comment_text: str,
    total_penalty: float,
    penalty_paid: float,
) -> dict[str, float | str]:
    total_penalty = max(0.0, float(total_penalty or 0.0))
    penalty_paid = max(0.0, float(penalty_paid or 0.0))
    remaining_before_waiver = max(0.0, total_penalty - penalty_paid)

    extracted = _extract_amounts_from_comment(comment_text)
    granted_waiver = 0.0
    status_text = ""

    if remaining_before_waiver <= 0:
        granted_waiver = 0.0
    elif bool(extracted.get("rejected")):
        granted_waiver = 0.0
        status_text = "No Payment" if penalty_paid <= 0 else "Partially Paid"
    elif bool(extracted.get("full_waiver")):
        granted_waiver = remaining_before_waiver
        status_text = "Fully Paid"
    elif isinstance(extracted.get("percent_waiver"), (int, float)):
        pct = max(0.0, min(100.0, float(extracted.get("percent_waiver") or 0.0)))
        granted_waiver = remaining_before_waiver * (pct / 100.0)
    elif extracted.get("waived"):
        granted_waiver = max(float(x) for x in extracted.get("waived", []) if float(x) > 0.0)
    elif extracted.get("to_pay"):
        to_pay = min(float(x) for x in extracted.get("to_pay", []) if float(x) >= 0.0)
        granted_waiver = max(0.0, remaining_before_waiver - to_pay)
    else:
        # Strict rule: only explicit committee-comment signals affect granted waiver.
        granted_waiver = 0.0

    granted_waiver = max(0.0, min(granted_waiver, remaining_before_waiver))
    balance = max(0.0, remaining_before_waiver - granted_waiver)

    if balance <= 1e-9:
        status_text = "Fully Paid"
    elif penalty_paid <= 0 and granted_waiver <= 0:
        status_text = "No Payment"
    else:
        status_text = "Partially Paid"

    return {
        "granted_waiver": granted_waiver,
        "balance": balance,
        "status": status_text,
    }


def _waiver_comment_ai_decision(comment_text: str, requested_amount: float) -> dict[str, object]:
    if not WAIVER_COMMENT_AI_ENABLED:
        return {"decision": "unknown", "status": "", "subtract_amount": 0.0}
    api_key = str(os.getenv("OPENAI_API_KEY", "")).strip()
    if not api_key:
        return {"decision": "unknown", "status": "", "subtract_amount": 0.0}

    cache_key = (str(comment_text or "").strip().lower(), round(float(requested_amount or 0.0), 2))
    with _waiver_ai_cache_lock:
        cached = _waiver_ai_cache.get(cache_key)
    if cached is not None:
        return dict(cached)

    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        prompt = (
            "You classify a waiver committee comment.\n"
            "Return JSON only with keys: decision, pay_amount, status.\n"
            "decision must be one of: full_waiver, rejected, to_pay_amount, unknown.\n"
            f"requested_amount={requested_amount}\n"
            f"comment={comment_text}\n"
        )
        response = client.responses.create(model=WAIVER_COMMENT_AI_MODEL, input=prompt, store=False)
        raw = str(getattr(response, "output_text", "") or "").strip()
        data = json.loads(raw)
        decision = str(data.get("decision", "unknown")).strip().lower()
        status = str(data.get("status", "")).strip()
        if decision == "full_waiver":
            out = {"decision": decision, "status": status or "Full Waiver", "subtract_amount": max(0.0, requested_amount)}
        elif decision == "rejected":
            out = {"decision": decision, "status": status or "Rejected", "subtract_amount": 0.0}
        elif decision == "to_pay_amount":
            pay_amount = parse_amount_value(data.get("pay_amount", 0))
            out = {
                "decision": decision,
                "status": status or "Partially Approved",
                "subtract_amount": min(pay_amount, max(0.0, requested_amount)),
            }
        else:
            out = {"decision": "unknown", "status": status, "subtract_amount": 0.0}
    except Exception:
        out = {"decision": "unknown", "status": "", "subtract_amount": 0.0}

    with _waiver_ai_cache_lock:
        _waiver_ai_cache[cache_key] = dict(out)
    return out


def evaluate_waiver_comment_decision(comment_text: str, requested_amount: float) -> dict[str, object]:
    rule = _waiver_comment_heuristic(comment_text, requested_amount)
    if str(rule.get("decision", "")) != "unknown":
        return rule
    ai_rule = _waiver_comment_ai_decision(comment_text, requested_amount)
    if str(ai_rule.get("decision", "")) != "unknown":
        return ai_rule
    return {
        "decision": "unknown",
        "status": "",
        "subtract_amount": 0.0,
    }


def apply_waiver_balance_and_status_to_values(
    values: dict[str, object],
    mapping: list[dict[str, str]],
    file_name: str,
    sheet_name: str,
) -> dict[str, object]:
    if not is_waiver_rule_target_sheet(file_name, sheet_name):
        return values
    keys = get_waiver_rule_keys(mapping)
    balance_key = keys.get("balance")
    total_penalty_key = keys.get("total_penalty")
    penalty_paid_key = keys.get("penalty_paid")
    granted_waiver_key = keys.get("granted_waiver")
    comment_key = keys.get("comment")
    comment_keys_raw = keys.get("comment_keys")
    comment_keys = [str(k) for k in (comment_keys_raw or []) if str(k).strip()]
    if not comment_keys and comment_key:
        comment_keys = [str(comment_key)]
    status_key = keys.get("status")
    if not balance_key:
        return values
    if not (total_penalty_key and penalty_paid_key):
        return values

    total_penalty = parse_amount_value(values.get(total_penalty_key, "")) if total_penalty_key else 0.0
    penalty_paid = parse_amount_value(values.get(penalty_paid_key, "")) if penalty_paid_key else 0.0
    comment_parts: list[str] = []
    for c_key in comment_keys:
        part = str(values.get(c_key, "") or "").strip()
        if not part:
            continue
        # Preserve column context so payment-style fields are parsed correctly.
        comment_parts.append(f"{c_key}: {part}")
    comment_text = "\n".join(comment_parts)
    calc = compute_waiver_financials(comment_text, total_penalty, penalty_paid)
    values[balance_key] = format_amount_value(float(calc["balance"]))
    if granted_waiver_key:
        values[granted_waiver_key] = format_amount_value(float(calc["granted_waiver"]))
    status_text = str(calc.get("status", "")).strip()
    if status_key and status_text:
        values[status_key] = status_text
    return values


def apply_waiver_balance_rule_to_updates(
    updates: list[dict],
    mapping: list[dict[str, str]],
    file_name: str,
    sheet_name: str,
) -> list[dict]:
    for item in updates:
        values = item.get("values", {})
        if not isinstance(values, dict):
            continue
        apply_waiver_balance_and_status_to_values(values, mapping, file_name, sheet_name)
    return updates


def drop_forbidden_unnamed_columns(df: pd.DataFrame) -> pd.DataFrame:
    drop_cols = [col for col in df.columns if is_unnamed_label(str(col))]
    if not drop_cols:
        return df
    return df.drop(columns=drop_cols)


def mapping_has_forbidden_unnamed_columns(mapping_json: str) -> bool:
    try:
        mapping = json.loads(mapping_json)
    except Exception:
        return False
    if not isinstance(mapping, list):
        return False
    for item in mapping:
        if not isinstance(item, dict):
            continue
        if is_unnamed_label(str(item.get("label", ""))):
            return True
    return False


def detect_header_row_index(
    path: Path,
    sheet_name: str,
    max_scan_rows: int = 20,
    xls: pd.ExcelFile | None = None,
) -> int:
    try:
        if xls is not None:
            preview = xls.parse(
                sheet_name=sheet_name,
                header=None,
                dtype=str,
                keep_default_na=False,
                nrows=max_scan_rows,
            )
        else:
            preview = pd.read_excel(
                path,
                sheet_name=sheet_name,
                header=None,
                dtype=str,
                keep_default_na=False,
                nrows=max_scan_rows,
            )
    except Exception:
        return 0
    if preview.empty:
        return 0

    keywords = {
        "name",
        "organization",
        "ngo",
        "year",
        "date",
        "status",
        "penalty",
        "fee",
        "return",
        "action",
        "comment",
        "waiver",
        "audit",
        "form",
        "income",
    }
    best_idx = 0
    best_score = float("-inf")
    for idx, row in preview.iterrows():
        vals = [str(v).strip() for v in row.tolist()]
        non_empty = [v for v in vals if v != ""]
        if not non_empty:
            continue
        alpha_cells = sum(bool(re.search(r"[A-Za-z]", v)) for v in non_empty)
        numeric_cells = sum(bool(re.fullmatch(r"[\d,.\-]+", v)) for v in non_empty)
        row_text = " ".join(non_empty).lower()
        keyword_hits = sum(1 for k in keywords if k in row_text)
        score = (len(non_empty) * 3) + (alpha_cells * 2) + (keyword_hits * 4) - (numeric_cells * 2)
        if len(non_empty) <= 2 and keyword_hits == 0:
            score -= 12
        if "summary" in row_text and len(non_empty) <= 3:
            score -= 10
        if score > best_score:
            best_score = score
            best_idx = int(idx)
    return best_idx


def extract_validation_for_sheet(
    path: Path,
    sheet_name: str,
    mapping: list[dict[str, str]],
    header_row_number: int = 1,
    workbook=None,
) -> dict[str, list[str]]:
    try:
        wb = workbook or load_workbook(path, data_only=False)
        ws = wb[sheet_name]
    except Exception:
        return {}

    map_by_label = {item["label"].strip().lower(): item["db_key"] for item in mapping}
    header_by_col: dict[int, str] = {}
    for col_idx in range(1, ws.max_column + 1):
        value = ws.cell(row=header_row_number, column=col_idx).value
        if value is None:
            continue
        label = str(value).strip()
        if label.lower() == "unnamed: 0":
            label = "Organization Name"
        header_by_col[col_idx] = label.lower()

    out: dict[str, list[str]] = {}
    for validation in ws.data_validations.dataValidation:
        if validation.type != "list":
            continue
        options = parse_validation_formula(validation.formula1, wb, ws.title)
        if not options:
            continue
        for sqref in str(validation.sqref).split():
            min_col, _, max_col, _ = range_boundaries(sqref)
            for col_idx in range(min_col, max_col + 1):
                label = header_by_col.get(col_idx)
                if not label:
                    continue
                db_key = map_by_label.get(label)
                if not db_key:
                    continue
                merged = out.setdefault(db_key, [])
                for option in options:
                    clean = str(option).strip()
                    if clean and clean not in merged:
                        merged.append(clean)
    return out


def create_uploaded_sheet_table(
    conn: sqlite3.Connection,
    file_id: int,
    sheet_index: int,
    sheet_name: str,
    mapping: list[dict[str, str]],
    validation_map: dict[str, list[str]],
    frame: pd.DataFrame,
    sheet_row_id: int | None = None,
) -> None:
    table_name = f"uploaded_data_{file_id}_{sheet_index}"
    table_ident = quote_identifier(table_name)

    db_cols_sql = ", ".join(f'{quote_identifier(item["db_key"])} TEXT' for item in mapping)
    conn.execute(
        f"""
        CREATE TABLE {table_ident} (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            excel_row_number INTEGER NOT NULL UNIQUE,
            {db_cols_sql}
        )
        """
    )

    db_keys = [item["db_key"] for item in mapping]
    key_sql = ", ".join(quote_identifier(k) for k in db_keys)
    placeholders = ", ".join("?" for _ in db_keys)
    insert_sql = f"INSERT INTO {table_ident} (excel_row_number, {key_sql}) VALUES (?, {placeholders})"
    records: list[tuple[object, ...]] = []
    for row in frame[["excel_row_number"] + db_keys].itertuples(index=False, name=None):
        records.append(
            tuple([int(row[0])] + [str(v) if v is not None else "" for v in row[1:]])
        )
    if records:
        conn.executemany(insert_sql, records)
    conn.execute(
        f"CREATE INDEX IF NOT EXISTS {quote_identifier(safe_index_name('idx', table_name, 'excel_row'))} "
        f"ON {table_ident}(excel_row_number)"
    )
    ensure_postgres_text_search_indexes(conn, table_name, db_keys)

    if sheet_row_id is None:
        conn.execute(
            """
            INSERT INTO uploaded_sheets
            (file_id, sheet_name, table_name, mapping_json, validation_json, row_count, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                file_id,
                sheet_name,
                table_name,
                json.dumps(mapping),
                json.dumps(validation_map),
                int(len(frame)),
                datetime.utcnow().isoformat(),
            ),
        )
    else:
        conn.execute(
            """
            INSERT INTO uploaded_sheets
            (id, file_id, sheet_name, table_name, mapping_json, validation_json, row_count, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                sheet_row_id,
                file_id,
                sheet_name,
                table_name,
                json.dumps(mapping),
                json.dumps(validation_map),
                int(len(frame)),
                datetime.utcnow().isoformat(),
            ),
        )


def import_workbook_sheets_into_file(
    conn: DbConnection, path: Path, file_id: int, existing_sheet_ids: list[int] | None = None
) -> None:
    start_total = time.perf_counter()
    xls = pd.ExcelFile(path)
    wb = load_workbook(path, data_only=False)
    for index, sheet_name in enumerate(xls.sheet_names, start=1):
        start_sheet = time.perf_counter()
        header_idx = detect_header_row_index(path, sheet_name, xls=xls)
        df = xls.parse(sheet_name=sheet_name, header=header_idx, dtype=str, keep_default_na=False)
        df = drop_forbidden_unnamed_columns(df)
        if len(df.columns) >= 2:
            first_col = df.columns[0]
            second_col = df.columns[1]
            first_text = str(first_col).strip().lower() if first_col is not None else ""
            second_text = normalize_label(str(second_col))
            first_numeric = pd.to_numeric(df[first_col], errors="coerce")
            first_numeric_ratio = float(first_numeric.notna().mean()) if len(df) else 0.0
            if first_text.startswith("unnamed") and first_numeric_ratio >= 0.8 and any(
                token in second_text for token in ["name", "organization", "ngo"]
            ):
                df = df.drop(columns=[first_col])

        drop_cols: list[str] = []
        for col in df.columns:
            col_text = str(col).strip().lower() if col is not None else ""
            if not col_text.startswith("unnamed"):
                continue
            numeric = pd.to_numeric(df[col], errors="coerce")
            non_na = numeric.dropna()
            if non_na.empty:
                continue
            seq = pd.Series(range(1, len(df) + 1), index=df.index, dtype=float)
            aligned = ((numeric == seq) | numeric.isna()).mean()
            numeric_fill_ratio = float(non_na.count() / len(df)) if len(df) else 0.0
            if aligned >= 0.8 and numeric_fill_ratio >= 0.8:
                drop_cols.append(col)
        if drop_cols:
            df = df.drop(columns=drop_cols)

        headers = []
        for col_idx, col in enumerate(df.columns, start=1):
            label = str(col).strip() if col is not None else ""
            if not label:
                label = f"Unnamed_{col_idx}"
            if label.lower() == "unnamed: 0":
                label = "Organization Name"
            headers.append(label)
        df.columns = headers

        mapping = column_mapping_from_headers(headers)
        rename_map = {item["label"]: item["db_key"] for item in mapping}
        clean_df = df.rename(columns=rename_map).fillna("")
        clean_df = clean_df[
            clean_df.apply(lambda row: any(str(value).strip() != "" for value in row.values), axis=1)
        ].copy()
        first_data_row = header_idx + 2
        clean_df.insert(0, "excel_row_number", range(first_data_row, first_data_row + len(clean_df)))
        validation_map = extract_validation_for_sheet(
            path,
            sheet_name,
            mapping,
            header_row_number=header_idx + 1,
            workbook=wb,
        )
        create_uploaded_sheet_table(
            conn=conn,
            file_id=file_id,
            sheet_index=index,
            sheet_name=sheet_name,
            mapping=mapping,
            validation_map=validation_map,
            frame=clean_df,
            sheet_row_id=existing_sheet_ids[index - 1] if existing_sheet_ids and index - 1 < len(existing_sheet_ids) else None,
        )
        elapsed_sheet = time.perf_counter() - start_sheet
        app.logger.info(
            "Imported sheet '%s' for file_id=%s in %.3fs (rows=%s, cols=%s)",
            sheet_name,
            file_id,
            elapsed_sheet,
            len(clean_df),
            len(mapping),
        )
    try:
        wb.close()
    except Exception:
        pass
    app.logger.info(
        "Completed workbook import for file_id=%s in %.3fs (%s sheet(s))",
        file_id,
        time.perf_counter() - start_total,
        len(xls.sheet_names),
    )
    bump_data_version()
    clear_sidebar_uploaded_links_cache()


def create_uploaded_file_record(
    path: Path, original_filename: str, import_status: str = "processing"
) -> int:
    ensure_uploaded_registry()
    with get_db() as conn:
        cursor = conn.execute(
            """
            INSERT INTO uploaded_files
            (original_filename, stored_filename, uploaded_at, import_status, import_started_at, import_completed_at, import_error)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                original_filename,
                path.name,
                datetime.utcnow().isoformat(),
                import_status,
                datetime.utcnow().isoformat(),
                None if import_status == "processing" else datetime.utcnow().isoformat(),
                None,
            ),
        )
        conn.commit()
        clear_sidebar_uploaded_links_cache()
        return int(cursor.lastrowid)


def mark_uploaded_file_import_status(file_id: int, status: str, error_text: str = "") -> None:
    with get_db() as conn:
        conn.execute(
            """
            UPDATE uploaded_files
            SET import_status = ?, import_completed_at = ?, import_error = ?
            WHERE id = ?
            """,
            (
                status,
                datetime.utcnow().isoformat() if status in {"completed", "failed"} else None,
                (error_text or "")[:5000],
                file_id,
            ),
        )
        conn.commit()


def import_uploaded_workbook(path: Path, original_filename: str) -> int:
    file_id = create_uploaded_file_record(path, original_filename, import_status="processing")
    with get_db() as conn:
        import_workbook_sheets_into_file(conn, path, file_id)
        conn.commit()
    enforce_waiver_balance_rule_for_existing_uploaded_sheets()
    mark_uploaded_file_import_status(file_id, "completed")
    return file_id


def import_uploaded_workbook_async(file_id: int, path: Path) -> None:
    try:
        with get_db() as conn:
            import_workbook_sheets_into_file(conn, path, file_id)
            conn.commit()
        enforce_waiver_balance_rule_for_existing_uploaded_sheets()
        mark_uploaded_file_import_status(file_id, "completed")
    except Exception:
        err = traceback.format_exc()
        with get_db() as conn:
            rows = conn.execute(
                "SELECT table_name FROM uploaded_sheets WHERE file_id = ? ORDER BY id ASC",
                (file_id,),
            ).fetchall()
            for row in rows:
                conn.execute(f"DROP TABLE IF EXISTS {quote_identifier(row['table_name'])}")
            conn.execute("DELETE FROM uploaded_sheets WHERE file_id = ?", (file_id,))
            conn.commit()
        mark_uploaded_file_import_status(file_id, "failed", err)


def sync_default_workbooks_from_static_files(force: bool = False) -> None:
    global _default_workbooks_synced
    if _default_workbooks_synced and not force:
        return

    with _default_workbooks_sync_lock:
        if _default_workbooks_synced and not force:
            return

        ensure_uploaded_registry()
        if not DEFAULT_WORKBOOKS_DIR.exists():
            _default_workbooks_synced = True
            return

        workbook_paths = sorted(
            p
            for p in DEFAULT_WORKBOOKS_DIR.iterdir()
            if p.is_file() and p.suffix.lower() == ".xlsx" and not p.name.startswith("~$")
        )
        if not workbook_paths:
            _default_workbooks_synced = True
            return

        for path in workbook_paths:
            file_id: int | None = None
            try:
                with get_db() as conn:
                    existing = conn.execute(
                        """
                        SELECT id, import_status
                        FROM uploaded_files
                        WHERE stored_filename = ?
                        ORDER BY id ASC
                        LIMIT 1
                        """,
                        (path.name,),
                    ).fetchone()
                    if existing is not None:
                        file_id = int(existing["id"])
                        sheet_rows = conn.execute(
                            "SELECT id, table_name, mapping_json FROM uploaded_sheets WHERE file_id = ? ORDER BY id ASC",
                            (file_id,),
                        ).fetchall()
                        has_sheets = bool(sheet_rows)
                        needs_cleanup = any(
                            mapping_has_forbidden_unnamed_columns(str(row["mapping_json"] or "[]"))
                            for row in sheet_rows
                        )
                        if has_sheets:
                            if needs_cleanup:
                                for row in sheet_rows:
                                    conn.execute(f"DROP TABLE IF EXISTS {quote_identifier(str(row['table_name']))}")
                                conn.execute("DELETE FROM uploaded_sheets WHERE file_id = ?", (file_id,))
                            else:
                                if str(existing["import_status"] or "") != "completed":
                                    conn.execute(
                                        """
                                        UPDATE uploaded_files
                                        SET import_status = 'completed', import_completed_at = ?, import_error = ''
                                        WHERE id = ?
                                        """,
                                        (datetime.utcnow().isoformat(), file_id),
                                    )
                                    conn.commit()
                                continue
                        conn.execute(
                            """
                            UPDATE uploaded_files
                            SET original_filename = ?, import_status = 'processing',
                                import_started_at = ?, import_completed_at = NULL, import_error = ''
                            WHERE id = ?
                            """,
                            (path.name, datetime.utcnow().isoformat(), file_id),
                        )
                        conn.commit()
                    else:
                        file_id = create_uploaded_file_record(path, path.name, import_status="processing")

                    import_workbook_sheets_into_file(conn, path, file_id)
                    conn.commit()
                enforce_waiver_balance_rule_for_existing_uploaded_sheets()
                mark_uploaded_file_import_status(file_id, "completed")
            except Exception:
                app.logger.exception("Default workbook auto-import failed for %s", path)
                if file_id is not None:
                    mark_uploaded_file_import_status(file_id, "failed", traceback.format_exc())

        _default_workbooks_synced = True


def get_uploaded_sheet_meta(sheet_id: int) -> sqlite3.Row | None:
    ensure_uploaded_registry()
    with get_db() as conn:
        return conn.execute(
            """
            SELECT s.*, f.original_filename, f.stored_filename, f.id AS file_id
            FROM uploaded_sheets s
            JOIN uploaded_files f ON f.id = s.file_id
            WHERE s.id = ?
            """,
            (sheet_id,),
        ).fetchone()


def get_uploaded_file_meta(file_id: int) -> sqlite3.Row | None:
    ensure_uploaded_registry()
    with get_db() as conn:
        return conn.execute(
            """
            SELECT id, original_filename, stored_filename, import_status, import_error, import_started_at, import_completed_at
            FROM uploaded_files
            WHERE id = ?
            """,
            (file_id,),
        ).fetchone()


def backup_uploaded_file_state(conn: sqlite3.Connection, file_id: int) -> int:
    rows = conn.execute(
        """
        SELECT id, sheet_name, table_name, mapping_json, validation_json
        FROM uploaded_sheets
        WHERE file_id = ?
        ORDER BY id ASC
        """,
        (file_id,),
    ).fetchall()
    if not rows:
        return 0

    stamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    snapshot_key = f"{file_id}_{stamp}_{uuid4().hex[:8]}"
    backed_up = 0
    for idx, row in enumerate(rows, start=1):
        source_table = str(row["table_name"])
        backup_table = f"backup_uploaded_{file_id}_{stamp}_{idx}"
        conn.execute(
            f"CREATE TABLE {quote_identifier(backup_table)} AS SELECT * FROM {quote_identifier(source_table)}"
        )
        conn.execute(
            """
            INSERT INTO uploaded_refresh_backups
            (file_id, source_sheet_id, source_sheet_name, source_table_name, backup_table_name, mapping_json, validation_json, snapshot_key, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                file_id,
                int(row["id"]),
                str(row["sheet_name"]),
                source_table,
                backup_table,
                row["mapping_json"],
                row["validation_json"],
                snapshot_key,
                datetime.utcnow().isoformat(),
            ),
        )
        backed_up += 1
    return backed_up


def get_backup_groups_for_file(file_id: int) -> list[dict[str, object]]:
    ensure_uploaded_registry()
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT snapshot_key, source_sheet_name, source_sheet_id, created_at
            FROM uploaded_refresh_backups
            WHERE file_id = ?
            ORDER BY id DESC
            """,
            (file_id,),
        ).fetchall()
    grouped: dict[str, dict[str, object]] = {}
    for row in rows:
        key = str(row["snapshot_key"] or f"legacy_{row['source_sheet_id']}")
        group = grouped.get(key)
        if not group:
            group = {"snapshot_key": key, "created_at": row["created_at"], "sheets": []}
            grouped[key] = group
        name = str(row["source_sheet_name"] or f"Sheet {row['source_sheet_id']}")
        if name not in group["sheets"]:
            group["sheets"].append(name)
    return list(grouped.values())


def restore_uploaded_file_snapshot(file_id: int, snapshot_key: str) -> int:
    with get_db() as conn:
        backup_rows = conn.execute(
            """
            SELECT source_sheet_id, source_sheet_name, source_table_name, backup_table_name, mapping_json, validation_json
            FROM uploaded_refresh_backups
            WHERE file_id = ? AND snapshot_key = ?
            ORDER BY source_sheet_id ASC
            """,
            (file_id, snapshot_key),
        ).fetchall()
        if not backup_rows:
            return 0

        backup_uploaded_file_state(conn, file_id)

        current_rows = conn.execute(
            "SELECT table_name FROM uploaded_sheets WHERE file_id = ? ORDER BY id ASC",
            (file_id,),
        ).fetchall()
        for row in current_rows:
            conn.execute(f"DROP TABLE IF EXISTS {quote_identifier(row['table_name'])}")
        conn.execute("DELETE FROM uploaded_sheets WHERE file_id = ?", (file_id,))

        for row in backup_rows:
            source_table = str(row["source_table_name"])
            backup_table = str(row["backup_table_name"])
            if not table_exists(conn, backup_table):
                continue
            conn.execute(f"DROP TABLE IF EXISTS {quote_identifier(source_table)}")
            conn.execute(
                f"CREATE TABLE {quote_identifier(source_table)} AS SELECT * FROM {quote_identifier(backup_table)}"
            )
            conn.execute(
                """
                INSERT INTO uploaded_sheets
                (id, file_id, sheet_name, table_name, mapping_json, validation_json, row_count, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    int(row["source_sheet_id"]),
                    file_id,
                    str(row["source_sheet_name"] or f"Sheet {row['source_sheet_id']}"),
                    source_table,
                    row["mapping_json"],
                    row["validation_json"],
                    int(
                        conn.execute(
                            f"SELECT COUNT(*) AS total FROM {quote_identifier(source_table)}"
                        ).fetchone()["total"]
                    ),
                    datetime.utcnow().isoformat(),
                ),
            )
            restored_mapping = json.loads(row["mapping_json"])
            restored_keys = [str(item.get("db_key", "")) for item in restored_mapping if isinstance(item, dict)]
            ensure_postgres_text_search_indexes(conn, source_table, restored_keys)
        conn.commit()
    bump_data_version()
    return len(backup_rows)


def bulk_update_rows(table_name: str, mapping: list[dict[str, str]], updates: list[dict]) -> int:
    table_ident = quote_identifier(table_name)
    allowed_keys = {item["db_key"] for item in mapping}
    updated_count = 0

    with get_db() as conn:
        for item in updates:
            row_id = item.get("row_id")
            values = item.get("values", {})
            if not isinstance(row_id, int) or not isinstance(values, dict):
                continue

            clean = {}
            for key, value in values.items():
                if key in allowed_keys:
                    clean[key] = str(value).strip() if value is not None else ""
            if not clean:
                continue

            set_clause = ", ".join(f"{quote_identifier(k)} = ?" for k in clean)
            params = list(clean.values()) + [row_id]
            conn.execute(
                f"UPDATE {table_ident} SET {set_clause} WHERE id = ?",
                params,
            )
            updated_count += 1
        conn.commit()
    if updated_count > 0:
        bump_data_version()
    return updated_count

def sanitize_column(name: str, index: int, used: dict[str, int]) -> str:
    key = re.sub(r"[^a-zA-Z0-9]+", "_", name.strip().lower()).strip("_")
    if not key:
        key = f"column_{index}"
    if key[0].isdigit():
        key = f"col_{key}"
    counter = used.get(key, 0)
    used[key] = counter + 1
    if counter > 0:
        key = f"{key}_{counter + 1}"
    return key


def build_csv_from_excel() -> list[dict[str, str]]:
    if not EXCEL_PATH.exists():
        raise FileNotFoundError(f"Excel file not found: {EXCEL_PATH}")

    source_df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, dtype=str, keep_default_na=False)
    original_columns = []
    orgname_counter = 0
    for idx, col in enumerate(source_df.columns, start=1):
        text = str(col).strip() if col is not None else ""
        if not text:
            text = f"Unnamed_{idx}"
        if is_unnamed_label(text):
            orgname_counter += 1
            text = "ORGNAME" if orgname_counter == 1 else f"ORGNAME_{orgname_counter}"
        original_columns.append(text)
    source_df.columns = original_columns

    used: dict[str, int] = {}
    mapping = []
    renamed = {}
    for idx, label in enumerate(original_columns, start=1):
        orgname_match = re.fullmatch(r"ORGNAME(?:_(\d+))?", str(label))
        if orgname_match:
            suffix = orgname_match.group(1)
            base = "orgname" if suffix is None else f"orgname_{suffix}"
            db_key = base
            if db_key in used:
                n = 2
                while f"{base}_{n}" in used:
                    n += 1
                db_key = f"{base}_{n}"
            used[db_key] = 1
        else:
            db_key = sanitize_column(label, idx, used)
        mapping.append({"label": label, "db_key": db_key})
        renamed[label] = db_key

    clean_df = source_df.rename(columns=renamed).fillna("")
    clean_df = clean_df[
        clean_df.apply(lambda row: any(str(value).strip() != "" for value in row.values), axis=1)
    ].copy()
    clean_df.insert(0, "excel_row_number", range(2, 2 + len(clean_df)))

    CSV_PATH.parent.mkdir(parents=True, exist_ok=True)
    clean_df.to_csv(CSV_PATH, index=False, encoding="utf-8")
    MAPPING_PATH.write_text(json.dumps(mapping, indent=2), encoding="utf-8")
    return mapping


def load_column_mapping() -> list[dict[str, str]]:
    if not MAPPING_PATH.exists() or not CSV_PATH.exists():
        return build_csv_from_excel()
    return json.loads(MAPPING_PATH.read_text(encoding="utf-8"))


def table_exists(conn: DbConnection, table: str) -> bool:
    if conn.backend == "postgres":
        row = conn.execute(
            """
            SELECT table_name
            FROM information_schema.tables
            WHERE table_schema = 'public' AND table_name = ?
            """,
            (table,),
        ).fetchone()
        return row is not None
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name = ?", (table,)
    ).fetchone()
    return row is not None


def expected_table_columns(mapping: list[dict[str, str]]) -> list[str]:
    return ["id", "excel_row_number"] + [item["db_key"] for item in mapping]


def current_table_columns(conn: DbConnection) -> list[str]:
    return [row["name"] for row in conn.execute("PRAGMA table_info(annual_returns)").fetchall()]


def is_table_data_empty(conn: DbConnection, mapping: list[dict[str, str]]) -> bool:
    count_row = conn.execute("SELECT COUNT(*) AS total FROM annual_returns").fetchone()
    if not count_row or count_row["total"] == 0:
        return True

    checks = [f'COALESCE(TRIM("{col["db_key"]}"), \'\') != \'\'' for col in mapping]
    where_clause = " OR ".join(checks)
    row = conn.execute(f"SELECT COUNT(*) AS non_empty_rows FROM annual_returns WHERE {where_clause}").fetchone()
    return not row or row["non_empty_rows"] == 0


def recreate_table_from_csv(conn: DbConnection, mapping: list[dict[str, str]]) -> None:
    conn.execute("DROP TABLE IF EXISTS annual_returns")
    db_cols_sql = ", ".join(f'"{item["db_key"]}" TEXT' for item in mapping)
    conn.execute(
        f"""
        CREATE TABLE annual_returns (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            excel_row_number INTEGER NOT NULL UNIQUE,
            {db_cols_sql}
        )
        """
    )

    frame = pd.read_csv(CSV_PATH, dtype=str, keep_default_na=False).fillna("")
    db_keys = [item["db_key"] for item in mapping]
    col_names = ", ".join(f'"{name}"' for name in db_keys)
    placeholders = ", ".join("?" for _ in db_keys)
    insert_sql = (
        f'INSERT INTO annual_returns (excel_row_number, {col_names}) VALUES (?, {placeholders})'
    )

    records: list[tuple[object, ...]] = []
    for row in frame[["excel_row_number"] + db_keys].itertuples(index=False, name=None):
        records.append(tuple([int(row[0])] + [str(v) for v in row[1:]]))
    if records:
        conn.executemany(insert_sql, records)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_annual_returns_excel_row_number ON annual_returns(excel_row_number)")
    ensure_postgres_text_search_indexes(conn, "annual_returns", db_keys)
    conn.commit()
    bump_data_version()


def init_database() -> list[dict[str, str]]:
    mapping = build_csv_from_excel()
    with get_db() as conn:
        expected = expected_table_columns(mapping)
        needs_recreate = True
        if table_exists(conn, "annual_returns"):
            current = current_table_columns(conn)
            needs_recreate = current != expected or is_table_data_empty(conn, mapping)
        if needs_recreate:
            recreate_table_from_csv(conn, mapping)
        else:
            if AUTO_ENSURE_TEXT_INDEXES:
                ensure_postgres_text_search_indexes(conn, "annual_returns", [item["db_key"] for item in mapping])
    return mapping


def get_mapping() -> list[dict[str, str]]:
    return load_column_mapping()


def parse_paging(data) -> tuple[int, int]:
    try:
        page = max(1, int(data.get("page", 1)))
    except (TypeError, ValueError):
        page = 1
    try:
        per_page = int(data.get("per_page", DEFAULT_PER_PAGE))
    except (TypeError, ValueError):
        per_page = DEFAULT_PER_PAGE
    per_page = max(1, min(per_page, MAX_PER_PAGE))
    return page, per_page


def to_date_input_value(value: str) -> str:
    text = str(value).strip() if value is not None else ""
    if not text:
        return ""
    parsed = pd.to_datetime(text, errors="coerce")
    if pd.isna(parsed):
        parsed = pd.to_datetime(text, errors="coerce", dayfirst=True)
    if pd.isna(parsed):
        return ""
    return parsed.strftime("%Y-%m-%d")


def is_date_field_label(label: str) -> bool:
    normalized = " ".join(label.strip().lower().split())
    date_labels = {
        "date filed by registry for action",
        "date assigned",
        "date acknoledged/notice sent",
    }
    if normalized in date_labels:
        return True
    return normalized.startswith("end of notice perio")


def parse_validation_formula(formula: str, workbook, default_sheet_name: str) -> list[str]:
    if not formula:
        return []

    value = str(formula).strip()
    if value.startswith('"') and value.endswith('"'):
        return [part.strip() for part in value.strip('"').split(",") if part.strip()]

    if value.startswith("="):
        value = value[1:]

    sheet_name = default_sheet_name
    cell_range = value
    if "!" in value:
        sheet_part, range_part = value.split("!", 1)
        sheet_name = sheet_part.strip("'")
        cell_range = range_part

    try:
        min_col, min_row, max_col, max_row = range_boundaries(cell_range)
        ws = workbook[sheet_name]
        options: list[str] = []
        for row in ws.iter_rows(
            min_row=min_row,
            max_row=max_row,
            min_col=min_col,
            max_col=max_col,
        ):
            for cell in row:
                cell_value = str(cell.value).strip() if cell.value is not None else ""
                if cell_value and cell_value not in options:
                    options.append(cell_value)
        return options
    except Exception:
        return []


@lru_cache(maxsize=1)
def get_excel_validation_options() -> dict[str, list[str]]:
    if not EXCEL_PATH.exists():
        return {}

    workbook = load_workbook(EXCEL_PATH, data_only=False)
    ws = workbook[SHEET_NAME]

    header_by_col: dict[int, str] = {}
    for col_idx in range(1, ws.max_column + 1):
        raw_header = ws.cell(row=1, column=col_idx).value
        if raw_header is None:
            continue
        header_by_col[col_idx] = str(raw_header).strip().lower()

    options_by_header: dict[str, list[str]] = {}
    for validation in ws.data_validations.dataValidation:
        if validation.type != "list":
            continue
        options = parse_validation_formula(validation.formula1, workbook, ws.title)
        if not options:
            continue
        for sqref in str(validation.sqref).split():
            min_col, _, max_col, _ = range_boundaries(sqref)
            for col_idx in range(min_col, max_col + 1):
                header = header_by_col.get(col_idx)
                if not header:
                    continue
                merged = options_by_header.setdefault(header, [])
                for option in options:
                    if option not in merged:
                        merged.append(option)
    return options_by_header


def get_select_config(mapping: list[dict[str, str]]) -> dict[str, dict[str, object]]:
    cache_key = ("annual", _data_version, _mapping_cache_fingerprint(mapping))
    cached = _cache_get(_select_config_cache, cache_key)
    if cached is not None:
        return cached

    config: dict[str, dict[str, object]] = {}
    excel_options = get_excel_validation_options()

    def merge_options(*groups: list[str]) -> list[str]:
        merged: list[str] = []
        for group in groups:
            for value in group:
                clean = str(value).strip()
                if clean and clean not in merged:
                    merged.append(clean)
        return merged

    with get_db() as conn:
        for item in mapping:
            key = item["db_key"]
            label = item["label"].lower()
            db_values = [
                row["value"]
                for row in conn.execute(
                    f'SELECT DISTINCT "{key}" AS value FROM annual_returns WHERE TRIM("{key}") != \'\' ORDER BY value'
                ).fetchall()
                if row["value"]
            ]

            if "received by" in label:
                config[key] = {"options": merge_options(db_values), "multiple": False}
            elif label == "filed by" or "filed for action by registry" in label:
                config[key] = {
                    "options": merge_options(excel_options.get(label, []), ["Filed"], db_values),
                    "multiple": False,
                }
            elif "designate" in label and "pco" in label:
                config[key] = {
                    "options": merge_options(excel_options.get(label, []), db_values),
                    "multiple": False,
                }
    _cache_set(_select_config_cache, cache_key, SELECT_CONFIG_CACHE_TTL_SECONDS, config)
    return config


def get_uploaded_file_groups() -> list[dict[str, object]]:
    if AUTO_SYNC_DEFAULT_WORKBOOKS:
        sync_default_workbooks_from_static_files()
    ensure_uploaded_registry()
    with get_db() as conn:
        files = conn.execute(
            """
            SELECT id, original_filename, uploaded_at, import_status, import_error, import_started_at, import_completed_at
            FROM uploaded_files
            ORDER BY id DESC
            """
        ).fetchall()
        sheets = conn.execute(
            "SELECT id, file_id, sheet_name, table_name, mapping_json, row_count FROM uploaded_sheets ORDER BY id ASC"
        ).fetchall()

        by_file = {
            row["id"]: {
                "file_id": row["id"],
                "title": row["original_filename"],
                "uploaded_at": row["uploaded_at"],
                "import_status": row["import_status"],
                "import_error": row["import_error"],
                "import_started_at": row["import_started_at"],
                "import_completed_at": row["import_completed_at"],
                "sheets": [],
            }
            for row in files
        }
        for sheet in sheets:
            group = by_file.get(sheet["file_id"])
            if not group:
                continue
            mapping = json.loads(sheet["mapping_json"])
            group["sheets"].append(
                {
                    "id": sheet["id"],
                    "name": sheet["sheet_name"],
                    "rows": int(sheet["row_count"] or 0),
                    "columns": len(mapping),
                }
            )
    return [group for group in by_file.values()]


def clear_sidebar_uploaded_links_cache() -> None:
    global _sidebar_uploaded_links_cache_expires_at
    with _sidebar_uploaded_links_cache_lock:
        _sidebar_uploaded_links_cache_expires_at = 0.0


def get_sidebar_uploaded_links() -> list[dict[str, object]]:
    global _sidebar_uploaded_links_cache
    global _sidebar_uploaded_links_cache_expires_at

    now = time.time()
    with _sidebar_uploaded_links_cache_lock:
        if now < _sidebar_uploaded_links_cache_expires_at and _sidebar_uploaded_links_cache:
            return _sidebar_uploaded_links_cache

    if AUTO_SYNC_DEFAULT_WORKBOOKS:
        sync_default_workbooks_from_static_files()
    ensure_uploaded_registry()
    with get_db() as conn:
        files = conn.execute(
            "SELECT id, original_filename FROM uploaded_files ORDER BY id DESC"
        ).fetchall()
        sheets = conn.execute(
            "SELECT id, file_id, sheet_name FROM uploaded_sheets ORDER BY file_id DESC, id ASC"
        ).fetchall()

    grouped: dict[int, dict[str, object]] = {
        int(row["id"]): {"title": row["original_filename"], "sheets": []} for row in files
    }
    for row in sheets:
        file_id = int(row["file_id"])
        if file_id not in grouped:
            continue
        grouped[file_id]["sheets"].append(
            {"id": int(row["id"]), "name": row["sheet_name"]}
        )
    result = [grouped[int(row["id"])] for row in files if int(row["id"]) in grouped]
    with _sidebar_uploaded_links_cache_lock:
        _sidebar_uploaded_links_cache = result
        _sidebar_uploaded_links_cache_expires_at = time.time() + SIDEBAR_LINKS_CACHE_TTL_SECONDS
    return result


@app.context_processor
def inject_sidebar_links() -> dict[str, object]:
    links: list[dict[str, object]] = [
        {
            "title": EXCEL_PATH.name,
            "sheets": [
                {
                    "name": SHEET_NAME,
                    "url": url_for("annual_returns_sheet", sheet_name=SHEET_NAME),
                    "active": request.endpoint == "annual_returns_sheet"
                    and request.view_args
                    and request.view_args.get("sheet_name") == SHEET_NAME,
                }
            ],
        }
    ]
    waiver_sheets: list[dict[str, object]] = []
    if WAIVER_SUMMARY_PATH.exists():
        try:
            wb = load_workbook(WAIVER_SUMMARY_PATH, read_only=True, data_only=True)
            for ws in wb.worksheets:
                sheet_name = str(ws.title)
                waiver_sheets.append(
                    {
                        "name": sheet_name,
                        "url": url_for(
                            "default_workbook_sheet_view",
                            workbook_key="waiver_summary_2025",
                            sheet_name=sheet_name,
                        ),
                        "active": request.endpoint == "default_workbook_sheet_view"
                        and request.view_args
                        and request.view_args.get("workbook_key") == "waiver_summary_2025"
                        and request.view_args.get("sheet_name") == sheet_name,
                    }
                )
            wb.close()
        except Exception:
            app.logger.exception("Could not load sidebar links from %s", WAIVER_SUMMARY_PATH)
    links.append({"title": WAIVER_SUMMARY_PATH.name, "sheets": waiver_sheets})
    return {"sidebar_default_links": links}


@app.before_request
def start_request_timer():
    # Upload features are disabled: app reads only from static/files defaults.
    if request.path == "/upload" or request.path.startswith("/uploaded/"):
        abort(410, description="Upload features are disabled. Use static/files sources only.")
    if not PERF_LOG_HTTP:
        return
    g.request_started_at = time.perf_counter()
    g.sql_count = 0
    g.sql_time_ms = 0.0
    g.sql_slow_count = 0


@app.after_request
def log_request_timing(response):
    if not PERF_LOG_HTTP:
        return response
    started = getattr(g, "request_started_at", None)
    if started is None:
        return response
    elapsed_ms = (time.perf_counter() - started) * 1000.0
    app.logger.info(
        "HTTP %s %s status=%s total_ms=%.2f sql_count=%s sql_ms=%.2f slow_sql=%s",
        request.method,
        request.path,
        response.status_code,
        elapsed_ms,
        int(getattr(g, "sql_count", 0)),
        float(getattr(g, "sql_time_ms", 0.0)),
        int(getattr(g, "sql_slow_count", 0)),
    )
    return response


def get_uploaded_select_config(
    table_name: str, mapping: list[dict[str, str]], validation_map: dict[str, list[str]]
) -> dict[str, dict[str, object]]:
    cache_key = (
        "uploaded",
        _data_version,
        table_name,
        _mapping_cache_fingerprint(mapping),
        sha1(json.dumps(validation_map, sort_keys=True, ensure_ascii=True).encode("utf-8")).hexdigest()[:16],
    )
    cached = _cache_get(_select_config_cache, cache_key)
    if cached is not None:
        return cached

    config: dict[str, dict[str, object]] = {}
    dropdown_keys = set(validation_map.keys())
    with get_db() as conn:
        for item in mapping:
            key = item["db_key"]
            if key not in dropdown_keys:
                continue
            options = list(validation_map.get(key, []))
            db_values = [
                row["value"]
                for row in conn.execute(
                    f"SELECT DISTINCT {quote_identifier(key)} AS value FROM {quote_identifier(table_name)} "
                    f"WHERE TRIM({quote_identifier(key)}) != '' ORDER BY value"
                ).fetchall()
                if row["value"]
            ]
            for value in db_values:
                clean = str(value).strip()
                if clean and clean not in options:
                    options.append(clean)
            if options:
                config[key] = {"options": options, "multiple": False}
    _cache_set(_select_config_cache, cache_key, SELECT_CONFIG_CACHE_TTL_SECONDS, config)
    return config


def date_keys_from_mapping(mapping: list[dict[str, str]]) -> list[str]:
    return [item["db_key"] for item in mapping if is_date_field_label(item["label"])]


def normalize_label(text: str) -> str:
    return " ".join(re.sub(r"[^a-z0-9]+", " ", str(text).lower()).split())


def find_column(mapping: list[dict[str, str]], required_terms: list[str], forbidden_terms: list[str] | None = None) -> str | None:
    forbidden_terms = forbidden_terms or []
    for item in mapping:
        label = normalize_label(item["label"])
        if all(term in label for term in required_terms) and all(term not in label for term in forbidden_terms):
            return item["db_key"]
    return None


def extract_year_from_text(value: str) -> int | None:
    m = re.search(r"(?<!\d)(19\d{2}|20\d{2})(?!\d)", str(value))
    return int(m.group(1)) if m else None


def extract_month_year_hint(text: str) -> tuple[int | None, int | None]:
    month = month_to_number(text)
    year = extract_year_from_text(text)
    return month, year


def normalize_year_token(year_text: str) -> int | None:
    raw = str(year_text).strip()
    if not raw.isdigit():
        return None
    if len(raw) == 4:
        year = int(raw)
        return year if 1900 <= year <= 2100 else None
    if len(raw) == 2:
        y2 = int(raw)
        return 2000 + y2 if y2 <= 49 else 1900 + y2
    return None


def extract_year_flexible(text: str) -> int | None:
    year = extract_year_from_text(text)
    if year:
        return year
    raw = str(text).lower()
    compact_match = re.search(
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s*[,._/-]?\s*(\d{2,4})(?!\d)",
        raw,
    )
    if compact_match:
        year_guess = normalize_year_token(compact_match.group(1))
        if year_guess:
            return year_guess

    q_year_match = re.search(r"(?<!\d)q[1-4]\s*[,._/-]?\s*(\d{2,4})(?!\d)", raw)
    if q_year_match:
        year_guess = normalize_year_token(q_year_match.group(1))
        if year_guess:
            return year_guess
    year_q_match = re.search(r"(?<!\d)(\d{2,4})\s*[,._/-]?\s*q[1-4](?!\d)", raw)
    if year_q_match:
        year_guess = normalize_year_token(year_q_match.group(1))
        if year_guess:
            return year_guess

    tail_year = re.search(r"(?:^|[,_./\-\s])(\d{2,4})$", raw)
    if tail_year:
        year_guess = normalize_year_token(tail_year.group(1))
        if year_guess:
            return year_guess
    return None


def parse_sheet_period_hint(sheet_name: str) -> dict[str, object]:
    raw = str(sheet_name).lower()
    token_to_month = {
        "jan": 1,
        "january": 1,
        "feb": 2,
        "february": 2,
        "mar": 3,
        "march": 3,
        "apr": 4,
        "april": 4,
        "may": 5,
        "jun": 6,
        "june": 6,
        "jul": 7,
        "july": 7,
        "aug": 8,
        "august": 8,
        "sep": 9,
        "sept": 9,
        "september": 9,
        "oct": 10,
        "october": 10,
        "nov": 11,
        "november": 11,
        "dec": 12,
        "december": 12,
    }

    months: list[int] = []
    for token in re.findall(r"[a-z0-9]+", raw):
        month_num = token_to_month.get(token)
        if month_num is not None:
            months.append(month_num)

    # Compact month-year names such as WSMAR_26, May25, Sep-Oct_2025.
    compact_hits = re.finditer(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s*[,._/-]?\s*(\d{2,4})(?!\d)",
        raw,
    )
    for hit in compact_hits:
        month_num = token_to_month.get(hit.group(1))
        if month_num is not None:
            months.append(month_num)

    quarter_num = None
    quarter_match = re.search(r"(?<!\d)q([1-4])\s*[,._/-]?\s*(\d{2,4})(?!\d)", raw)
    if quarter_match is None:
        quarter_match = re.search(r"(?<!\d)(\d{2,4})\s*[,._/-]?\s*q([1-4])(?!\d)", raw)
        if quarter_match is not None:
            quarter_num = int(quarter_match.group(2))
    else:
        quarter_num = int(quarter_match.group(1))

    year = extract_year_flexible(sheet_name) or 0
    if quarter_match is not None:
        year_group = quarter_match.group(2) if quarter_num == int(quarter_match.group(1)) else quarter_match.group(1)
        quarter_year = normalize_year_token(year_group)
        if quarter_year:
            year = quarter_year
    if quarter_num in {1, 2, 3, 4} and not months:
        q_start = (quarter_num - 1) * 3 + 1
        months.extend([q_start, q_start + 1, q_start + 2])

    months = sorted(set(months), key=lambda m: m)
    if not months:
        m = month_to_number(sheet_name)
        if m:
            months = [m]

    start_month = min(months) if months else 0
    end_month = max(months) if months else start_month
    if year <= 0:
        year = 9999

    if quarter_num in {1, 2, 3, 4} and start_month and end_month == start_month + 2:
        label = f"Q{quarter_num} {year}"
    elif start_month and end_month and end_month != start_month:
        label = f"{number_to_month_name(start_month)}-{number_to_month_name(end_month)} {year}"
    elif start_month:
        label = f"{number_to_month_name(start_month)} {year}"
    else:
        label = str(sheet_name)

    return {
        "year": int(year),
        "start_month": int(start_month or 0),
        "end_month": int(end_month or 0),
        "label": label,
    }


def to_numeric_series(series: pd.Series | None) -> pd.Series:
    if series is None:
        return pd.Series(dtype=float)
    cleaned = (
        series.astype(str)
        .str.replace(",", "", regex=False)
        .str.replace("$", "", regex=False)
        .str.strip()
    )
    return pd.to_numeric(cleaned, errors="coerce").fillna(0.0)


def to_bool_series(series: pd.Series | None) -> pd.Series:
    if series is None:
        return pd.Series(dtype=bool)
    text = series.astype(str).str.strip().str.lower()
    true_vals = {"yes", "y", "true", "1", "filed", "submitted", "done", "paid", "complete", "completed"}
    false_vals = {"", "nan", "no", "n", "false", "0", "none", "not filed", "pending"}

    numeric = pd.to_numeric(text, errors="coerce")
    out = numeric.gt(0)
    out[text.isin(true_vals)] = True
    out[text.isin(false_vals)] = False
    return out.fillna(False)


def to_date_series(series: pd.Series | None) -> pd.Series:
    if series is None:
        return pd.Series(dtype="datetime64[ns]")
    raw = series.astype(str).str.strip()
    try:
        parsed = pd.to_datetime(raw, errors="coerce", format="mixed")
    except TypeError:
        parsed = pd.to_datetime(raw, errors="coerce")
    remaining = parsed.isna()
    if remaining.any():
        try:
            parsed.loc[remaining] = pd.to_datetime(
                raw[remaining], errors="coerce", dayfirst=True, format="mixed"
            )
        except TypeError:
            parsed.loc[remaining] = pd.to_datetime(raw[remaining], errors="coerce", dayfirst=True)
    remaining = parsed.isna()
    if remaining.any():
        numeric = pd.to_numeric(clean_numeric_text(raw[remaining]), errors="coerce")
        serial_mask = numeric.between(20000, 60000, inclusive="both")
        if serial_mask.any():
            serial_index = numeric[serial_mask].index
            parsed.loc[serial_index] = pd.to_datetime(
                numeric[serial_mask], unit="D", origin="1899-12-30", errors="coerce"
            )
    return parsed


def month_to_number(value: str) -> int | None:
    months = {
        "jan": 1,
        "feb": 2,
        "mar": 3,
        "apr": 4,
        "may": 5,
        "jun": 6,
        "jul": 7,
        "aug": 8,
        "sep": 9,
        "oct": 10,
        "nov": 11,
        "dec": 12,
    }
    text = normalize_label(value)
    if not text:
        return None
    for token in text.split():
        if token[:3] in months:
            return months[token[:3]]
    numeric = re.search(r"\b(1[0-2]|[1-9])\b", text)
    if numeric:
        return int(numeric.group(1))
    return None


def number_to_month_name(month_number: int) -> str:
    try:
        return datetime(2000, int(month_number), 1).strftime("%B")
    except Exception:
        return str(month_number)


def fmt_pct(value: float) -> str:
    return f"{value:.1f}%"


def fmt_num(value: float) -> str:
    return f"{value:,.2f}"


def paragraphize_ai_text(text: str, sentences_per_paragraph: int = 2) -> str:
    clean = " ".join(str(text or "").split())
    if not clean:
        return ""
    parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", clean) if p.strip()]
    if not parts:
        return clean
    blocks: list[str] = []
    chunk_size = max(1, int(sentences_per_paragraph))
    for i in range(0, len(parts), chunk_size):
        blocks.append(" ".join(parts[i : i + chunk_size]).strip())
    return "\n\n".join(blocks)


def build_submission_summary(
    sheet_name_hint: str,
    total_orgs: int,
    compliance_rate: float,
    late_rate: float,
    total_revenue: float,
    total_penalty_paid: float,
    total_penalty_out: float,
    recovery_rate: float,
    avg_process_days: float | int,
    dynamic_dashboard: dict[str, object] | None = None,
    kpi_targets: list[dict[str, object]] | None = None,
    kpi_target_variance_interpretation: str = "",
    snapshot_delta: dict[str, object] | None = None,
    snapshot_delta_interpretation: str = "",
) -> str:
    dynamic_dashboard = dynamic_dashboard or {}
    kpi_targets = kpi_targets or []
    snapshot_delta = snapshot_delta or {}
    overview = dynamic_dashboard.get("overview", {}) if isinstance(dynamic_dashboard, dict) else {}
    insights = dynamic_dashboard.get("insights", []) if isinstance(dynamic_dashboard, dict) else []
    insight_text = " ".join(str(x) for x in insights[:3]) if isinstance(insights, list) else ""
    completeness = float(overview.get("overall_completeness_pct", 0) or 0.0)
    on_target = sum(1 for row in kpi_targets if str(row.get("status", "")).strip().lower() == "on target")
    off_target_rows = [row for row in kpi_targets if str(row.get("status", "")).strip().lower() != "on target"]
    off_target_names = ", ".join(str(row.get("kpi", "")) for row in off_target_rows[:4]) if off_target_rows else "None"

    delta_context = "No snapshot baseline available."
    mode = str(snapshot_delta.get("mode", "")).strip().lower()
    if mode == "period_sheet":
        delta_context = (
            f"Current period {snapshot_delta.get('current_period')} ({snapshot_delta.get('current_sheet')}) "
            f"compared with previous period {snapshot_delta.get('previous_period')} ({snapshot_delta.get('previous_sheet')})."
        )
    elif mode == "timeseries":
        rows = snapshot_delta.get("rows", [])
        if isinstance(rows, list) and rows:
            first = rows[0]
            last = rows[-1]
            delta_context = (
                f"Time-series baseline from {first.get('period')} to {last.get('period')} "
                f"across {len(rows)} chronological sheet snapshots."
            )
    elif snapshot_delta.get("snapshot_key"):
        delta_context = f"Compared against backup snapshot {snapshot_delta.get('snapshot_key')}."

    return (
        f"Management Submission Summary - {sheet_name_hint or 'Current Sheet'}\n\n"
        f"This report covers {total_orgs} records captured in the current compliance dataset. "
        f"Overall compliance is {compliance_rate:.2f}%, while late filing stands at {late_rate:.2f}%. "
        f"Total value realized from filing fees and penalties is {total_revenue + total_penalty_paid:,.2f}, "
        f"with outstanding penalties of {total_penalty_out:,.2f}. Current penalty recovery performance is "
        f"{recovery_rate:.2f}%, indicating {'strong' if recovery_rate >= 70 else 'moderate' if recovery_rate >= 40 else 'weak'} conversion "
        "of liabilities into collections.\n\n"
        f"Operationally, average processing time is {float(avg_process_days):.2f} days. Data completeness is "
        f"{completeness:.2f}%, which should be considered when interpreting trends and risk positions. "
        f"{insight_text}\n\n"
        "KPI Target vs Variance - AI Interpretation:\n"
        f"{on_target} of {len(kpi_targets)} KPI targets are currently on track. "
        f"Off-target KPIs: {off_target_names}. "
        f"{kpi_target_variance_interpretation}\n\n"
        "Snapshot Delta Reporting - AI Interpretation:\n"
        f"{delta_context} "
        f"{snapshot_delta_interpretation}\n\n"
        "Recommended actions for supervisory review:\n"
        "1. Prioritize follow-up for late and high-risk records with clear ownership and deadlines.\n"
        "2. Focus collection activity on high-value outstanding penalties to improve recovery rate.\n"
        "3. Strengthen data-quality controls for frequently missing or inconsistent fields before the next reporting cycle.\n\n"
        "Prepared for management submission."
    )


def build_kpi_target_variance(metrics: dict[str, object], thresholds: dict[str, float] | None = None) -> list[dict[str, object]]:
    thresholds = normalize_thresholds(thresholds or {})

    def get_num(key: str) -> float:
        return float(metrics.get(key, 0) or 0.0)

    targets = [
        ("Late filing rate (%)", "late_filing_rate_pct", thresholds["late_filing_target_pct"], "lower_better"),
        ("Penalty recovery rate (%)", "recovery_rate_pct", thresholds["recovery_target_pct"], "higher_better"),
        ("Outstanding penalties", "total_penalty_outstanding", thresholds["outstanding_target"], "lower_better"),
        ("Avg processing days", "avg_processing_days", thresholds["avg_processing_target_days"], "lower_better"),
    ]
    rows: list[dict[str, object]] = []
    for label, key, target, mode in targets:
        actual = get_num(key)
        variance = actual - target
        if mode == "higher_better":
            status = "On target" if actual >= target else "Below target"
        else:
            status = "On target" if actual <= target else "Above target"
        rows.append(
            {
                "kpi": label,
                "actual": round(actual, 2),
                "target": round(target, 2),
                "variance": round(variance, 2),
                "status": status,
            }
        )
    return rows


def build_kpi_target_variance_interpretation(kpi_targets: list[dict[str, object]]) -> str:
    if not kpi_targets:
        text = (
            "KPI target interpretation is not available because no KPI target rows were generated for this sheet. "
            "This usually means required fields for late filing, recovery, outstanding penalties, or processing time "
            "were missing or could not be parsed reliably. Once those fields are present, this section will explain "
            "which KPIs are meeting target, where variance is highest, and what that means for operational follow-up."
        )
        if len(text.split()) < 70:
            text += (
                " Refresh the data after confirming core numeric fields are mapped correctly."
            )
        return text

    on_target = [row for row in kpi_targets if str(row.get("status", "")).strip().lower() == "on target"]
    off_target = [row for row in kpi_targets if row not in on_target]
    largest_gap = max(kpi_targets, key=lambda row: abs(float(row.get("variance", 0) or 0.0)))
    largest_kpi = str(largest_gap.get("kpi", "Unknown KPI"))
    largest_var = float(largest_gap.get("variance", 0) or 0.0)
    off_names = ", ".join(str(row.get("kpi", "")) for row in off_target[:4])

    text = (
        f"This KPI target review shows {len(on_target)} of {len(kpi_targets)} indicators currently on target, "
        f"with {len(off_target)} requiring attention. The largest variance is in {largest_kpi} at {largest_var:,.2f}, "
        "which represents the biggest gap between current performance and the expected control level. "
        + (
            f"Priority follow-up should focus first on {off_names}. "
            if off_names
            else "At this point, no KPI is materially off target. "
        )
        + "For management use, interpret this table as a prioritization tool: on-target KPIs indicate stable controls, "
        "while off-target KPIs identify where corrective actions, resourcing, or process adjustments are most likely to "
        "improve compliance outcomes in the next reporting cycle."
    )
    if len(text.split()) < 70:
        text += (
            " Reassess these KPI variances after each refresh to confirm whether interventions are reducing the gap."
        )
    return text


def build_trend_normalization(ym_group: pd.DataFrame) -> list[dict[str, object]]:
    if ym_group is None or ym_group.empty:
        return []
    work = ym_group.copy().sort_values(["year", "month"]).reset_index(drop=True)
    work["late_rate_pct"] = work["late_rate"].astype(float) * 100.0
    rows: list[dict[str, object]] = []
    for i, row in work.iterrows():
        prev = work.iloc[i - 1] if i > 0 else None
        prev_year = work[(work["year"] == int(row["year"]) - 1) & (work["month"] == int(row["month"]))]
        mom_filings = None
        mom_late = None
        yoy_filings = None
        yoy_late = None
        if prev is not None and float(prev["filings"]) != 0:
            mom_filings = ((float(row["filings"]) - float(prev["filings"])) / float(prev["filings"])) * 100.0
            mom_late = float(row["late_rate_pct"]) - float(prev["late_rate_pct"])
        if not prev_year.empty and float(prev_year.iloc[0]["filings"]) != 0:
            p = prev_year.iloc[0]
            yoy_filings = ((float(row["filings"]) - float(p["filings"])) / float(p["filings"])) * 100.0
            yoy_late = float(row["late_rate_pct"]) - float(float(p["late_rate"]) * 100.0)
        rows.append(
            {
                "period": f"{int(row['year'])}-{int(row['month']):02d}",
                "filings": int(row["filings"]),
                "late_rate_pct": round(float(row["late_rate_pct"]), 2),
                "mom_filings_pct": round(mom_filings, 2) if mom_filings is not None else None,
                "mom_late_rate_pp": round(mom_late, 2) if mom_late is not None else None,
                "yoy_filings_pct": round(yoy_filings, 2) if yoy_filings is not None else None,
                "yoy_late_rate_pp": round(yoy_late, 2) if yoy_late is not None else None,
            }
        )
    return rows[-24:]


def build_trend_normalization_interpretation(trend_rows: list[dict[str, object]]) -> str:
    if not trend_rows:
        return (
            "Trend normalization interpretation is unavailable because there are not enough monthly periods to compute "
            "month-over-month or year-over-year movement. Once at least two valid periods are available, this section "
            "will explain whether filing volume is accelerating or slowing, and whether late-filing behavior is improving "
            "or worsening over time."
        )

    latest = trend_rows[-1]
    recent_window = trend_rows[-6:] if len(trend_rows) >= 6 else trend_rows
    mom_filings_vals = [float(r["mom_filings_pct"]) for r in recent_window if r.get("mom_filings_pct") is not None]
    mom_late_vals = [float(r["mom_late_rate_pp"]) for r in recent_window if r.get("mom_late_rate_pp") is not None]
    yoy_filings_vals = [float(r["yoy_filings_pct"]) for r in recent_window if r.get("yoy_filings_pct") is not None]
    yoy_late_vals = [float(r["yoy_late_rate_pp"]) for r in recent_window if r.get("yoy_late_rate_pp") is not None]

    def avg_or_none(values: list[float]) -> float | None:
        return (sum(values) / len(values)) if values else None

    avg_mom_filings = avg_or_none(mom_filings_vals)
    avg_mom_late = avg_or_none(mom_late_vals)
    avg_yoy_filings = avg_or_none(yoy_filings_vals)
    avg_yoy_late = avg_or_none(yoy_late_vals)

    mom_volume_text = (
        "is rising"
        if avg_mom_filings is not None and avg_mom_filings > 0
        else "is falling"
        if avg_mom_filings is not None and avg_mom_filings < 0
        else "is broadly stable"
    )
    mom_late_text = (
        "improving"
        if avg_mom_late is not None and avg_mom_late < 0
        else "worsening"
        if avg_mom_late is not None and avg_mom_late > 0
        else "holding steady"
    )
    yoy_volume_text = (
        "higher"
        if avg_yoy_filings is not None and avg_yoy_filings > 0
        else "lower"
        if avg_yoy_filings is not None and avg_yoy_filings < 0
        else "flat"
    )
    yoy_late_text = (
        "improving"
        if avg_yoy_late is not None and avg_yoy_late < 0
        else "worsening"
        if avg_yoy_late is not None and avg_yoy_late > 0
        else "stable"
    )

    return (
        f"Trend normalization compares changes across periods so management can see direction, not just raw totals. "
        f"For the latest period {latest.get('period')}, filings are {int(latest.get('filings', 0) or 0)} and late filing is "
        f"{float(latest.get('late_rate_pct', 0) or 0.0):.2f}%. Over the recent window, month-over-month filing volume {mom_volume_text} "
        f"and month-over-month late-rate movement is {mom_late_text}. Year-over-year, filing volume is {yoy_volume_text} and "
        f"late-rate direction is {yoy_late_text}. Use this section to confirm whether compliance actions are producing sustained "
        f"improvement rather than one-off period changes."
    )


def build_cohort_summary(org_series: pd.Series, late_bool: pd.Series, filing_year: pd.Series) -> dict[str, object]:
    if org_series.empty:
        return {"organizations": 0, "never_late": 0, "one_time_late": 0, "repeat_late": 0}
    frame = pd.DataFrame(
        {
            "org": org_series.astype(str).str.strip().replace("", "Unknown Organization"),
            "late": late_bool.astype(bool),
            "year": filing_year,
        }
    )
    by_org = frame.groupby("org")["late"].agg(["sum", "count"])
    repeat_late = int((by_org["sum"] >= 2).sum())
    one_time_late = int((by_org["sum"] == 1).sum())
    never_late = int((by_org["sum"] == 0).sum())
    return {
        "organizations": int(len(by_org)),
        "never_late": never_late,
        "one_time_late": one_time_late,
        "repeat_late": repeat_late,
    }


def build_anomaly_classification(df: pd.DataFrame, mapping: list[dict[str, str]]) -> list[dict[str, object]]:
    out: list[dict[str, object]] = []
    for item in mapping:
        key = item["db_key"]
        if key not in df.columns:
            continue
        series = pd.to_numeric(clean_numeric_text(df[key]), errors="coerce").dropna()
        if len(series) < 8:
            continue
        q1 = float(series.quantile(0.25))
        q3 = float(series.quantile(0.75))
        iqr = q3 - q1
        if iqr <= 0:
            continue
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        outlier_vals = series[(series < lower) | (series > upper)]
        if outlier_vals.empty:
            continue
        median = float(series.median())
        likely_error = int(((outlier_vals < 0) | ((median > 0) & (outlier_vals > median * 50))).sum())
        exceptional_case = int(len(outlier_vals) - likely_error)
        out.append(
            {
                "field": item["label"],
                "outliers": int(len(outlier_vals)),
                "likely_error": likely_error,
                "exceptional_case": exceptional_case,
            }
        )
    out.sort(key=lambda row: row["outliers"], reverse=True)
    return out[:10]


def build_anomaly_classification_interpretation(anomalies: list[dict[str, object]]) -> str:
    if not anomalies:
        return (
            "No significant anomaly clusters were detected in the current numeric fields. "
            "This suggests values are mostly within expected ranges based on the IQR method, so there is no immediate high-risk "
            "data-quality hotspot. Continue routine checks during each refresh to catch new outliers early, especially after imports "
            "or bulk edits."
        )

    top = anomalies[0]
    total_outliers = int(sum(int(row.get("outliers", 0) or 0) for row in anomalies))
    total_likely_error = int(sum(int(row.get("likely_error", 0) or 0) for row in anomalies))
    total_exceptional = int(sum(int(row.get("exceptional_case", 0) or 0) for row in anomalies))
    likely_share = (total_likely_error / total_outliers * 100.0) if total_outliers else 0.0
    priority_fields = ", ".join(str(row.get("field", "")) for row in anomalies[:3])

    return (
        f"Anomaly classification found {total_outliers} flagged outlier records across {len(anomalies)} fields. "
        f"The highest concentration is in {top.get('field')} with {int(top.get('outliers', 0) or 0)} outliers, "
        f"so this field should be reviewed first. About {likely_share:.1f}% of flagged outliers are currently classified as likely errors "
        f"({total_likely_error} records), while {total_exceptional} appear to be exceptional but potentially valid cases. "
        f"Operational priority should focus on {priority_fields}: first verify data-entry accuracy, then separate true exceptions for "
        f"documented business justification so reporting quality and enforcement decisions remain reliable."
    )


def build_management_templates(dataset_title: str, metrics: dict[str, object], quality: dict[str, object]) -> dict[str, str]:
    late = float(metrics.get("late_filing_rate_pct", 0) or 0.0)
    recovery = float(metrics.get("recovery_rate_pct", 0) or 0.0)
    out_pen = float(metrics.get("total_penalty_outstanding", 0) or 0.0)
    score = float(quality.get("score", 0) or 0.0)
    return {
        "executive": (
            f"Executive Brief ({dataset_title}): late filing is {late:.2f}%, recovery is {recovery:.2f}%, "
            f"and outstanding penalties are {out_pen:,.2f}. Data quality score is {score:.2f}/100."
        ),
        "compliance_ops": (
            f"Compliance Operations Note: prioritize repeat late filers and unresolved notices; "
            f"current late rate is {late:.2f}% with data quality at {score:.2f}/100."
        ),
        "finance_recovery": (
            f"Finance Recovery Note: penalty recovery is {recovery:.2f}% against outstanding exposure "
            f"of {out_pen:,.2f}. Focus first on highest-value aged balances."
        ),
    }


def build_alerts(
    metrics: dict[str, object],
    quality: dict[str, object],
    kpi_targets: list[dict[str, object]],
    thresholds: dict[str, float] | None = None,
) -> list[dict[str, str]]:
    thresholds = normalize_thresholds(thresholds or {})
    alerts: list[dict[str, str]] = []
    late = float(metrics.get("late_filing_rate_pct", 0) or 0.0)
    recovery = float(metrics.get("recovery_rate_pct", 0) or 0.0)
    out_pen = float(metrics.get("total_penalty_outstanding", 0) or 0.0)
    q_score = float(quality.get("score", 0) or 0.0)
    if late > thresholds["late_filing_alert_pct"]:
        alerts.append({"severity": "high", "message": f"Late filing rate is high at {late:.2f}%."})
    if recovery < thresholds["recovery_alert_pct_min"]:
        alerts.append({"severity": "high", "message": f"Penalty recovery is low at {recovery:.2f}%."})
    if out_pen > thresholds["outstanding_alert_max"]:
        alerts.append({"severity": "medium", "message": f"Outstanding penalties are elevated at {out_pen:,.2f}."})
    if q_score < thresholds["quality_alert_min"]:
        alerts.append({"severity": "medium", "message": f"Data quality score is {q_score:.2f}/100; verify missing/invalid fields."})
    below = [row["kpi"] for row in kpi_targets if row.get("status") in {"Below target", "Above target"}]
    if below:
        alerts.append({"severity": "low", "message": f"KPI variance flags: {', '.join(below[:4])}."})
    return alerts[:8]


def build_compliance_numeric_interpretation(
    numeric_summary: list[dict[str, object]], metrics: dict[str, object]
) -> str:
    if not numeric_summary:
        return (
            "This section is meant to explain money and count fields in a practical way for compliance monitoring. "
            "Right now, there are not enough usable numeric values to give a strong operational reading. That usually "
            "means key amount columns were left blank, entered as text, or are inconsistent across rows. Before taking "
            "decisions, first clean and complete the important numeric fields, then regenerate this dashboard. Once numeric "
            "coverage improves, this interpretation will help you identify which values dominate total exposure, where the "
            "typical case sits, and which records look unusual enough to require validation or escalation."
        )

    top_total = max(numeric_summary, key=lambda r: float(r.get("sum", 0) or 0.0))
    top_unusual = max(numeric_summary, key=lambda r: int(r.get("outliers_iqr", 0) or 0))
    labels = [str(r.get("label", "")).lower() for r in numeric_summary]
    financial_hits = sum(any(k in lbl for k in ["fee", "penalty", "income", "waiver", "amount", "balance"]) for lbl in labels)
    count_hits = sum(any(k in lbl for k in ["count", "returns", "cases", "files", "notice"]) for lbl in labels)
    if financial_hits >= max(2, count_hits):
        intro = (
            "This sheet is primarily financial, so this numeric summary focuses on money exposure, collections, and outstanding risk."
        )
    elif count_hits > financial_hits:
        intro = (
            "This sheet is primarily volume-driven, so this numeric summary focuses on case counts, workload levels, and operational pressure."
        )
    else:
        intro = (
            "This sheet has a mixed numeric profile, so this summary combines financial values and operational counts for a balanced view."
        )
    late_rate = float(metrics.get("late_filing_rate_pct", 0) or 0.0)
    recovery = float(metrics.get("recovery_rate_pct", 0) or 0.0)
    outstanding = float(metrics.get("total_penalty_outstanding", 0) or 0.0)

    return (
        f"{intro} "
        f"The largest total value currently appears in {top_total.get('label')} at {float(top_total.get('sum', 0) or 0):,.2f}, "
        "so that field should be treated as a priority when reviewing financial exposure and follow-up actions. "
        "The typical value column gives a stable baseline for day-to-day decision making, because it is less sensitive "
        "to extreme entries than a simple average. Data quality flags also matter: "
        f"{top_unusual.get('label')} has the highest unusual-record count at {int(top_unusual.get('outliers_iqr', 0) or 0)}, "
        "which means records in that field should be reviewed before relying on them for enforcement or reporting. "
        f"At a high level, late filing is {late_rate:.2f}%, recovery is {recovery:.2f}%, and outstanding penalties are "
        f"{outstanding:,.2f}. Operationally, this suggests focusing first on high-value fields, validating unusual entries, "
        "and tracking whether future updates reduce flagged records and improve numeric consistency over time."
    )


def build_categorical_statistics_interpretation(
    categorical_summary: list[dict[str, object]],
) -> str:
    if not categorical_summary:
        text = (
            "No strong categorical fields were detected for this sheet, so category-based interpretation is limited. "
            "To improve this section, ensure text/status fields are populated consistently and avoid mixing many spelling variants "
            "for the same category."
        )
        if len(text.split()) < 100:
            text += (
                " At the moment, this means management decisions should rely more on validated numeric and date indicators "
                "until label quality improves. A practical next step is to standardize key category values through dropdown "
                "controls, normalize capitalization and punctuation, and merge duplicate category spellings that refer to the "
                "same meaning. After cleanup, rerun the dashboard to get stronger distribution patterns and more reliable subgroup analysis."
            )
        return text

    by_unique = sorted(categorical_summary, key=lambda r: int(r.get("unique_count", 0) or 0), reverse=True)
    by_top = sorted(categorical_summary, key=lambda r: int(r.get("top_count", 0) or 0), reverse=True)
    most_diverse = by_unique[0]
    most_dominant = by_top[0]
    dominant_unique = int(most_dominant.get("unique_count", 0) or 0)

    diversity_msg = (
        f"The most diverse field is {most_diverse.get('label')} with {int(most_diverse.get('unique_count', 0) or 0)} unique values, "
        "which suggests varied case types that may need segmented handling."
    )
    dominance_msg = (
        f"The strongest concentration appears in {most_dominant.get('label')}, where the top category "
        f"('{most_dominant.get('top_value')}') occurs {int(most_dominant.get('top_count', 0) or 0)} times. "
    )
    if dominant_unique <= 3:
        dominance_msg += "This indicates a highly uniform pattern, so exceptions outside that category may warrant quick review."
    else:
        dominance_msg += "This indicates meaningful category spread, so policy responses should be tailored by subgroup."

    text = (
        "This categorical interpretation explains how records are distributed across labels such as statuses, reasons, or action groups. "
        + diversity_msg
        + " "
        + dominance_msg
        + " In operational terms, use high-diversity fields to identify where specialized workflows are needed, and use dominant categories "
        "to streamline standard handling paths and monitor deviations."
    )
    if len(text.split()) < 100:
        text += (
            " For compliance teams, the immediate value is prioritization: fields with many unique labels often indicate process "
            "fragmentation, while fields with one dominant label are better candidates for standardized handling. Track these patterns "
            "across refresh cycles to confirm whether interventions are reducing category sprawl and improving consistency in case routing, "
            "review quality, and response time."
        )
    return text


def build_snapshot_delta_interpretation(snapshot_delta: dict[str, object]) -> str:
    if not isinstance(snapshot_delta, dict) or not snapshot_delta.get("available"):
        return "Snapshot delta interpretation is unavailable because no comparable historical baseline was found."

    if snapshot_delta.get("mode") == "period_sheet":
        rows = snapshot_delta.get("rows", [])
        if not isinstance(rows, list) or not rows:
            return "Sheet-to-sheet period comparison is available, but KPI rows are missing."
        late = next((r for r in rows if r.get("kpi") == "Late filing rate (%)"), None)
        recovery = next((r for r in rows if r.get("kpi") == "Recovery rate (%)"), None)
        outstanding = next((r for r in rows if r.get("kpi") == "Outstanding penalties"), None)
        parts = []
        if late:
            d = float(late.get("delta", 0) or 0.0)
            parts.append(
                f"late filing {'improved' if d < 0 else 'worsened' if d > 0 else 'was stable'} by {abs(d):.2f} percentage points"
            )
        if recovery:
            d = float(recovery.get("delta", 0) or 0.0)
            parts.append(
                f"recovery {'improved' if d > 0 else 'declined' if d < 0 else 'was stable'} by {abs(d):.2f} percentage points"
            )
        if outstanding:
            d = float(outstanding.get("delta", 0) or 0.0)
            parts.append(
                f"outstanding penalties {'reduced' if d < 0 else 'increased' if d > 0 else 'stayed stable'} by {abs(d):,.2f}"
            )
        core = "; ".join(parts) if parts else "key KPIs moved moderately between periods"
        return (
            f"This compares {snapshot_delta.get('current_period')} ({snapshot_delta.get('current_sheet')}) "
            f"against the previous period {snapshot_delta.get('previous_period')} ({snapshot_delta.get('previous_sheet')}). "
            f"Overall, {core}. This is the correct month-to-month operational delta for workbook trend tracking."
        )

    if snapshot_delta.get("mode") == "timeseries":
        rows = snapshot_delta.get("rows", [])
        if not isinstance(rows, list) or len(rows) < 2:
            return "Cross-sheet time-series comparison is available, but there are not enough periods to interpret movement."
        first = rows[0]
        last = rows[-1]
        late_delta = float(last.get("late_filing_rate_pct", 0) or 0.0) - float(first.get("late_filing_rate_pct", 0) or 0.0)
        rec_delta = float(last.get("recovery_rate_pct", 0) or 0.0) - float(first.get("recovery_rate_pct", 0) or 0.0)
        out_delta = float(last.get("total_penalty_outstanding", 0) or 0.0) - float(first.get("total_penalty_outstanding", 0) or 0.0)
        trend_late = "improved" if late_delta < 0 else "worsened" if late_delta > 0 else "remained flat"
        trend_rec = "improved" if rec_delta > 0 else "declined" if rec_delta < 0 else "remained flat"
        out_text = "reduced" if out_delta < 0 else "increased" if out_delta > 0 else "stayed stable"
        return (
            f"This time-series delta view compares sheets chronologically from {first.get('period')} to {last.get('period')}. "
            f"Late filing has {trend_late} by {abs(late_delta):.2f} percentage points, recovery has {trend_rec} by {abs(rec_delta):.2f} "
            f"percentage points, and outstanding penalties have {out_text} by {abs(out_delta):,.2f}. "
            "Use this as a directional trend signal: sustained reductions in late filing and outstanding balances, combined with improving recovery, "
            "indicate stronger compliance control. If movement is volatile between periods, review whether policy changes, data quality shifts, or "
            "case-mix differences are driving the variance before escalating conclusions."
        )

    rows = snapshot_delta.get("rows", [])
    if not isinstance(rows, list) or not rows:
        return "Snapshot comparison is available, but no KPI rows were returned for interpretation."
    late = next((r for r in rows if r.get("kpi") == "Late filing rate (%)"), None)
    recovery = next((r for r in rows if r.get("kpi") == "Recovery rate (%)"), None)
    outstanding = next((r for r in rows if r.get("kpi") == "Outstanding penalties"), None)

    parts = []
    if late:
        d = float(late.get("delta", 0) or 0.0)
        parts.append(
            f"late filing {'improved' if d < 0 else 'worsened' if d > 0 else 'was stable'} by {abs(d):.2f} percentage points"
        )
    if recovery:
        d = float(recovery.get("delta", 0) or 0.0)
        parts.append(
            f"recovery {'improved' if d > 0 else 'declined' if d < 0 else 'was stable'} by {abs(d):.2f} percentage points"
        )
    if outstanding:
        d = float(outstanding.get("delta", 0) or 0.0)
        parts.append(
            f"outstanding penalties {'reduced' if d < 0 else 'increased' if d > 0 else 'stayed stable'} by {abs(d):,.2f}"
        )
    core = "; ".join(parts) if parts else "key KPIs moved modestly between snapshots"
    return (
        f"This snapshot delta compares the current sheet to backup {snapshot_delta.get('snapshot_key')}. "
        f"Overall, {core}. Treat positive movement as an operational signal, not proof of causality. "
        "Confirm whether changes are driven by real case progress, backlog clearing, or data corrections before final reporting."
    )


def build_data_quality_score(
    df: pd.DataFrame,
    mapping: list[dict[str, str]],
    kinds: dict[str, str],
    profile: list[dict[str, object]],
) -> dict[str, object]:
    if df.empty or not mapping:
        return {
            "score": 0.0,
            "grade": "D",
            "components": {
                "completeness": 0.0,
                "duplicate_rows": 0.0,
                "invalid_dates": 0.0,
                "numeric_outliers": 0.0,
            },
        }

    total_rows = len(df)
    total_cells = total_rows * len(mapping)
    non_empty_cells = 0
    for item in mapping:
        key = item["db_key"]
        s = df[key].astype(str).str.strip()
        non_empty_cells += int((s != "").sum())
    completeness_pct = (non_empty_cells / total_cells * 100.0) if total_cells else 0.0

    dedup = df.copy()
    for item in mapping:
        k = item["db_key"]
        dedup[k] = dedup[k].astype(str).str.strip()
    duplicate_rate = float(dedup.duplicated(subset=[item["db_key"] for item in mapping]).mean() * 100.0) if total_rows else 0.0

    date_cols = [c for c, kind in kinds.items() if kind == "date"]
    invalid_date_ratios: list[float] = []
    for col in date_cols:
        raw = df[col].astype(str).str.strip()
        non_empty = raw[raw != ""]
        if non_empty.empty:
            continue
        parsed = to_date_series(non_empty)
        invalid_ratio = float(parsed.isna().mean() * 100.0)
        invalid_date_ratios.append(invalid_ratio)
    invalid_date_pct = float(sum(invalid_date_ratios) / len(invalid_date_ratios)) if invalid_date_ratios else 0.0

    numeric_cols = [c for c, kind in kinds.items() if kind == "numeric"]
    outlier_ratios: list[float] = []
    for col in numeric_cols:
        s = pd.to_numeric(clean_numeric_text(df[col]), errors="coerce").dropna()
        if len(s) < 4:
            continue
        q1 = float(s.quantile(0.25))
        q3 = float(s.quantile(0.75))
        iqr = q3 - q1
        if iqr <= 0:
            continue
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        outlier_ratios.append(float(((s < lower) | (s > upper)).mean() * 100.0))
    outlier_pct = float(sum(outlier_ratios) / len(outlier_ratios)) if outlier_ratios else 0.0

    score = (
        completeness_pct * 0.45
        + (100.0 - duplicate_rate) * 0.20
        + (100.0 - invalid_date_pct) * 0.20
        + (100.0 - outlier_pct) * 0.15
    )
    score = max(0.0, min(100.0, score))
    grade = "A" if score >= 85 else "B" if score >= 70 else "C" if score >= 55 else "D"
    return {
        "score": round(score, 2),
        "grade": grade,
        "components": {
            "completeness": round(completeness_pct, 2),
            "duplicate_rows": round(duplicate_rate, 2),
            "invalid_dates": round(invalid_date_pct, 2),
            "numeric_outliers": round(outlier_pct, 2),
        },
    }


def build_snapshot_delta_for_uploaded_sheet(
    sheet_id: int, current_metrics: dict[str, object]
) -> dict[str, object]:
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        return {"available": False, "reason": "Sheet not found."}

    series_rows = build_file_period_series(int(meta["file_id"]))
    current_idx = next((i for i, row in enumerate(series_rows) if int(row["sheet_id"]) == int(sheet_id)), -1)
    if current_idx > 0:
        current_row = series_rows[current_idx]
        previous_row = series_rows[current_idx - 1]
        keys = [
            ("Late filing rate (%)", "late_filing_rate_pct"),
            ("Penalty paid", "total_penalty_paid"),
            ("Outstanding penalties", "total_penalty_outstanding"),
            ("Recovery rate (%)", "recovery_rate_pct"),
            ("Filing fee total", "total_revenue"),
            ("Requested waiver total", "total_requested_waiver"),
            ("Last income total", "total_last_income"),
        ]
        rows: list[dict[str, object]] = []
        for label, key in keys:
            current = float(current_row.get(key, 0) or 0.0)
            previous = float(previous_row.get(key, 0) or 0.0)
            rows.append(
                {
                    "kpi": label,
                    "current": round(current, 2),
                    "previous": round(previous, 2),
                    "delta": round(current - previous, 2),
                }
            )
        return {
            "available": True,
            "mode": "period_sheet",
            "reason": "Showing chronological sheet-to-sheet delta for this workbook.",
            "current_sheet": str(current_row["sheet_name"]),
            "current_period": str(current_row["period_label"]),
            "previous_sheet": str(previous_row["sheet_name"]),
            "previous_period": str(previous_row["period_label"]),
            "rows": rows,
            "file_name": str(meta["original_filename"]),
        }

    with get_db() as conn:
        latest = conn.execute(
            """
            SELECT backup_table_name, mapping_json, created_at, snapshot_key
            FROM uploaded_refresh_backups
            WHERE source_sheet_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (sheet_id,),
        ).fetchone()
        older = conn.execute(
            """
            SELECT backup_table_name, mapping_json, created_at, snapshot_key
            FROM uploaded_refresh_backups
            WHERE source_sheet_id = ?
            ORDER BY id DESC
            LIMIT 1 OFFSET 1
            """,
            (sheet_id,),
        ).fetchone()
    if latest is None:
        return build_cross_sheet_timeseries_delta(sheet_id)

    latest_table = str(latest["backup_table_name"])
    if not USE_POSTGRES and not Path(DB_PATH).exists():
        return {"available": False, "reason": "Database unavailable."}
    with get_db() as conn:
        if not table_exists(conn, latest_table):
            return {"available": False, "reason": "Snapshot table missing."}

    latest_mapping = json.loads(latest["mapping_json"])
    previous_payload = get_cached_dashboard_payload(
        latest_table,
        latest_mapping,
        sheet_name_hint=str(meta["sheet_name"]),
        include_snapshot_delta=False,
    )
    prev_metrics = previous_payload.get("metrics", {}) if isinstance(previous_payload, dict) else {}

    keys = [
        ("Late filing rate (%)", "late_filing_rate_pct"),
        ("Penalty paid", "total_penalty_paid"),
        ("Outstanding penalties", "total_penalty_outstanding"),
        ("Recovery rate (%)", "recovery_rate_pct"),
        ("Filing fee total", "total_revenue"),
        ("Requested waiver total", "total_requested_waiver"),
    ]
    rows: list[dict[str, object]] = []
    for label, key in keys:
        current = float(current_metrics.get(key, 0) or 0.0)
        previous = float(prev_metrics.get(key, 0) or 0.0)
        rows.append(
            {
                "kpi": label,
                "current": round(current, 2),
                "previous": round(previous, 2),
                "delta": round(current - previous, 2),
            }
        )

    result: dict[str, object] = {
        "available": True,
        "snapshot_key": latest["snapshot_key"],
        "compared_at": latest["created_at"],
        "rows": rows,
    }

    if older is not None:
        older_table = str(older["backup_table_name"])
        with get_db() as conn:
            has_older = table_exists(conn, older_table)
        if has_older:
            older_mapping = json.loads(older["mapping_json"])
            older_payload = get_cached_dashboard_payload(
                older_table,
                older_mapping,
                sheet_name_hint=str(meta["sheet_name"]),
                include_snapshot_delta=False,
            )
            older_metrics = older_payload.get("metrics", {}) if isinstance(older_payload, dict) else {}
            history_rows: list[dict[str, object]] = []
            for label, key in keys:
                latest_prev = float(prev_metrics.get(key, 0) or 0.0)
                older_prev = float(older_metrics.get(key, 0) or 0.0)
                history_rows.append(
                    {
                        "kpi": label,
                        "latest_snapshot": round(latest_prev, 2),
                        "older_snapshot": round(older_prev, 2),
                        "delta": round(latest_prev - older_prev, 2),
                    }
                )
            result["history_available"] = True
            result["history_snapshot_key"] = older["snapshot_key"]
            result["history_compared_at"] = older["created_at"]
            result["history_rows"] = history_rows
        else:
            result["history_available"] = False
    else:
        result["history_available"] = False
    return result


def build_file_period_series(file_id: int) -> list[dict[str, object]]:
    with get_db() as conn:
        sheets = conn.execute(
            """
            SELECT id, sheet_name, table_name, mapping_json
            FROM uploaded_sheets
            WHERE file_id = ?
            ORDER BY id ASC
            """,
            (file_id,),
        ).fetchall()
    if len(sheets) < 2:
        return []

    series_rows: list[dict[str, object]] = []
    for row in sheets:
        period = parse_sheet_period_hint(str(row["sheet_name"]))
        mapping = json.loads(row["mapping_json"])
        payload = get_cached_dashboard_payload(
            str(row["table_name"]),
            mapping,
            sheet_name_hint=str(row["sheet_name"]),
            include_snapshot_delta=False,
        )
        m = payload.get("metrics", {}) if isinstance(payload, dict) else {}
        year = int(period["year"])
        start_month = int(period["start_month"])
        end_month = int(period["end_month"])
        if year <= 0 or year >= 9999 or start_month <= 0:
            continue
        if end_month <= 0:
            end_month = start_month
        series_rows.append(
            {
                "sheet_id": int(row["id"]),
                "sheet_name": str(row["sheet_name"]),
                "period_label": str(period["label"]),
                "sort_key": (year, start_month, end_month, int(row["id"])),
                "late_filing_rate_pct": float(m.get("late_filing_rate_pct", 0) or 0.0),
                "recovery_rate_pct": float(m.get("recovery_rate_pct", 0) or 0.0),
                "total_penalty_paid": float(m.get("total_penalty_paid", 0) or 0.0),
                "total_penalty_outstanding": float(m.get("total_penalty_outstanding", 0) or 0.0),
                "total_revenue": float(m.get("total_revenue", 0) or 0.0),
                "total_requested_waiver": float(m.get("total_requested_waiver", 0) or 0.0),
                "total_last_income": float(m.get("total_last_income", 0) or 0.0),
            }
        )

    series_rows.sort(key=lambda r: r["sort_key"])
    return series_rows


def build_cross_sheet_timeseries_delta(sheet_id: int) -> dict[str, object]:
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        return {"available": False, "reason": "Sheet not found."}

    series_rows = build_file_period_series(int(meta["file_id"]))
    if len(series_rows) < 2:
        return {"available": False, "reason": "No previous snapshot found."}
    output_rows: list[dict[str, object]] = []
    for i, row in enumerate(series_rows):
        prev = series_rows[i - 1] if i > 0 else None
        out = {
            "sheet_id": row["sheet_id"],
            "sheet_name": row["sheet_name"],
            "period": row["period_label"],
            "late_filing_rate_pct": round(row["late_filing_rate_pct"], 2),
            "recovery_rate_pct": round(row["recovery_rate_pct"], 2),
            "total_penalty_paid": round(row["total_penalty_paid"], 2),
            "total_penalty_outstanding": round(row["total_penalty_outstanding"], 2),
            "total_requested_waiver": round(row["total_requested_waiver"], 2),
            "is_current": row["sheet_id"] == sheet_id,
            "delta_late_rate_pp": None,
            "delta_recovery_rate_pp": None,
            "delta_penalty_paid": None,
            "delta_outstanding_penalty": None,
            "delta_requested_waiver": None,
        }
        if prev is not None:
            out["delta_late_rate_pp"] = round(row["late_filing_rate_pct"] - prev["late_filing_rate_pct"], 2)
            out["delta_recovery_rate_pp"] = round(row["recovery_rate_pct"] - prev["recovery_rate_pct"], 2)
            out["delta_penalty_paid"] = round(row["total_penalty_paid"] - prev["total_penalty_paid"], 2)
            out["delta_outstanding_penalty"] = round(
                row["total_penalty_outstanding"] - prev["total_penalty_outstanding"], 2
            )
            out["delta_requested_waiver"] = round(
                row["total_requested_waiver"] - prev["total_requested_waiver"], 2
            )
        output_rows.append(out)

    return {
        "available": True,
        "mode": "timeseries",
        "reason": "No previous backup snapshot found; showing chronological cross-sheet delta reporting.",
        "rows": output_rows,
        "file_name": str(meta["original_filename"]),
    }


def summary_text_to_docx(summary_text: str, title: str) -> BytesIO:
    from docx import Document

    doc = Document()
    title_clean = (title or "Management Submission Summary").strip()
    doc.add_heading(title_clean, level=1)

    blocks = [block.strip() for block in str(summary_text).split("\n\n") if block.strip()]
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        if lines[0].endswith(":") and len(lines) == 1:
            doc.add_paragraph(lines[0]).runs[0].bold = True
            continue
        if lines[0].startswith(("1.", "2.", "3.")):
            for line in lines:
                if re.match(r"^\d+\.\s+", line):
                    doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
                else:
                    doc.add_paragraph(line)
            continue
        for line in lines:
            doc.add_paragraph(line)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def clean_numeric_text(series: pd.Series) -> pd.Series:
    return (
        series.astype(str)
        .str.replace(",", "", regex=False)
        .str.replace("$", "", regex=False)
        .str.replace("%", "", regex=False)
        .str.replace("(", "-", regex=False)
        .str.replace(")", "", regex=False)
        .str.strip()
    )


def is_date_like_label(label: str) -> bool:
    text = normalize_label(label)
    return any(token in text for token in [" date", "date ", "date", "registration", "registered", "dob", "birth"])


def is_temporal_field_label(label: str) -> bool:
    text = normalize_label(label)
    temporal_tokens = [
        "date",
        "period",
        "year",
        "month",
        "quarter",
        "fy",
        "financial year",
        "registration",
        "registered",
    ]
    return any(token in text for token in temporal_tokens)


def infer_column_kind(series: pd.Series, label: str = "") -> str:
    non_empty = series.astype(str).str.strip()
    non_empty = non_empty[non_empty != ""]
    if non_empty.empty:
        return "empty"

    normalized = non_empty.str.lower()
    bool_tokens = {
        "yes",
        "no",
        "y",
        "n",
        "true",
        "false",
        "1",
        "0",
        "filed",
        "not filed",
        "pending",
        "complete",
        "completed",
        "paid",
        "unpaid",
    }
    unique_tokens = set(normalized.unique().tolist())
    if len(unique_tokens) <= 6 and unique_tokens.issubset(bool_tokens):
        return "boolean"

    numeric_clean = clean_numeric_text(non_empty)
    numeric_parsed = pd.to_numeric(numeric_clean, errors="coerce")
    numeric_ratio = float(numeric_parsed.notna().mean())
    numeric_non_null = numeric_parsed.dropna()

    try:
        date_parsed = pd.to_datetime(non_empty, errors="coerce", format="mixed")
    except TypeError:
        date_parsed = pd.to_datetime(non_empty, errors="coerce")
    date_ratio = float(date_parsed.notna().mean())
    if date_ratio < 0.7:
        try:
            date_parsed_alt = pd.to_datetime(non_empty, errors="coerce", dayfirst=True, format="mixed")
        except TypeError:
            date_parsed_alt = pd.to_datetime(non_empty, errors="coerce", dayfirst=True)
        date_ratio = max(date_ratio, float(date_parsed_alt.notna().mean()))

    excel_serial_ratio = float(numeric_parsed.between(20000, 60000, inclusive="both").mean()) if len(numeric_parsed) else 0.0
    label_suggests_date = is_date_like_label(label)
    label_suggests_temporal = is_temporal_field_label(label)
    if len(numeric_non_null):
        integer_like = (numeric_non_null.round().eq(numeric_non_null)).mean()
        year_like = (numeric_non_null.between(1900, 2100, inclusive="both")).mean()
        year_like_ratio = float(integer_like * year_like)
    else:
        year_like_ratio = 0.0

    if label_suggests_temporal and (date_ratio >= 0.35 or excel_serial_ratio >= 0.6 or year_like_ratio >= 0.6):
        return "date"
    if label_suggests_date and (date_ratio >= 0.4 or excel_serial_ratio >= 0.6):
        return "date"
    if year_like_ratio >= 0.9 and label_suggests_temporal:
        return "date"

    if date_ratio >= 0.75 and date_ratio >= numeric_ratio:
        return "date"
    if numeric_ratio >= 0.75:
        return "numeric"
    return "categorical"


def build_dynamic_dashboard(df: pd.DataFrame, mapping: list[dict[str, str]]) -> dict[str, object]:
    labels_by_key = {item["db_key"]: item["label"] for item in mapping}
    total_rows = int(len(df))
    total_columns = int(len(mapping))
    profile: list[dict[str, object]] = []
    kinds: dict[str, str] = {}

    for item in mapping:
        key = item["db_key"]
        series = df[key] if key in df.columns else pd.Series(dtype=str)
        missing_count = int(series.astype(str).str.strip().eq("").sum()) if len(series) else 0
        non_empty_count = total_rows - missing_count
        kind = infer_column_kind(series, item["label"]) if len(series) else "empty"
        kinds[key] = kind
        distinct_count = int(
            series.astype(str).str.strip().replace("", pd.NA).dropna().nunique()
        ) if len(series) else 0
        profile.append(
            {
                "key": key,
                "label": labels_by_key.get(key, key),
                "kind": kind,
                "non_empty_count": non_empty_count,
                "missing_count": missing_count,
                "missing_pct": round((missing_count / total_rows * 100.0), 2) if total_rows else 0.0,
                "distinct_count": distinct_count,
            }
        )

    numeric_cols = [c for c, k in kinds.items() if k == "numeric"]
    categorical_cols = [c for c, k in kinds.items() if k == "categorical"]
    date_cols = [c for c, k in kinds.items() if k == "date"]
    boolean_cols = [c for c, k in kinds.items() if k == "boolean"]
    temporal_numeric_cols = [
        c for c in numeric_cols if is_temporal_field_label(labels_by_key.get(c, c))
    ]
    if temporal_numeric_cols:
        blocked = set(temporal_numeric_cols)
        numeric_cols = [c for c in numeric_cols if c not in blocked]
        for col in temporal_numeric_cols:
            if col not in date_cols:
                date_cols.append(col)

    numeric_summary: list[dict[str, object]] = []
    categorical_summary: list[dict[str, object]] = []
    date_summary: list[dict[str, object]] = []
    boolean_summary: list[dict[str, object]] = []
    charts: list[dict[str, object]] = []
    insights: list[str] = []
    chart_id = 1

    for col in numeric_cols:
        s = pd.to_numeric(clean_numeric_text(df[col]), errors="coerce").dropna()
        if s.empty:
            continue
        q1 = float(s.quantile(0.25))
        q3 = float(s.quantile(0.75))
        iqr = q3 - q1
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        outliers = int(((s < lower) | (s > upper)).sum())
        numeric_summary.append(
            {
                "label": labels_by_key.get(col, col),
                "count": int(s.count()),
                "mean": round(float(s.mean()), 2),
                "median": round(float(s.median()), 2),
                "std": round(float(s.std(ddof=0)), 2),
                "min": round(float(s.min()), 2),
                "max": round(float(s.max()), 2),
                "sum": round(float(s.sum()), 2),
                "outliers_iqr": outliers,
            }
        )
        if chart_id <= 4:
            bins = min(12, max(4, int(s.nunique() ** 0.5)))
            hist = pd.cut(s, bins=bins, include_lowest=True).value_counts(sort=False)
            labels = [str(idx) for idx in hist.index]
            peak_label = str(hist.idxmax()) if not hist.empty else "N/A"
            bucket_ranges = [
                {"left": float(interval.left), "right": float(interval.right), "label": str(interval)}
                for interval in hist.index
            ]
            outlier_rate = (outliers / int(s.count()) * 100.0) if int(s.count()) else 0.0
            skew_gap = abs(float(s.mean()) - float(s.median()))
            if outlier_rate == 0:
                meaning = "Meaning: the values are very consistent, so this metric looks stable across records."
            elif outlier_rate <= 5:
                meaning = "Meaning: most values are stable, with a small set of unusual cases worth checking."
            else:
                meaning = "Meaning: there are many unusual values, so this metric likely mixes very different cases."
            if skew_gap > (float(s.std(ddof=0)) * 0.5 if float(s.std(ddof=0)) > 0 else 0):
                meaning += " The gap between average and middle value suggests some large values are pulling the average up or down."
            charts.append(
                {
                    "id": f"dynamic_chart_{chart_id}",
                    "title": f"{labels_by_key.get(col, col)} distribution",
                    "type": "bar",
                    "labels": labels,
                    "datasets": [{"label": "Count", "data": [int(v) for v in hist.tolist()]}],
                    "drilldown": {"type": "numeric_bucket", "column": col, "bucket_ranges": bucket_ranges},
                    "interpretation": (
                        f"Most values are grouped around {peak_label}. "
                        f"The average is {float(s.mean()):,.2f}, the middle value is {float(s.median()):,.2f}, "
                        f"and {outliers} values are far from the normal range. "
                        f"{meaning}"
                    ),
                }
            )
            chart_id += 1

    for col in categorical_cols:
        s = df[col].astype(str).str.strip()
        s = s[s != ""]
        if s.empty:
            continue
        counts = s.value_counts().head(10)
        top_value = str(counts.index[0]) if len(counts) else ""
        top_count = int(counts.iloc[0]) if len(counts) else 0
        categorical_summary.append(
            {
                "label": labels_by_key.get(col, col),
                "unique_count": int(s.nunique()),
                "top_value": top_value,
                "top_count": top_count,
            }
        )
        if chart_id <= 8:
            dominant_share = (top_count / int(s.count()) * 100.0) if int(s.count()) else 0.0
            chart_kind = "pie" if len(counts) <= 8 else "bar"
            if dominant_share >= 70:
                meaning = "Meaning: one category dominates strongly, so most records are concentrated in a single group."
            elif dominant_share >= 40:
                meaning = "Meaning: one category leads, but there is still a visible mix of other categories."
            else:
                meaning = "Meaning: the data is spread across multiple categories, so there is no single dominant pattern."
            charts.append(
                {
                    "id": f"dynamic_chart_{chart_id}",
                    "title": f"Top values: {labels_by_key.get(col, col)}",
                    "type": chart_kind,
                    "labels": [str(v) for v in counts.index.tolist()],
                    "datasets": [{"label": "Frequency", "data": [int(v) for v in counts.tolist()]}],
                    "drilldown": {"type": "categorical", "column": col},
                    "interpretation": (
                        f"The leading category is '{top_value}' with {top_count} records "
                        f"({dominant_share:.1f}% of non-empty rows), indicating "
                        f"{'high' if dominant_share >= 50 else 'moderate'} concentration. "
                        f"{meaning}"
                    ),
                }
            )
            chart_id += 1

    for col in date_cols:
        raw_non_empty = df[col].astype(str).str.strip()
        raw_non_empty = raw_non_empty[raw_non_empty != ""]
        year_token_ratio = (
            float(raw_non_empty.str.match(r"^(19\d{2}|20\d{2})$").mean())
            if len(raw_non_empty)
            else 0.0
        )
        parsed = to_date_series(df[col]).dropna()
        if parsed.empty:
            continue
        min_date = parsed.min()
        max_date = parsed.max()
        span_days = int((max_date - min_date).days)
        is_year_only = year_token_ratio >= 0.8
        if is_year_only:
            min_year = int(min_date.year)
            max_year = int(max_date.year)
            span_label = "Year span"
            span_value = max(0, max_year - min_year)
            start_text = str(min_year)
            end_text = str(max_year)
        else:
            span_label = "Span (days)"
            span_value = span_days
            start_text = min_date.strftime("%Y-%m-%d")
            end_text = max_date.strftime("%Y-%m-%d")
        date_summary.append(
            {
                "label": labels_by_key.get(col, col),
                "start": start_text,
                "end": end_text,
                "span_days": span_days,
                "span_label": span_label,
                "span_value": span_value,
                "is_year_only": is_year_only,
            }
        )
        if chart_id <= 10:
            monthly = parsed.dt.to_period("M").astype(str).value_counts().sort_index().tail(18)
            trend_note = "stable"
            if len(monthly) >= 2:
                if int(monthly.iloc[-1]) > int(monthly.iloc[0]):
                    trend_note = "upward"
                elif int(monthly.iloc[-1]) < int(monthly.iloc[0]):
                    trend_note = "downward"
            if len(monthly) >= 2:
                change = int(monthly.iloc[-1]) - int(monthly.iloc[0])
                if trend_note == "upward":
                    meaning = f"Meaning: activity is increasing over time ({'+' if change >= 0 else ''}{change} from first to last visible month)."
                elif trend_note == "downward":
                    meaning = f"Meaning: activity is decreasing over time ({change} from first to last visible month)."
                else:
                    meaning = "Meaning: activity is fairly steady over time, with no strong rise or drop."
            else:
                meaning = "Meaning: there is not enough time-series variation yet to infer a trend."
            charts.append(
                {
                    "id": f"dynamic_chart_{chart_id}",
                    "title": f"Monthly trend: {labels_by_key.get(col, col)}",
                    "type": "line",
                    "labels": [str(v) for v in monthly.index.tolist()],
                    "datasets": [{"label": "Records", "data": [int(v) for v in monthly.tolist()]}],
                    "drilldown": {"type": "date_month", "column": col},
                    "interpretation": (
                        f"Records span from {start_text} to {end_text} "
                        f"({span_value} {'years' if is_year_only else 'days'}). Recent monthly movement appears {trend_note}. "
                        f"{meaning}"
                    ),
                }
            )
            chart_id += 1

    for col in boolean_cols:
        b = to_bool_series(df[col]).fillna(False)
        true_count = int(b.sum())
        false_count = int((~b).sum())
        boolean_summary.append(
            {
                "label": labels_by_key.get(col, col),
                "true_count": true_count,
                "false_count": false_count,
                "true_pct": round((true_count / len(b) * 100.0), 2) if len(b) else 0.0,
            }
        )

    def find_exact_chart_column(target: str) -> str | None:
        target_norm = normalize_label(target)
        for item in mapping:
            if normalize_label(item["label"]) == target_norm:
                return item["db_key"]
        for item in mapping:
            if str(item["db_key"]).strip().lower() == target.lower():
                return item["db_key"]
        return None

    status_col = find_exact_chart_column("status")
    if status_col and status_col in df.columns:
        raw_status = df[status_col].astype(str).str.strip()
        raw_status = raw_status[raw_status != ""]
        if not raw_status.empty:
            normalized_status = raw_status.str.lower().replace(
                {
                    "paid": "fully paid",
                    "partly paid": "partially paid",
                    "partial payment": "partially paid",
                    "unpaid": "no payment",
                    "not paid": "no payment",
                }
            )
            status_counts = normalized_status.value_counts().head(10)
            status_total = int(status_counts.sum())
            dominant_status = str(status_counts.index[0]) if len(status_counts) else "unknown"
            dominant_count = int(status_counts.iloc[0]) if len(status_counts) else 0
            dominant_pct = (dominant_count / status_total * 100.0) if status_total else 0.0
            status_label_map = {
                "fully paid": "Fully Paid",
                "partially paid": "Partially Paid",
                "no payment": "No Payment",
            }
            chart_kind = "pie" if len(status_counts) <= 8 else "bar"
            charts.append(
                {
                    "id": f"status_outcomes_{status_col}",
                    "title": f"Status outcomes: {labels_by_key.get(status_col, status_col)}",
                    "type": chart_kind,
                    "labels": [status_label_map.get(str(v), str(v).title()) for v in status_counts.index.tolist()],
                    "datasets": [{"label": "Records", "data": [int(v) for v in status_counts.tolist()]}],
                    "drilldown": {"type": "categorical", "column": status_col},
                    "interpretation": (
                        f"The most common status is '{status_label_map.get(dominant_status, dominant_status.title())}' "
                        f"with {dominant_count} records ({dominant_pct:.1f}% of rows that have status values). "
                        "Meaning: this chart shows payment completion mix and helps identify where follow-up should focus."
                    ),
                }
            )

    balance_col = find_exact_chart_column("balance")
    if balance_col and balance_col in df.columns:
        balance_series = pd.to_numeric(clean_numeric_text(df[balance_col]), errors="coerce").dropna()
        if not balance_series.empty:
            bins = min(10, max(4, int(balance_series.nunique() ** 0.5)))
            hist = pd.cut(balance_series, bins=bins, include_lowest=True).value_counts(sort=False)
            if not hist.empty:
                peak_label = str(hist.idxmax())
                bucket_ranges = [
                    {"left": float(interval.left), "right": float(interval.right), "label": str(interval)}
                    for interval in hist.index
                ]
                positive_total = float(balance_series[balance_series > 0].sum())
                settled_rows = int((balance_series <= 0).sum())
                charts.append(
                    {
                        "id": f"balance_distribution_{balance_col}",
                        "title": f"Balance distribution: {labels_by_key.get(balance_col, balance_col)}",
                        "type": "bar",
                        "labels": [str(idx) for idx in hist.index],
                        "datasets": [{"label": "Records", "data": [int(v) for v in hist.tolist()]}],
                        "drilldown": {"type": "numeric_bucket", "column": balance_col, "bucket_ranges": bucket_ranges},
                        "interpretation": (
                            f"The largest concentration of balances is in {peak_label}. "
                            f"Rows with zero or negative balance are {settled_rows}, while positive outstanding balance totals {positive_total:,.2f}. "
                            "Meaning: use this to prioritize high-balance buckets for recovery actions."
                        ),
                    }
                )

    best_profile = [row for row in profile if row["kind"] != "empty"]
    if best_profile:
        least_missing = min(best_profile, key=lambda row: row["missing_pct"])
        most_missing = max(best_profile, key=lambda row: row["missing_pct"])
        insights.append(
            f"Most complete field: {least_missing['label']} has values in {100 - least_missing['missing_pct']:.2f}% of records. "
            "Meaning: this is reliable information because it is filled in for almost everyone."
        )
        insights.append(
            f"Most missing field: {most_missing['label']} is blank in {most_missing['missing_pct']:.2f}% of records. "
            "Meaning: we should be careful using this field because there is not enough data to trust strong conclusions."
        )
    if numeric_summary:
        top_outlier = max(numeric_summary, key=lambda row: row["outliers_iqr"])
        if top_outlier["outliers_iqr"] > 0:
            insights.append(
                f"Possible unusual values: {top_outlier['label']} has {top_outlier['outliers_iqr']} values that are far from the normal range. "
                "Meaning: some entries may be errors, or they may be special cases that need manual review."
            )
    completeness_values = [100.0 - float(row["missing_pct"]) for row in profile]
    quality = build_data_quality_score(df, mapping, kinds, profile)
    return {
        "overview": {
            "rows": total_rows,
            "columns": total_columns,
            "overall_completeness_pct": round(
                float(sum(completeness_values) / len(completeness_values)), 2
            )
            if completeness_values
            else 0.0,
            "numeric_fields": len(numeric_cols),
            "categorical_fields": len(categorical_cols),
            "date_fields": len(date_cols),
            "boolean_fields": len(boolean_cols),
        },
        "quality": quality,
        "column_profile": profile,
        "insights": insights,
        "stats": {
            "numeric_summary": numeric_summary,
            "categorical_summary": categorical_summary,
            "date_summary": date_summary,
            "boolean_summary": boolean_summary,
            "correlations": [],
        },
        "charts": charts,
    }


def build_dashboard_payload(
    table_name: str,
    mapping: list[dict[str, str]],
    sheet_name_hint: str = "",
    snapshot_sheet_id: int | None = None,
    include_snapshot_delta: bool = True,
    thresholds: dict[str, float] | None = None,
) -> dict[str, object]:
    thresholds = normalize_thresholds(thresholds or {})
    db_keys = [item["db_key"] for item in mapping]
    labels_by_key = {item["db_key"]: item["label"] for item in mapping}
    table_ident = quote_identifier(table_name)

    with get_db() as conn:
        rows = conn.execute(
            f"SELECT {', '.join(quote_identifier(k) for k in db_keys)} FROM {table_ident}"
        ).fetchall()

    df = pd.DataFrame([{k: row[k] for k in db_keys} for row in rows], columns=db_keys)
    total_orgs = int(len(df))
    dynamic_dashboard = build_dynamic_dashboard(df, mapping)

    col_filing_period = (
        find_column(mapping, ["fil", "period"])
        or find_column(mapping, ["year"])
        or find_column(mapping, ["fy"])
    )
    col_fy_month = (
        find_column(mapping, ["financial", "year", "ending", "month"])
        or find_column(mapping, ["ending", "month"])
        or find_column(mapping, ["month"])
    )
    col_late = find_column(mapping, ["late", "return"])
    col_form14 = find_column(mapping, ["form", "14"])
    col_audit = find_column(mapping, ["audit", "report"])
    col_waiver = find_column(mapping, ["request", "waiver"])
    col_fee = find_column(mapping, ["filling", "fee"]) or find_column(mapping, ["filing", "fee"])
    col_penalty_paid = find_column(mapping, ["penalty", "paid"])
    col_penalty_outstanding = find_column(mapping, ["outstanding", "penalty"])
    col_total_penalty = find_column(mapping, ["total", "penalty"])
    col_requested_waiver = find_column(mapping, ["requested", "waiver"]) or find_column(mapping, ["request", "waiver"])
    col_last_income = find_column(mapping, ["last", "income"]) or find_column(mapping, ["income"])
    col_reason_waiver = find_column(mapping, ["reason", "waiver"])
    col_files_status = find_column(mapping, ["files", "status"]) or find_column(mapping, ["file", "status"])
    col_date_received = find_column(mapping, ["date", "received"])
    col_date_filed = find_column(mapping, ["date", "filed", "registry"])
    col_date_assigned = find_column(mapping, ["date", "assigned"])
    col_date_ack = find_column(mapping, ["acknoledged"]) or find_column(mapping, ["acknowledged"])
    col_end_notice = find_column(mapping, ["end", "notice", "period"])
    col_notice_countdown = find_column(mapping, ["notice", "countdown"])
    col_received_by = find_column(mapping, ["received", "by"])
    col_designate = find_column(mapping, ["designate", "pco"])
    col_org = find_column(mapping, ["organization", "name"])
    col_action = find_column(mapping, ["action"], ["filed", "registry"])
    has_form14 = col_form14 is not None
    has_audit = col_audit is not None
    has_waiver = col_waiver is not None or col_requested_waiver is not None

    late_bool = to_bool_series(df[col_late]) if col_late else pd.Series([False] * total_orgs)
    form14_bool = to_bool_series(df[col_form14]) if col_form14 else pd.Series([False] * total_orgs)
    audit_bool = to_bool_series(df[col_audit]) if col_audit else pd.Series([False] * total_orgs)
    waiver_bool = to_bool_series(df[col_waiver]) if col_waiver else pd.Series([False] * total_orgs)

    on_time_pct = (100.0 * (1 - late_bool.mean())) if total_orgs else 0.0
    form14_pct = (100.0 * form14_bool.mean()) if total_orgs else 0.0
    audit_pct = (100.0 * audit_bool.mean()) if total_orgs else 0.0
    waiver_freq = int(waiver_bool.sum())

    fully_compliant = (~late_bool) & form14_bool & audit_bool
    compliance_rate = (100.0 * fully_compliant.mean()) if total_orgs else 0.0
    late_rate = (100.0 * late_bool.mean()) if total_orgs else 0.0

    filing_fee = to_numeric_series(df[col_fee]) if col_fee else pd.Series([0.0] * total_orgs)
    penalty_paid = to_numeric_series(df[col_penalty_paid]) if col_penalty_paid else pd.Series([0.0] * total_orgs)
    penalty_out = to_numeric_series(df[col_penalty_outstanding]) if col_penalty_outstanding else pd.Series([0.0] * total_orgs)
    if (not col_penalty_outstanding) and col_total_penalty:
        total_penalty_series = to_numeric_series(df[col_total_penalty])
        penalty_out = (total_penalty_series - penalty_paid).clip(lower=0)
    requested_waiver_amount = (
        to_numeric_series(df[col_requested_waiver]) if col_requested_waiver else pd.Series([0.0] * total_orgs)
    )
    last_income_amount = (
        to_numeric_series(df[col_last_income]) if col_last_income else pd.Series([0.0] * total_orgs)
    )

    total_revenue = float(filing_fee.sum())
    total_penalty_paid = float(penalty_paid.sum())
    total_penalty_out = float(penalty_out.sum())
    late_count = int(late_bool.sum())
    avg_penalty_late = float(total_penalty_paid / late_count) if late_count else 0.0
    recovery_rate = (
        (total_penalty_paid / (total_penalty_paid + total_penalty_out)) * 100
        if (total_penalty_paid + total_penalty_out) > 0
        else 0.0
    )

    date_received = to_date_series(df[col_date_received]) if col_date_received else pd.Series(pd.NaT, index=df.index)
    date_filed = to_date_series(df[col_date_filed]) if col_date_filed else pd.Series(pd.NaT, index=df.index)
    date_assigned = to_date_series(df[col_date_assigned]) if col_date_assigned else pd.Series(pd.NaT, index=df.index)
    date_ack = to_date_series(df[col_date_ack]) if col_date_ack else pd.Series(pd.NaT, index=df.index)
    end_notice = to_date_series(df[col_end_notice]) if col_end_notice else pd.Series(pd.NaT, index=df.index)

    avg_process_days = (date_filed - date_received).dt.days.dropna().mean()
    avg_receipt_to_assign = (date_assigned - date_received).dt.days.dropna().mean()
    avg_assign_to_ack = (date_ack - date_assigned).dt.days.dropna().mean()

    notice_mask = date_ack.notna() & end_notice.notna()
    notice_compliance = (
        ((end_notice[notice_mask] >= date_ack[notice_mask]).mean() * 100) if notice_mask.any() else 0.0
    )

    stage_delays = {
        "Receipt to Assignment": float(avg_receipt_to_assign) if pd.notna(avg_receipt_to_assign) else 0.0,
        "Assignment to Acknowledgement": float(avg_assign_to_ack) if pd.notna(avg_assign_to_ack) else 0.0,
        "Receipt to Registry Filed": float(avg_process_days) if pd.notna(avg_process_days) else 0.0,
    }
    bottleneck_stage = max(stage_delays, key=stage_delays.get) if stage_delays else "N/A"

    officer_col = col_designate or col_received_by
    officer_series = (
        df[officer_col].astype(str).str.strip().replace("", "Unassigned")
        if officer_col
        else pd.Series(["Unassigned"] * total_orgs)
    )
    cases_per_officer = officer_series.value_counts().head(10)
    pending_mask = date_ack.isna()
    pending_per_officer = officer_series[pending_mask].value_counts().head(10)

    resolution_by_officer = pd.Series(dtype=float)
    if officer_col and col_date_assigned and col_date_ack:
        resolution = (date_ack - date_assigned).dt.days
        resolution_by_officer = (
            pd.DataFrame({"officer": officer_series, "days": resolution})
            .dropna()
            .groupby("officer")["days"]
            .mean()
            .sort_values(ascending=False)
            .head(10)
        )

    action_series = (
        df[col_action].astype(str).str.strip().replace("", "Unspecified").value_counts().head(10)
        if col_action
        else pd.Series(dtype=int)
    )

    org_series = (
        df[col_org].astype(str).str.strip().replace("", "Unknown Organization")
        if col_org
        else pd.Series(["Unknown Organization"] * total_orgs)
    )

    # Date Received quality + latest valid record (<= today)
    latest_by_date_received = None
    missing_date_received = int(date_received.isna().sum()) if col_date_received else total_orgs
    future_date_received_orgs: list[dict[str, object]] = []
    if col_date_received:
        today_date = pd.Timestamp.today().normalize()
        valid_received_mask = date_received.notna() & (date_received <= today_date)
        if valid_received_mask.any():
            latest_date = date_received[valid_received_mask].max()
            latest_idx = date_received[date_received == latest_date].index
            latest_org = (
                org_series.loc[latest_idx].astype(str).str.strip().replace("", "Unknown Organization").value_counts()
            )
            latest_by_date_received = {
                "date": latest_date.strftime("%Y-%m-%d"),
                "organization": latest_org.index[0] if len(latest_org) else "Unknown Organization",
                "count": int(len(latest_idx)),
            }

        future_mask_received = date_received.notna() & (date_received > today_date)
        if future_mask_received.any():
            bad = (
                org_series.loc[future_mask_received]
                .astype(str)
                .str.strip()
                .replace("", "Unknown Organization")
                .value_counts()
                .head(20)
            )
            future_date_received_orgs = [{"name": k, "count": int(v)} for k, v in bad.items()]
    repeat_late = org_series[late_bool].value_counts()
    high_risk_orgs = repeat_late[repeat_late > 1].head(10)

    notice_countdown = (
        to_numeric_series(df[col_notice_countdown]) if col_notice_countdown else pd.Series([9999.0] * total_orgs)
    )
    nearing_enforcement = int(((notice_countdown >= 0) & (notice_countdown <= 7)).sum())
    expired_notice = int((end_notice.dropna() < pd.Timestamp.today().normalize()).sum())
    waiver_by_org = org_series[waiver_bool].value_counts()
    waiver_abuse_orgs = waiver_by_org[waiver_by_org > 1].head(10)
    reason_dist = (
        df[col_reason_waiver].astype(str).str.strip().replace("", "Unspecified").value_counts().head(10)
        if col_reason_waiver
        else pd.Series(dtype=int)
    )
    file_status_dist = (
        df[col_files_status].astype(str).str.strip().replace("", "Unspecified").value_counts().head(10)
        if col_files_status
        else pd.Series(dtype=int)
    )

    # Yearly + monthly analysis: Filing Period year matched with FY ending month
    filing_period = df[col_filing_period].astype(str) if col_filing_period else pd.Series([""] * total_orgs)
    filing_year = filing_period.apply(extract_year_from_text)

    hint_month, hint_year = extract_month_year_hint(sheet_name_hint)
    if col_fy_month:
        fy_month = df[col_fy_month].astype(str).map(month_to_number)
    else:
        fy_month = pd.Series([hint_month] * total_orgs)
    if filing_year.isna().all() and hint_year:
        filing_year = pd.Series([hint_year] * total_orgs)
    trend_df = pd.DataFrame({"year": filing_year, "month": fy_month, "late": late_bool})
    trend_df = trend_df.dropna(subset=["year", "month"])
    trend_df["year"] = trend_df["year"].astype(int)
    trend_df["month"] = trend_df["month"].astype(int)

    trend_fin_df = trend_df.copy()
    trend_fin_df["filing_fee"] = filing_fee.loc[trend_fin_df.index].values if len(trend_fin_df) else []
    trend_fin_df["penalty_paid"] = penalty_paid.loc[trend_fin_df.index].values if len(trend_fin_df) else []
    trend_fin_df["penalty_outstanding"] = penalty_out.loc[trend_fin_df.index].values if len(trend_fin_df) else []
    trend_fin_df["organization"] = org_series.loc[trend_fin_df.index].values if len(trend_fin_df) else []

    today = datetime.today()
    current_year = today.year
    current_month = today.month
    valid_mask = (trend_fin_df["year"] < current_year) | (
        (trend_fin_df["year"] == current_year) & (trend_fin_df["month"] <= current_month)
    )
    future_mask = ~valid_mask

    ym_group = (
        trend_df.groupby(["year", "month"])
        .agg(filings=("late", "size"), late_rate=("late", "mean"))
        .reset_index()
        .sort_values(["year", "month"])
    )
    ym_labels = [f"{row.year}-{row.month:02d}" for row in ym_group.itertuples()]
    ym_filings = [int(v) for v in ym_group["filings"].tolist()]
    ym_late_rates = [round(float(v) * 100, 2) for v in ym_group["late_rate"].tolist()]

    month_group = (
        trend_df.groupby("month")
        .agg(filings=("late", "size"), late_rate=("late", "mean"))
        .reset_index()
        .sort_values("month")
    )
    month_labels = [datetime(2000, int(m), 1).strftime("%b") for m in month_group["month"].tolist()]
    month_filings = [int(v) for v in month_group["filings"].tolist()]
    month_late_rates = [round(float(v) * 100, 2) for v in month_group["late_rate"].tolist()]
    trend_normalized = build_trend_normalization(ym_group)
    cohort_summary = build_cohort_summary(org_series, late_bool, filing_year)
    anomaly_classification = build_anomaly_classification(df, mapping)

    ym_fin_group = (
        trend_fin_df.groupby(["year", "month"])
        .agg(
            filings=("late", "size"),
            late_count=("late", "sum"),
            filing_fee=("filing_fee", "sum"),
            penalty_paid=("penalty_paid", "sum"),
            penalty_outstanding=("penalty_outstanding", "sum"),
        )
        .reset_index()
        .sort_values(["year", "month"])
    )
    year_month_financial: list[dict[str, object]] = []
    for row in ym_fin_group.itertuples():
        paid = float(row.penalty_paid)
        out = float(row.penalty_outstanding)
        recovery = (paid / (paid + out) * 100.0) if (paid + out) > 0 else 0.0
        late_rate_ym = (float(row.late_count) / float(row.filings) * 100.0) if row.filings else 0.0
        year_month_financial.append(
            {
                "year": int(row.year),
                "month": int(row.month),
                "month_name": number_to_month_name(int(row.month)),
                "filings": int(row.filings),
                "late_count": int(row.late_count),
                "late_rate_pct": round(late_rate_ym, 2),
                "filing_fee": round(float(row.filing_fee), 2),
                "penalty_paid": round(paid, 2),
                "penalty_outstanding": round(out, 2),
                "recovery_rate_pct": round(recovery, 2),
            }
        )

    valid_trend_fin_df = trend_fin_df[valid_mask].copy()
    latest_valid_filing = None
    if not valid_trend_fin_df.empty:
        max_year = int(valid_trend_fin_df["year"].max())
        max_month = int(valid_trend_fin_df[valid_trend_fin_df["year"] == max_year]["month"].max())
        latest_slice = valid_trend_fin_df[
            (valid_trend_fin_df["year"] == max_year) & (valid_trend_fin_df["month"] == max_month)
        ]
        org_counts = (
            latest_slice["organization"].astype(str).str.strip().replace("", "Unknown Organization").value_counts()
        )
        latest_org = org_counts.index[0] if len(org_counts) else "Unknown Organization"
        latest_valid_filing = {
            "year": max_year,
            "month": max_month,
            "month_name": number_to_month_name(max_month),
            "filings": int(len(latest_slice)),
            "organization": latest_org,
        }

    future_orgs = []
    if future_mask.any():
        flagged = (
            trend_fin_df[future_mask]["organization"]
            .astype(str)
            .str.strip()
            .replace("", "Unknown Organization")
            .value_counts()
            .head(20)
        )
        future_orgs = [{"name": k, "count": int(v)} for k, v in flagged.items()]

    executive_summary = {
        "total_organizations": total_orgs,
        "compliance_rate": fmt_pct(compliance_rate),
        "revenue_collected": fmt_num(total_revenue + total_penalty_paid),
        "outstanding_penalties": fmt_num(total_penalty_out),
        "avg_processing_time_days": round(float(avg_process_days), 2) if pd.notna(avg_process_days) else 0.0,
    }

    ai_interpretations = {
        "executive": (
            f"{fmt_pct(compliance_rate)} complied fully while {fmt_pct(late_rate)} filed late. "
            f"Collected revenue is {fmt_num(total_revenue + total_penalty_paid)} with "
            f"{fmt_num(total_penalty_out)} still outstanding."
        ),
        "compliance": "",
        "financial": (
            f"Penalty recovery rate is {fmt_pct(recovery_rate)}. Average penalty per late filer is "
            f"{fmt_num(avg_penalty_late)}; this indicates {'strong' if recovery_rate >= 70 else 'weak'} recovery."
        ),
        "registry": (
            f"Average processing time is {round(float(avg_process_days),2) if pd.notna(avg_process_days) else 0.0} days. "
            f"The main bottleneck is {bottleneck_stage}."
        ),
        "risk": (
            f"{len(high_risk_orgs)} organizations are repeat late filers. "
            f"{nearing_enforcement} cases are near enforcement and {expired_notice} notices have expired."
        ),
        "waiver": (
            f"Total requested waiver value is {fmt_num(float(requested_waiver_amount.sum()))}. "
            f"Top waiver reasons and file statuses are included for {sheet_name_hint or 'this sheet'}."
        ),
    }

    compliance_items: list[dict[str, object]] = [
        {"label": "% on-time filings", "value": round(on_time_pct, 2), "is_percent": True},
        {"label": "Late filing rate", "value": round(late_rate, 2), "is_percent": True},
    ]
    if has_form14:
        compliance_items.append({"label": "% Form 14 submissions", "value": round(form14_pct, 2), "is_percent": True})
    if has_audit:
        compliance_items.append({"label": "% Audit Report submissions", "value": round(audit_pct, 2), "is_percent": True})
    if has_waiver:
        compliance_items.append({"label": "Waiver request frequency", "value": waiver_freq, "is_percent": False})

    if has_form14 and has_audit:
        compliance_statement = f"“{fmt_pct(compliance_rate)} complied fully, while {fmt_pct(late_rate)} filed late.”"
        ai_interpretations["compliance"] = (
            f"On-time filings are {fmt_pct(on_time_pct)}. Form 14 submission is {fmt_pct(form14_pct)}, "
            f"audit report submission is {fmt_pct(audit_pct)}, and waiver requests appear in {waiver_freq} cases."
        )
    else:
        compliance_statement = (
            f"“This sheet is waiver-focused. {fmt_pct(on_time_pct)} are on time and "
            f"{fmt_pct(late_rate)} are late across available records.”"
        )
        ai_interpretations["compliance"] = (
            f"Adaptive compliance view: on-time filings are {fmt_pct(on_time_pct)} and late filings are {fmt_pct(late_rate)}. "
            f"{'Waiver requests are tracked in this sheet.' if has_waiver else 'Waiver request fields are not available in this sheet.'}"
        )

    payload = {
        "columns": labels_by_key,
        "metrics": {
            "on_time_pct": round(on_time_pct, 2),
            "form14_pct": round(form14_pct, 2),
            "audit_pct": round(audit_pct, 2),
            "waiver_freq": waiver_freq,
            "total_revenue": round(total_revenue, 2),
            "total_penalty_paid": round(total_penalty_paid, 2),
            "total_penalty_outstanding": round(total_penalty_out, 2),
            "avg_penalty_per_late": round(avg_penalty_late, 2),
            "recovery_rate_pct": round(recovery_rate, 2),
            "late_filing_rate_pct": round(late_rate, 2),
            "total_requested_waiver": round(float(requested_waiver_amount.sum()), 2),
            "total_last_income": round(float(last_income_amount.sum()), 2),
            "avg_processing_days": round(float(avg_process_days), 2) if pd.notna(avg_process_days) else 0.0,
            "avg_receipt_to_assignment_days": round(float(avg_receipt_to_assign), 2)
            if pd.notna(avg_receipt_to_assign)
            else 0.0,
            "avg_assignment_to_ack_days": round(float(avg_assign_to_ack), 2) if pd.notna(avg_assign_to_ack) else 0.0,
            "notice_compliance_rate_pct": round(float(notice_compliance), 2),
            "bottleneck_stage": bottleneck_stage,
            "nearing_enforcement_count": nearing_enforcement,
            "expired_notice_count": expired_notice,
            "repeat_late_org_count": int(len(high_risk_orgs)),
            "sheet_hint_month": hint_month or 0,
            "sheet_hint_year": hint_year or 0,
        },
        "charts": {
            "year_month_labels": ym_labels,
            "year_month_filings": ym_filings,
            "year_month_late_rates": ym_late_rates,
            "month_labels": month_labels,
            "month_filings": month_filings,
            "month_late_rates": month_late_rates,
            "action_labels": action_series.index.tolist(),
            "action_values": [int(v) for v in action_series.tolist()],
            "reason_labels": reason_dist.index.tolist(),
            "reason_values": [int(v) for v in reason_dist.tolist()],
            "file_status_labels": file_status_dist.index.tolist(),
            "file_status_values": [int(v) for v in file_status_dist.tolist()],
            "officer_labels": cases_per_officer.index.tolist(),
            "officer_values": [int(v) for v in cases_per_officer.tolist()],
            "pending_labels": pending_per_officer.index.tolist(),
            "pending_values": [int(v) for v in pending_per_officer.tolist()],
            "revenue_labels": ["Filing Fee", "Penalty Paid", "Outstanding Penalty"],
            "revenue_values": [round(total_revenue, 2), round(total_penalty_paid, 2), round(total_penalty_out, 2)],
        },
        "tables": {
            "high_risk_orgs": [{"name": k, "count": int(v)} for k, v in high_risk_orgs.items()],
            "waiver_abuse_orgs": [{"name": k, "count": int(v)} for k, v in waiver_abuse_orgs.items()],
            "waiver_reasons": [{"name": k, "count": int(v)} for k, v in reason_dist.items()],
            "files_status": [{"name": k, "count": int(v)} for k, v in file_status_dist.items()],
            "resolution_by_officer": [
                {"name": k, "days": round(float(v), 2)} for k, v in resolution_by_officer.items()
            ],
            "year_month_financial": year_month_financial,
            "latest_valid_filing": latest_valid_filing,
            "future_date_orgs": future_orgs,
            "latest_by_date_received": latest_by_date_received,
            "future_date_received_orgs": future_date_received_orgs,
        },
        "adaptive": {
            "has_form14": has_form14,
            "has_audit": has_audit,
            "has_waiver": has_waiver,
            "has_date_received": col_date_received is not None,
            "compliance_items": compliance_items,
            "compliance_statement": compliance_statement,
        },
        "executive_summary": executive_summary,
        "ai_interpretations": ai_interpretations,
        "dynamic": dynamic_dashboard,
        "trend_normalized": trend_normalized,
        "cohort_summary": cohort_summary,
        "anomaly_classification": anomaly_classification,
        "kpi_targets": [],
        "snapshot_delta": (
            build_snapshot_delta_for_uploaded_sheet(snapshot_sheet_id, {
                "late_filing_rate_pct": round(late_rate, 2),
                "total_penalty_paid": round(total_penalty_paid, 2),
                "total_penalty_outstanding": round(total_penalty_out, 2),
                "recovery_rate_pct": round(recovery_rate, 2),
                "total_revenue": round(total_revenue, 2),
                "total_requested_waiver": round(float(requested_waiver_amount.sum()), 2),
            })
            if include_snapshot_delta and snapshot_sheet_id is not None
            else {"available": False, "reason": "Snapshot comparison available for uploaded sheets only."}
        ),
        "snapshot_delta_interpretation": "",
        "management_templates": {},
        "alerts": [],
        "thresholds": thresholds,
        "kpi_target_variance_interpretation": "",
        "trend_normalized_interpretation": "",
        "anomaly_classification_interpretation": "",
        "compliance_numeric_interpretation": "",
        "categorical_statistics_interpretation": "",
        "submission_summary": "",
    }
    payload["kpi_targets"] = build_kpi_target_variance(payload["metrics"], thresholds)
    payload["management_templates"] = build_management_templates(
        sheet_name_hint or table_name,
        payload["metrics"],
        dynamic_dashboard.get("quality", {}) if isinstance(dynamic_dashboard, dict) else {},
    )
    payload["alerts"] = build_alerts(
        payload["metrics"],
        dynamic_dashboard.get("quality", {}) if isinstance(dynamic_dashboard, dict) else {},
        payload["kpi_targets"],
        thresholds,
    )
    payload["kpi_target_variance_interpretation"] = build_kpi_target_variance_interpretation(payload["kpi_targets"])
    payload["trend_normalized_interpretation"] = build_trend_normalization_interpretation(payload["trend_normalized"])
    payload["anomaly_classification_interpretation"] = build_anomaly_classification_interpretation(
        payload["anomaly_classification"]
    )
    payload["compliance_numeric_interpretation"] = build_compliance_numeric_interpretation(
        dynamic_dashboard.get("stats", {}).get("numeric_summary", []) if isinstance(dynamic_dashboard, dict) else [],
        payload["metrics"],
    )
    payload["categorical_statistics_interpretation"] = build_categorical_statistics_interpretation(
        dynamic_dashboard.get("stats", {}).get("categorical_summary", []) if isinstance(dynamic_dashboard, dict) else []
    )
    payload["snapshot_delta_interpretation"] = build_snapshot_delta_interpretation(payload.get("snapshot_delta", {}))
    payload["kpi_target_variance_interpretation"] = paragraphize_ai_text(
        payload["kpi_target_variance_interpretation"], 2
    )
    payload["trend_normalized_interpretation"] = paragraphize_ai_text(
        payload["trend_normalized_interpretation"], 2
    )
    payload["anomaly_classification_interpretation"] = paragraphize_ai_text(
        payload["anomaly_classification_interpretation"], 2
    )
    payload["compliance_numeric_interpretation"] = paragraphize_ai_text(
        payload["compliance_numeric_interpretation"], 2
    )
    payload["categorical_statistics_interpretation"] = paragraphize_ai_text(
        payload["categorical_statistics_interpretation"], 2
    )
    payload["snapshot_delta_interpretation"] = paragraphize_ai_text(payload["snapshot_delta_interpretation"], 2)
    charts = payload.get("dynamic", {}).get("charts", []) if isinstance(payload.get("dynamic"), dict) else []
    if isinstance(charts, list):
        for chart in charts:
            if isinstance(chart, dict) and chart.get("interpretation"):
                chart["interpretation"] = paragraphize_ai_text(str(chart.get("interpretation", "")), 2)
    ai_blocks = payload.get("ai_interpretations", {})
    if isinstance(ai_blocks, dict):
        for key, val in list(ai_blocks.items()):
            ai_blocks[key] = paragraphize_ai_text(str(val), 2)
    payload["submission_summary"] = build_submission_summary(
        sheet_name_hint=sheet_name_hint,
        total_orgs=total_orgs,
        compliance_rate=compliance_rate,
        late_rate=late_rate,
        total_revenue=total_revenue,
        total_penalty_paid=total_penalty_paid,
        total_penalty_out=total_penalty_out,
        recovery_rate=recovery_rate,
        avg_process_days=round(float(avg_process_days), 2) if pd.notna(avg_process_days) else 0.0,
        dynamic_dashboard=dynamic_dashboard,
        kpi_targets=payload["kpi_targets"],
        kpi_target_variance_interpretation=payload["kpi_target_variance_interpretation"],
        snapshot_delta=payload.get("snapshot_delta", {}),
        snapshot_delta_interpretation=payload["snapshot_delta_interpretation"],
    )
    return payload


def get_cached_dashboard_payload(
    table_name: str,
    mapping: list[dict[str, str]],
    sheet_name_hint: str = "",
    snapshot_sheet_id: int | None = None,
    include_snapshot_delta: bool = True,
    thresholds: dict[str, float] | None = None,
) -> dict[str, object]:
    thresholds_key = json.dumps(thresholds or {}, sort_keys=True, ensure_ascii=True)
    cache_key = (
        _data_version,
        table_name,
        _mapping_cache_fingerprint(mapping),
        str(sheet_name_hint or ""),
        int(snapshot_sheet_id) if snapshot_sheet_id is not None else None,
        bool(include_snapshot_delta),
        thresholds_key,
    )
    cached = _cache_get(_dashboard_payload_cache, cache_key)
    if cached is not None:
        return cached
    payload = build_dashboard_payload(
        table_name,
        mapping,
        sheet_name_hint=sheet_name_hint,
        snapshot_sheet_id=snapshot_sheet_id,
        include_snapshot_delta=include_snapshot_delta,
        thresholds=thresholds,
    )
    _cache_set(_dashboard_payload_cache, cache_key, DASHBOARD_CACHE_TTL_SECONDS, payload)
    return deepcopy(payload)


def extract_month_year_from_question(question: str) -> tuple[int | None, int | None]:
    q = question.lower()
    year_match = re.search(r"\b(19\d{2}|20\d{2})\b", q)
    year = int(year_match.group(1)) if year_match else None

    months = {
        "january": 1,
        "february": 2,
        "march": 3,
        "april": 4,
        "may": 5,
        "june": 6,
        "july": 7,
        "august": 8,
        "september": 9,
        "october": 10,
        "november": 11,
        "december": 12,
        "jan": 1,
        "feb": 2,
        "mar": 3,
        "apr": 4,
        "jun": 6,
        "jul": 7,
        "aug": 8,
        "sep": 9,
        "oct": 10,
        "nov": 11,
        "dec": 12,
    }
    month = None
    for key, value in months.items():
        if re.search(rf"\b{re.escape(key)}\b", q):
            month = value
            break
    return month, year


def schema_name_match_report(columns_map: dict[str, object]) -> dict[str, object]:
    if not isinstance(columns_map, dict):
        return {"filing_period": None, "fy_month": None, "score": 0}
    labels = [str(v) for v in columns_map.values()]
    best_filing = None
    best_fy = None
    filing_score = -1
    fy_score = -1
    for label in labels:
        n = normalize_label(label)
        score_filing = 0
        if "fil" in n:
            score_filing += 2
        if "period" in n:
            score_filing += 2
        if "year" in n:
            score_filing += 1
        if score_filing > filing_score:
            filing_score = score_filing
            best_filing = label

        score_fy = 0
        if "financial" in n:
            score_fy += 2
        if "ending" in n:
            score_fy += 1
        if "month" in n:
            score_fy += 2
        if score_fy > fy_score:
            fy_score = score_fy
            best_fy = label
    score = max(0, filing_score) + max(0, fy_score)
    if filing_score < 3:
        best_filing = None
    if fy_score < 3:
        best_fy = None
    return {"filing_period": best_filing, "fy_month": best_fy, "score": score}


def nearest_period_row(
    ym: list[dict[str, object]],
    month: int,
    year: int,
) -> tuple[dict[str, object] | None, str]:
    if not ym:
        return None, "none"
    same_month = [r for r in ym if int(r.get("month", 0) or 0) == month]
    if same_month:
        nearest = min(same_month, key=lambda r: abs(int(r.get("year", 0) or 0) - year))
        return nearest, "same_month"
    nearest = min(
        ym,
        key=lambda r: abs(int(r.get("year", 0) or 0) - year) * 12 + abs(int(r.get("month", 0) or 0) - month),
    )
    return nearest, "nearest_period"


def cross_sheet_period_answer_for_uploaded_file(
    file_id: int, month: int, year: int, dataset_title: str
) -> str | None:
    with get_db() as conn:
        sheets = conn.execute(
            """
            SELECT id, sheet_name, table_name, mapping_json
            FROM uploaded_sheets
            WHERE file_id = ?
            ORDER BY id ASC
            """,
            (file_id,),
        ).fetchall()
    if not sheets:
        return None

    timeline: list[dict[str, object]] = []
    for row in sheets:
        period = parse_sheet_period_hint(str(row["sheet_name"]))
        mapping = json.loads(row["mapping_json"])
        payload = get_cached_dashboard_payload(
            str(row["table_name"]),
            mapping,
            sheet_name_hint=str(row["sheet_name"]),
            include_snapshot_delta=False,
        )
        m = payload.get("metrics", {}) if isinstance(payload, dict) else {}
        timeline.append(
            {
                "sheet": str(row["sheet_name"]),
                "year": int(period.get("year", 0) or 0),
                "month": int(period.get("start_month", 0) or 0),
                "period_label": str(period.get("label", row["sheet_name"])),
                "late_rate": float(m.get("late_filing_rate_pct", 0) or 0.0),
                "recovery": float(m.get("recovery_rate_pct", 0) or 0.0),
                "filing_fee": float(m.get("total_revenue", 0) or 0.0),
                "penalty_paid": float(m.get("total_penalty_paid", 0) or 0.0),
                "outstanding": float(m.get("total_penalty_outstanding", 0) or 0.0),
                "requested_waiver": float(m.get("total_requested_waiver", 0) or 0.0),
                "last_income": float(m.get("total_last_income", 0) or 0.0),
            }
        )
    timeline = [t for t in timeline if t["month"] > 0 and t["year"] > 0]
    if not timeline:
        return None
    timeline.sort(key=lambda t: (t["year"], t["month"]))

    exact = next((t for t in timeline if t["year"] == year and t["month"] == month), None)
    if exact:
        return (
            f"Using cross-sheet timeline for {dataset_title}: exact match found at {exact['period_label']} "
            f"(sheet '{exact['sheet']}'). Filing fee {exact['filing_fee']:.2f}, penalty paid {exact['penalty_paid']:.2f}, "
            f"outstanding penalty {exact['outstanding']:.2f}, recovery rate {exact['recovery']:.2f}%, "
            f"late filing rate {exact['late_rate']:.2f}%, requested waiver {exact['requested_waiver']:.2f}, "
            f"last income {exact['last_income']:.2f}."
        )

    nearest = min(timeline, key=lambda t: abs(t["year"] - year) * 12 + abs(t["month"] - month))
    return (
        f"No exact records for {number_to_month_name(month)} {year} in {dataset_title}. "
        f"I retried with workbook-level chronological sheet matching and selected nearest period {nearest['period_label']} "
        f"(sheet '{nearest['sheet']}'). Filing fee {nearest['filing_fee']:.2f}, penalty paid {nearest['penalty_paid']:.2f}, "
        f"outstanding penalty {nearest['outstanding']:.2f}, recovery rate {nearest['recovery']:.2f}%, "
        f"late filing rate {nearest['late_rate']:.2f}%, requested waiver {nearest['requested_waiver']:.2f}, "
        f"last income {nearest['last_income']:.2f}."
    )


def format_top_waiver_reasons(tables: dict[str, object], limit: int = 3) -> str:
    reasons = tables.get("waiver_reasons", []) if isinstance(tables, dict) else []
    if not isinstance(reasons, list) or not reasons:
        return "No waiver reasons were captured in this sheet."
    parts = []
    for row in reasons[:limit]:
        try:
            name = str(row.get("name", "Unspecified")).strip() or "Unspecified"
            count = int(row.get("count", 0))
            parts.append(f"{name} ({count})")
        except Exception:
            continue
    return ", ".join(parts) if parts else "No waiver reasons were captured in this sheet."


def build_printable_report(
    dataset_title: str,
    metrics: dict[str, object],
    tables: dict[str, object],
    month: int | None = None,
    year: int | None = None,
    summary_words: int = 500,
) -> str:
    ym = tables.get("year_month_financial", []) if isinstance(tables, dict) else []
    ym = ym if isinstance(ym, list) else []
    period_text = "Current sheet snapshot"
    filings = late_count = 0
    late_rate = filing_fee = penalty_paid = penalty_outstanding = recovery_rate = 0.0
    period_note = ""

    if month and year:
        exact = next(
            (row for row in ym if int(row.get("year", 0) or 0) == year and int(row.get("month", 0) or 0) == month),
            None,
        )
        if exact:
            period_text = f"{number_to_month_name(month)} {year}"
            filings = int(exact.get("filings", 0) or 0)
            late_count = int(exact.get("late_count", 0) or 0)
            late_rate = float(exact.get("late_rate_pct", 0) or 0.0)
            filing_fee = float(exact.get("filing_fee", 0) or 0.0)
            penalty_paid = float(exact.get("penalty_paid", 0) or 0.0)
            penalty_outstanding = float(exact.get("penalty_outstanding", 0) or 0.0)
            recovery_rate = float(exact.get("recovery_rate_pct", 0) or 0.0)
        else:
            month_rows = [row for row in ym if int(row.get("month", 0) or 0) == month]
            if month_rows:
                period_text = f"{number_to_month_name(month)} {year} request (historical month-only view)"
                filings = sum(int(row.get("filings", 0) or 0) for row in month_rows)
                late_count = sum(int(row.get("late_count", 0) or 0) for row in month_rows)
                filing_fee = sum(float(row.get("filing_fee", 0) or 0.0) for row in month_rows)
                penalty_paid = sum(float(row.get("penalty_paid", 0) or 0.0) for row in month_rows)
                penalty_outstanding = sum(float(row.get("penalty_outstanding", 0) or 0.0) for row in month_rows)
                late_rate = (late_count / filings * 100.0) if filings else 0.0
                recovery_rate = (
                    penalty_paid / (penalty_paid + penalty_outstanding) * 100.0
                    if (penalty_paid + penalty_outstanding) > 0
                    else 0.0
                )
                years = sorted({int(row.get("year", 0) or 0) for row in month_rows if int(row.get("year", 0) or 0) > 0})
                if years:
                    period_note = (
                        f"No exact records were found for {number_to_month_name(month)} {year}. "
                        f"Available month data spans years {years[0]} to {years[-1]}."
                    )
            else:
                period_text = f"{number_to_month_name(month)} {year}"
                period_note = "No matching records were found for this month and year."
    elif ym:
        latest = max(ym, key=lambda row: (int(row.get("year", 0) or 0), int(row.get("month", 0) or 0)))
        period_text = f"Latest available period ({latest.get('month_name')} {latest.get('year')})"
        filings = int(latest.get("filings", 0) or 0)
        late_count = int(latest.get("late_count", 0) or 0)
        late_rate = float(latest.get("late_rate_pct", 0) or 0.0)
        filing_fee = float(latest.get("filing_fee", 0) or 0.0)
        penalty_paid = float(latest.get("penalty_paid", 0) or 0.0)
        penalty_outstanding = float(latest.get("penalty_outstanding", 0) or 0.0)
        recovery_rate = float(latest.get("recovery_rate_pct", 0) or 0.0)

    total_last_income = float(metrics.get("total_last_income", 0) or 0.0)
    total_requested_waiver = float(metrics.get("total_requested_waiver", 0) or 0.0)
    reasons_text = format_top_waiver_reasons(tables, limit=4)
    kpi_summary = (
        f"This report provides a focused management interpretation of the key KPIs for {period_text}. "
        f"Total filings are {filings}, and all {late_count} were late, producing a late-filing rate of {late_rate:.2f}%. "
        "This indicates the current filing cycle has full timeliness non-compliance and requires immediate enforcement follow-up, "
        "including targeted reminders and tighter escalation timelines. Filing fee performance shows "
        f"{filing_fee:,.2f} collected, while penalty collections stand at {penalty_paid:,.2f}. "
        f"Outstanding penalties remain high at {penalty_outstanding:,.2f}, which keeps recovery at only {recovery_rate:.2f}%. "
        "From a cash-realization perspective, the gap between paid and outstanding penalties is the primary revenue risk and "
        "should be tracked weekly until collection conversion improves. Last income totals "
        f"{total_last_income:,.2f}, which gives context on the economic scale of organizations under review and can inform "
        "proportional compliance actions. Requested waiver value is "
        f"{total_requested_waiver:,.2f}, signaling significant relief demand that may materially reduce collectible penalties "
        "if approvals are not tightly controlled. The top recorded reasons for waiver are "
        f"{reasons_text}. These reasons point to recurring financial and administrative constraints, suggesting the need for a "
        "consistent waiver decision framework with documented evidence thresholds and post-waiver compliance conditions. "
        "Overall, this sheet indicates a high-risk compliance posture, weak recovery efficiency, and concentrated waiver pressure. "
        "Management should prioritize three actions: clear the late-return backlog through time-bound case ownership, intensify "
        "collection on high-value outstanding penalties, and standardize waiver governance to protect fairness, transparency, and "
        "future compliance behavior."
    )

    report = (
        f"Monthly Work Report for Management\n"
        f"Dataset: {dataset_title}\n"
        f"Reporting period: {period_text}.\n\n"
        f"The dataset shows {filings} filings, with {late_count} late returns ({late_rate:.2f}%). "
        f"Filing fee recorded is {filing_fee:,.2f}, penalty paid is {penalty_paid:,.2f}, and outstanding penalty is "
        f"{penalty_outstanding:,.2f}, giving a recovery rate of {recovery_rate:.2f}%. "
        f"Across the full sheet, cumulative Last income is {total_last_income:,.2f} and total Requested waiver is "
        f"{total_requested_waiver:,.2f}. The leading REASONS FOR WAIVER are {reasons_text}. "
        f"Overall late filing rate in this dataset is {float(metrics.get('late_filing_rate_pct', 0) or 0.0):.2f}% "
        f"with penalty recovery at {float(metrics.get('recovery_rate_pct', 0) or 0.0):.2f}%. "
        "Priority actions are to reduce the late-return backlog, fast-track high-value outstanding penalty cases, and "
        "review recurring waiver justifications for consistency and policy compliance."
    )
    report += f"\n\n{summary_words}-Word KPI Summary\n{kpi_summary}"
    if period_note:
        report += f"\n\nNote: {period_note}"
    return report


def is_management_summary_query(question: str) -> bool:
    q = (question or "").lower()
    return any(
        token in q
        for token in [
            "summary",
            "monthly work report",
            "report",
            "management",
            "print",
            "print out",
            "300 word",
            "500 word",
            "words",
            "last income",
            "requested waiver",
            "reasons for waiver",
            "reason for waiver",
        ]
    )


def local_dataset_answer(dataset_title: str, payload: dict[str, object], question: str) -> str | None:
    q = question.lower().strip()
    metrics = payload.get("metrics", {}) if isinstance(payload, dict) else {}
    tables = payload.get("tables", {}) if isinstance(payload, dict) else {}
    columns_map = payload.get("columns", {}) if isinstance(payload, dict) else {}
    context = payload.get("context", {}) if isinstance(payload, dict) else {}
    ym = tables.get("year_month_financial", []) if isinstance(tables, dict) else []
    if not isinstance(ym, list):
        ym = []
    latest_valid = tables.get("latest_valid_filing") if isinstance(tables, dict) else None
    future_orgs = tables.get("future_date_orgs", []) if isinstance(tables, dict) else []
    latest_received = tables.get("latest_by_date_received") if isinstance(tables, dict) else None
    future_received_orgs = tables.get("future_date_received_orgs", []) if isinstance(tables, dict) else []
    adaptive = payload.get("adaptive", {}) if isinstance(payload, dict) else {}
    has_date_received = bool(adaptive.get("has_date_received")) if isinstance(adaptive, dict) else False
    if not isinstance(future_orgs, list):
        future_orgs = []
    if not isinstance(future_received_orgs, list):
        future_received_orgs = []

    month, year = extract_month_year_from_question(question)
    if not month or not year:
        title_month, title_year = extract_month_year_hint(dataset_title)
        month = month or title_month
        year = year or title_year
    summary_requested = any(
        token in q
        for token in [
            "summary",
            "monthly work report",
            "report",
            "management",
            "print",
            "print out",
            "200 word",
        ]
    )
    extra_fields_requested = any(
        token in q for token in ["last income", "requested waiver", "reasons for waiver", "reason for waiver"]
    )
    if summary_requested or extra_fields_requested:
        return build_printable_report(dataset_title, metrics, tables, month=month, year=year)

    if "date received" in q:
        if not has_date_received:
            return (
                f"Date Received column is not available in {dataset_title}. "
                "I used the available filing-period fields instead."
            )
        if latest_received:
            ans = (
                f"Using Date Received in {dataset_title}: latest valid received date (<= today) is "
                f"{latest_received['date']}. Organization: {latest_received['organization']} "
                f"(records at this latest date: {latest_received['count']})."
            )
        else:
            ans = f"Using Date Received in {dataset_title}: no valid received dates found up to today."
        if future_received_orgs:
            preview = ", ".join(item["name"] for item in future_received_orgs[:5])
            ans += (
                f" Flagged organizations with future Date Received values: {preview}"
                + ("." if len(future_received_orgs) <= 5 else " (and others).")
            )
        return ans

    if ("most recent" in q or "last" in q or "latest" in q) and ("filing" in q or "annual return" in q):
        if not latest_valid:
            return "No valid filing period up to today is available for this sheet."
        base = (
            f"Most recent valid filing period in {dataset_title} is {latest_valid['month_name']} "
            f"{latest_valid['year']} (capped at today). "
            f"Filings: {latest_valid['filings']}. "
            f"Company/organization: {latest_valid['organization']}."
        )
        if future_orgs:
            preview = ", ".join(item["name"] for item in future_orgs[:5])
            base += (
                f" Flagged wrong-data organizations with future filing dates: {preview}"
                + ("." if len(future_orgs) <= 5 else " (and others).")
            )
        return base

    if month and year:
        matched = next((row for row in ym if row.get("year") == year and row.get("month") == month), None)
        if not matched:
            hint_month = int(metrics.get("sheet_hint_month", 0) or 0)
            hint_year = int(metrics.get("sheet_hint_year", 0) or 0)
            if hint_month == month and hint_year == year and ym:
                month_rows = [r for r in ym if int(r.get("month", 0) or 0) == month]
                if not month_rows:
                    return (
                        f"No matched records for {number_to_month_name(month)} {year} in {dataset_title}. "
                        "The sheet name encodes this month/year, but row-level filing period values do not match it."
                    )
                filings = sum(int(r.get("filings", 0) or 0) for r in month_rows)
                late_count = sum(int(r.get("late_count", 0) or 0) for r in month_rows)
                filing_fee = sum(float(r.get("filing_fee", 0) or 0.0) for r in month_rows)
                penalty_paid = sum(float(r.get("penalty_paid", 0) or 0.0) for r in month_rows)
                penalty_out = sum(float(r.get("penalty_outstanding", 0) or 0.0) for r in month_rows)
                recovery = (penalty_paid / (penalty_paid + penalty_out) * 100.0) if (penalty_paid + penalty_out) > 0 else 0.0
                late_rate = (late_count / filings * 100.0) if filings else 0.0
                years = sorted({int(r.get("year", 0) or 0) for r in month_rows if int(r.get("year", 0) or 0) > 0})
                reason_text = format_top_waiver_reasons(tables, limit=3)
                return (
                    f"{number_to_month_name(month)} {year} requested snapshot for {dataset_title}: "
                    f"no exact {year} filing-period rows were found. Historical {number_to_month_name(month)} rows "
                    f"span {years[0]} to {years[-1] if years else 'N/A'}. "
                    f"Combined figures: Filings {filings}, Late filings {late_count} ({late_rate:.2f}%), "
                    f"Filing fee {filing_fee:.2f}, Penalty paid {penalty_paid:.2f}, "
                    f"Outstanding penalty {penalty_out:.2f}, Recovery rate {recovery:.2f}%, "
                    f"Last income total {float(metrics.get('total_last_income', 0) or 0.0):.2f}, "
                    f"Requested waiver total {float(metrics.get('total_requested_waiver', 0) or 0.0):.2f}. "
                    f"Top REASONS FOR WAIVER: {reason_text}."
                )
            schema_match = schema_name_match_report(columns_map if isinstance(columns_map, dict) else {})
            nearest, mode = nearest_period_row(ym, month, year)
            if nearest:
                reason_text = format_top_waiver_reasons(tables, limit=3)
                nearest_month = int(nearest.get("month", 0) or 0)
                nearest_year = int(nearest.get("year", 0) or 0)
                period_label = f"{number_to_month_name(nearest_month)} {nearest_year}"
                mode_text = (
                    "same month with nearest available year"
                    if mode == "same_month"
                    else "nearest available period in the dataset"
                )
                return (
                    f"No exact records were found for {number_to_month_name(month)} {year} in {dataset_title}. "
                    f"I retried using schema-aware variable matching (Filing Period matched to '{schema_match.get('filing_period')}', "
                    f"FY month matched to '{schema_match.get('fy_month')}') and selected the {mode_text}: {period_label}. "
                    f"Snapshot: Filings {nearest.get('filings', 0)}, Late filings {nearest.get('late_count', 0)} "
                    f"({nearest.get('late_rate_pct', 0)}%), Filing fee {nearest.get('filing_fee', 0)}, "
                    f"Penalty paid {nearest.get('penalty_paid', 0)}, Outstanding penalty {nearest.get('penalty_outstanding', 0)}, "
                    f"Recovery rate {nearest.get('recovery_rate_pct', 0)}%, Last income total {float(metrics.get('total_last_income', 0) or 0.0):.2f}, "
                    f"Requested waiver total {float(metrics.get('total_requested_waiver', 0) or 0.0):.2f}. "
                    f"Top REASONS FOR WAIVER: {reason_text}."
                )
            if isinstance(context, dict) and context.get("uploaded_file_id"):
                try:
                    file_id = int(context.get("uploaded_file_id"))
                    cross = cross_sheet_period_answer_for_uploaded_file(file_id, month, year, dataset_title)
                    if cross:
                        return cross
                except Exception:
                    pass
            return (
                f"No exact records were found for {number_to_month_name(month)} {year} in {dataset_title}. "
                f"Schema check matched Filing Period to '{schema_match.get('filing_period')}' and FY month to "
                f"'{schema_match.get('fy_month')}', but no period-level rows were available for retry."
            )
        reason_text = format_top_waiver_reasons(tables, limit=3)
        return (
            f"{number_to_month_name(month)} {year} financial and compliance snapshot for {dataset_title}: "
            f"Filings {matched['filings']}, Late filings {matched['late_count']} ({matched['late_rate_pct']}%), "
            f"Filing fee {matched['filing_fee']}, Penalty paid {matched['penalty_paid']}, "
            f"Outstanding penalty {matched['penalty_outstanding']}, Recovery rate {matched['recovery_rate_pct']}%, "
            f"Last income total {float(metrics.get('total_last_income', 0) or 0.0):.2f}, "
            f"Requested waiver total {float(metrics.get('total_requested_waiver', 0) or 0.0):.2f}. "
            f"Top REASONS FOR WAIVER: {reason_text}."
        )
    return None


def waiver_ranking_query(question: str) -> bool:
    q = question.lower()
    return ("waiver" in q) and ("most" in q or "top" in q or "highest" in q) and (
        "organization" in q or "ngo" in q
    )


def compute_top_waiver_organizations_for_file(file_id: int, limit: int = 5) -> list[dict[str, object]]:
    with get_db() as conn:
        sheets = conn.execute(
            "SELECT table_name, mapping_json, sheet_name FROM uploaded_sheets WHERE file_id = ?",
            (file_id,),
        ).fetchall()

    all_rows: list[dict[str, object]] = []
    for sheet in sheets:
        mapping = json.loads(sheet["mapping_json"])
        table_name = sheet["table_name"]
        db_keys = [item["db_key"] for item in mapping]
        with get_db() as conn:
            rows = conn.execute(
                f"SELECT {', '.join(quote_identifier(k) for k in db_keys)} FROM {quote_identifier(table_name)}"
            ).fetchall()
        frame = pd.DataFrame([{k: row[k] for k in db_keys} for row in rows], columns=db_keys)
        if frame.empty:
            continue

        candidate_org_cols = []
        for item in mapping:
            label = normalize_label(item["label"])
            if ("name ngo" in label) or ("name of the ngo" in label) or ("organization name" in label):
                candidate_org_cols.append(item["db_key"])
        if not candidate_org_cols:
            fallback = find_column(mapping, ["name"])
            if fallback:
                candidate_org_cols.append(fallback)

        col_org = None
        best_score = -1.0
        for candidate in candidate_org_cols:
            s = frame[candidate].astype(str).str.strip()
            non_empty = s[s != ""]
            if non_empty.empty:
                continue
            alpha_ratio = non_empty.str.contains(r"[A-Za-z]", regex=True).mean()
            numeric_only_ratio = non_empty.str.fullmatch(r"\d+").mean()
            score = float(alpha_ratio) - float(numeric_only_ratio)
            if score > best_score:
                best_score = score
                col_org = candidate

        col_requested_waiver = find_column(mapping, ["requested", "waiver"]) or find_column(mapping, ["request", "waiver"])
        col_reason_waiver = find_column(mapping, ["reason", "waiver"])

        if not col_org:
            continue

        org = frame[col_org].astype(str).str.strip()
        org = org.replace("", "Unknown Organization")

        if col_requested_waiver:
            waiver_numeric = to_numeric_series(frame[col_requested_waiver])
            if waiver_numeric.sum() > 0:
                waiver_value = waiver_numeric
            else:
                waiver_value = to_bool_series(frame[col_requested_waiver]).astype(int)
        elif col_reason_waiver:
            waiver_value = frame[col_reason_waiver].astype(str).str.strip().ne("").astype(int)
        else:
            waiver_value = pd.Series([0] * len(frame))

        temp = pd.DataFrame({"organization": org, "waiver_value": waiver_value})
        grouped = temp.groupby("organization", dropna=False)["waiver_value"].sum().reset_index()
        for row in grouped.itertuples():
            all_rows.append(
                {
                    "organization": row.organization,
                    "waiver_value": float(row.waiver_value),
                    "sheet_name": sheet["sheet_name"],
                }
            )

    if not all_rows:
        return []

    combined = pd.DataFrame(all_rows)
    top = (
        combined.groupby("organization", dropna=False)["waiver_value"]
        .sum()
        .sort_values(ascending=False)
        .head(limit)
    )
    return [{"organization": org, "waiver_value": float(val)} for org, val in top.items()]


def ask_dataset_ai(dataset_title: str, payload: dict[str, object], question: str) -> str:
    metrics = payload.get("metrics", {})
    tables = payload.get("tables", {})
    month, year = extract_month_year_from_question(question)
    if not month or not year:
        title_month, title_year = extract_month_year_hint(dataset_title)
        month = month or title_month
        year = year or title_year

    if is_management_summary_query(question):
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            try:
                from openai import OpenAI

                ym = tables.get("year_month_financial", []) if isinstance(tables, dict) else []
                reasons = tables.get("waiver_reasons", []) if isinstance(tables, dict) else []
                client = OpenAI(api_key=api_key)
                response = client.responses.create(
                    model="gpt-5-nano",
                    input=(
                        "You are a compliance analytics assistant writing a management report.\n"
                        "Write approximately 500 words in clear, formal executive language.\n"
                        "Use only the data provided below. Do not invent values.\n"
                        "Must discuss these KPIs explicitly: filings, late filings, filing fee, penalty paid, "
                        "outstanding penalty, recovery rate, Last income, Requested waiver, and top reasons for waiver.\n"
                        "If the requested month/year has no exact record, clearly state that and explain the historical month-only fallback.\n"
                        "End with 3 concise action points.\n\n"
                        f"Dataset title: {dataset_title}\n"
                        f"Requested month: {month}\n"
                        f"Requested year: {year}\n"
                        f"Metrics: {json.dumps(metrics)}\n"
                        f"Year-month financial rows: {json.dumps(ym)}\n"
                        f"Waiver reasons: {json.dumps(reasons)}\n\n"
                        f"User request: {question}"
                    ),
                    store=False,
                )
                text = (response.output_text or "").strip()
                if text:
                    return text
            except Exception as exc:
                fallback = build_printable_report(
                    dataset_title, metrics, tables, month=month, year=year, summary_words=500
                )
                return f"{fallback}\n\nNote: AI service unavailable ({exc}). Returned local analytics summary."

        return build_printable_report(dataset_title, metrics, tables, month=month, year=year, summary_words=500)

    local = local_dataset_answer(dataset_title, payload, question)
    if local:
        return local

    api_key = os.getenv("OPENAI_API_KEY")
    ym = tables.get("year_month_financial", []) if isinstance(tables, dict) else []
    ym_sample = ym[-24:] if isinstance(ym, list) else []
    brief = (
        f"Dataset: {dataset_title}\n"
        f"Compliance rate: {metrics.get('late_filing_rate_pct', 0)}% late filing rate\n"
        f"Recovery rate: {metrics.get('recovery_rate_pct', 0)}%\n"
        f"Avg processing days: {metrics.get('avg_processing_days', 0)}\n"
        f"Nearing enforcement: {metrics.get('nearing_enforcement_count', 0)}\n"
        f"Year-month financial breakdown sample (latest): {json.dumps(ym_sample)}"
    )

    if not api_key:
        return build_printable_report(dataset_title, metrics, tables, month=month, year=year, summary_words=500)

    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        response = client.responses.create(
            model="gpt-5-nano",
            input=(
                "You are a compliance analytics assistant. Use only the dataset summary below.\n"
                f"{brief}\n\nUser question: {question}"
            ),
            store=False,
        )
        return response.output_text
    except Exception as exc:
        fallback = build_printable_report(dataset_title, metrics, tables, month=month, year=year, summary_words=500)
        return f"{fallback}\n\nNote: AI service unavailable ({exc}). Returned local analytics summary."


def answer_annual_question(sheet_name: str, question: str) -> str:
    mapping = get_mapping()
    payload = get_cached_dashboard_payload("annual_returns", mapping, sheet_name_hint=sheet_name)
    return ask_dataset_ai(f"{EXCEL_PATH.name} / {sheet_name}", payload, question)


def answer_uploaded_question(sheet_id: int, question: str) -> str:
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        return "Uploaded sheet not found."
    if waiver_ranking_query(question):
        top = compute_top_waiver_organizations_for_file(meta["file_id"], limit=5)
        if top:
            first = top[0]
            tail = ", ".join(f"{item['organization']} ({item['waiver_value']:.0f})" for item in top[1:])
            return (
                f"Across all sheets in {meta['original_filename']}, the organization with the most waivers is "
                f"{first['organization']} with {first['waiver_value']:.0f}. "
                + (f"Next: {tail}." if tail else "")
            )
        return (
            f"I checked all sheets in {meta['original_filename']}, but no waiver values were found "
            "to rank organizations."
        )
    mapping = json.loads(meta["mapping_json"])
    payload = get_cached_dashboard_payload(
        meta["table_name"],
        mapping,
        sheet_name_hint=meta["sheet_name"],
        snapshot_sheet_id=sheet_id,
    )
    payload["context"] = {"uploaded_file_id": int(meta["file_id"]), "sheet_id": int(sheet_id)}
    return ask_dataset_ai(f"{meta['original_filename']} / {meta['sheet_name']}", payload, question)


def infer_ai_dataset_from_question(question: str) -> dict[str, object] | None:
    q = normalize_label(question)
    if not q:
        return None

    if "annual return" in q or "annual returns" in q or "annual database" in q:
        return {"type": "annual", "sheet_name": SHEET_NAME}

    ensure_uploaded_registry()
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT s.id AS sheet_id, s.sheet_name, f.original_filename, s.file_id
            FROM uploaded_sheets s
            JOIN uploaded_files f ON f.id = s.file_id
            ORDER BY s.id DESC
            """
        ).fetchall()

    if not rows:
        return None

    stopwords = {"file", "sheet", "data", "dataset", "report", "summary", "the", "and", "for"}
    best_row = None
    best_score = 0
    for row in rows:
        file_text = normalize_label(Path(row["original_filename"]).stem)
        sheet_text = normalize_label(row["sheet_name"])
        score = 0

        file_tokens = [t for t in file_text.split() if len(t) >= 3 and t not in stopwords]
        sheet_tokens = [t for t in sheet_text.split() if len(t) >= 3 and t not in stopwords]

        for token in file_tokens:
            if re.search(rf"\b{re.escape(token)}\b", q):
                score += 2
        for token in sheet_tokens:
            if re.search(rf"\b{re.escape(token)}\b", q):
                score += 1

        if "waiver" in q and "waiver" in file_text:
            score += 5
        if "sheet" in q and any(re.search(rf"\b{re.escape(t)}\b", q) for t in sheet_tokens):
            score += 2
        if "file" in q and any(re.search(rf"\b{re.escape(t)}\b", q) for t in file_tokens):
            score += 2

        if score > best_score:
            best_score = score
            best_row = row

    if best_row and best_score > 0:
        return {"type": "uploaded", "sheet_id": int(best_row["sheet_id"])}
    return None

@app.route("/")
def index():
    try:
        mapping = get_mapping()
    except FileNotFoundError as exc:
        abort(500, description=str(exc))

    with get_db() as conn:
        total_rows = conn.execute("SELECT COUNT(*) AS total FROM annual_returns").fetchone()["total"]

    sheets: list[dict[str, object]] = [
        {
            "workbook": EXCEL_PATH.name,
            "name": SHEET_NAME,
            "rows": int(total_rows),
            "columns": len(mapping),
            "open_url": url_for("annual_returns_sheet", sheet_name=SHEET_NAME),
        }
    ]
    if WAIVER_SUMMARY_PATH.exists():
        try:
            wb = load_workbook(WAIVER_SUMMARY_PATH, read_only=True, data_only=True)
            for ws in wb.worksheets:
                sheet_name = str(ws.title)
                rows = max(0, int(ws.max_row or 0) - 1)
                cols = int(ws.max_column or 0)
                sheets.append(
                    {
                        "workbook": WAIVER_SUMMARY_PATH.name,
                        "name": sheet_name,
                        "rows": rows,
                        "columns": cols,
                        "open_url": url_for(
                            "default_workbook_sheet_view",
                            workbook_key="waiver_summary_2025",
                            sheet_name=sheet_name,
                        ),
                    }
                )
            wb.close()
        except Exception:
            app.logger.exception("Could not read default workbook %s", WAIVER_SUMMARY_PATH)
    refreshed = (request.args.get("refreshed") or "").strip() == "1"
    return render_template(
        "index.html",
        ai_context={"type": "annual", "sheet_name": SHEET_NAME},
        excel_name=EXCEL_PATH.name,
        excel_path=str(EXCEL_PATH),
        csv_path=str(CSV_PATH),
        db_path=db_display_target(),
        sheets=sheets,
        refreshed=refreshed,
    )


@app.get("/default-workbooks/<workbook_key>/sheet/<sheet_name>")
def default_workbook_sheet_view(workbook_key: str, sheet_name: str):
    key = str(workbook_key or "").strip().lower()
    if key != "waiver_summary_2025":
        abort(404, description="Default workbook not found.")
    if not WAIVER_SUMMARY_PATH.exists():
        abort(404, description="Workbook file not found.")

    xls = pd.ExcelFile(WAIVER_SUMMARY_PATH)
    if sheet_name not in xls.sheet_names:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    frame = pd.read_excel(WAIVER_SUMMARY_PATH, sheet_name=sheet_name, dtype=str, keep_default_na=False).fillna("")
    frame = frame[
        [
            c
            for c in frame.columns
            if not re.fullmatch(r"unnamed(?:[\s_:]*\d+)?", str(c).strip().lower())
        ]
    ]
    computed, _ = _prepare_waiver_sheet_computed_frame(frame, sheet_name)
    headers = [str(c).strip() for c in computed.columns]
    rows_data = computed.to_dict(orient="records")

    page, per_page = parse_paging(request.args)
    total_rows = len(rows_data)
    total_pages = max(1, (total_rows + per_page - 1) // per_page)
    page = max(1, min(page, total_pages))
    start = (page - 1) * per_page
    end = start + per_page
    page_rows = rows_data[start:end]

    return render_template(
        "default_sheet_readonly.html",
        workbook_name=WAIVER_SUMMARY_PATH.name,
        sheet_name=sheet_name,
        workbook_key=workbook_key,
        headers=headers,
        rows=page_rows,
        page=page,
        per_page=per_page,
        total_rows=total_rows,
        total_pages=total_pages,
        first_url=url_for(
            "default_workbook_sheet_view",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
            page=1,
            per_page=per_page,
        )
        if page > 1
        else None,
        prev_url=url_for(
            "default_workbook_sheet_view",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
            page=page - 1,
            per_page=per_page,
        )
        if page > 1
        else None,
        next_url=url_for(
            "default_workbook_sheet_view",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
            page=page + 1,
            per_page=per_page,
        )
        if page < total_pages
        else None,
        last_url=url_for(
            "default_workbook_sheet_view",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
            page=total_pages,
            per_page=per_page,
        )
        if page < total_pages
        else None,
        dashboard_url=url_for(
            "default_workbook_sheet_dashboard",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
        ),
        export_url=url_for(
            "default_workbook_sheet_export",
            workbook_key=workbook_key,
            sheet_name=sheet_name,
        ),
    )


def _load_default_waiver_sheet_frame(workbook_key: str, sheet_name: str) -> pd.DataFrame:
    key = str(workbook_key or "").strip().lower()
    if key != "waiver_summary_2025":
        abort(404, description="Default workbook not found.")
    if not WAIVER_SUMMARY_PATH.exists():
        abort(404, description="Workbook file not found.")
    xls = pd.ExcelFile(WAIVER_SUMMARY_PATH)
    if sheet_name not in xls.sheet_names:
        abort(404, description=f"Sheet '{sheet_name}' not found.")
    frame = pd.read_excel(WAIVER_SUMMARY_PATH, sheet_name=sheet_name, dtype=str, keep_default_na=False).fillna("")
    frame = frame[
        [
            c
            for c in frame.columns
            if not re.fullmatch(r"unnamed(?:[\s_:]*\d+)?", str(c).strip().lower())
        ]
    ].copy()
    frame.columns = [str(c).strip() for c in frame.columns]
    return frame


def _label_for_key(mapping: list[dict[str, str]], db_key: str | None) -> str | None:
    key = str(db_key or "").strip()
    if not key:
        return None
    for item in mapping:
        if str(item.get("db_key", "")).strip() == key:
            return str(item.get("label", "")).strip() or key
    return key


def _prepare_waiver_sheet_computed_frame(frame: pd.DataFrame, sheet_name: str) -> tuple[pd.DataFrame, dict[str, str]]:
    headers = [str(c).strip() for c in frame.columns]
    mapping = column_mapping_from_headers(headers)
    db_keys = [item["db_key"] for item in mapping]
    data = frame.copy()
    data.columns = db_keys

    key_info = get_waiver_rule_keys(mapping)
    if not key_info.get("granted_waiver"):
        used = {item["db_key"] for item in mapping}
        grant_key = "granted_waiver_amount"
        idx = 2
        while grant_key in used:
            grant_key = f"granted_waiver_amount_{idx}"
            idx += 1
        mapping.append({"label": "Granted Waiver Amount", "db_key": grant_key})
        data[grant_key] = ""
        key_info = get_waiver_rule_keys(mapping)
    if not key_info.get("status"):
        used = {item["db_key"] for item in mapping}
        status_key = "status"
        idx = 2
        while status_key in used:
            status_key = f"status_{idx}"
            idx += 1
        mapping.append({"label": "status", "db_key": status_key})
        data[status_key] = ""
        key_info = get_waiver_rule_keys(mapping)
    if not key_info.get("balance"):
        used = {item["db_key"] for item in mapping}
        balance_key = "balance"
        idx = 2
        while balance_key in used:
            balance_key = f"balance_{idx}"
            idx += 1
        mapping.append({"label": "Balance", "db_key": balance_key})
        data[balance_key] = ""
        key_info = get_waiver_rule_keys(mapping)

    for item in mapping:
        key = str(item.get("db_key", "")).strip()
        if key and key not in data.columns:
            data[key] = ""
    data = data[[item["db_key"] for item in mapping]].copy()

    updated_rows: list[dict[str, object]] = []
    for row in data.to_dict(orient="records"):
        values = {str(k): row.get(k, "") for k in data.columns}
        apply_waiver_balance_and_status_to_values(values, mapping, WAIVER_RULE_FILE_NAME, sheet_name)
        updated_rows.append(values)
    computed_db = pd.DataFrame(updated_rows, columns=[item["db_key"] for item in mapping]).fillna("")

    label_by_key = {item["db_key"]: item["label"] for item in mapping}
    computed = computed_db.rename(columns=label_by_key)
    key_map = {
        "org": _label_for_key(mapping, find_column(mapping, ["name", "ngo"])),
        "outstanding_returns": _label_for_key(mapping, find_column(mapping, ["outstanding", "return"])),
        "reason": _label_for_key(mapping, find_column(mapping, ["reason", "waiver"])),
        "total_penalty": _label_for_key(mapping, key_info.get("total_penalty")),
        "penalty_paid": _label_for_key(mapping, key_info.get("penalty_paid")),
        "comment": _label_for_key(mapping, key_info.get("comment")),
        "status": _label_for_key(mapping, key_info.get("status")),
        "balance": _label_for_key(mapping, key_info.get("balance")),
        "granted_waiver": _label_for_key(mapping, key_info.get("granted_waiver")),
    }
    return computed, key_map


def _ensure_min_words(text: str, min_words: int = 100) -> str:
    base = str(text or "").strip()
    words = base.split()
    if len(words) >= min_words:
        return base
    filler = (
        "This interpretation should be reviewed alongside the row-level table because individual cases can still differ "
        "in timing, document quality, and committee judgment language. Use this section to guide prioritization, then "
        "validate outliers and edge records in operational review before final management reporting or enforcement action."
    )
    out = base
    while len(out.split()) < min_words:
        out = f"{out} {filler}".strip()
    return out


def _paragraph_every_two_sentences(text: str) -> str:
    raw = str(text or "").strip()
    if not raw:
        return raw
    # Insert a paragraph break after every second sentence terminator.
    matches = list(re.finditer(r"[.!?](?:\s+|$)", raw))
    if len(matches) < 3:
        return raw

    chunks: list[str] = []
    start = 0
    i = 0
    while i < len(matches):
        end_idx = matches[min(i + 1, len(matches) - 1)].end()
        chunk = raw[start:end_idx].strip()
        if chunk:
            chunks.append(chunk)
        start = end_idx
        i += 2
    tail = raw[start:].strip()
    if tail:
        chunks.append(tail)
    return "\n\n".join(chunks)


def _build_waiver_dashboard_interpretations(
    sheet_name: str,
    summary_cards: list[dict[str, str]],
    chart_data: dict[str, dict[str, object]],
) -> dict[str, str]:
    summary_map = {str(i.get("label", "")): str(i.get("value", "")) for i in summary_cards}
    status_labels = chart_data.get("status", {}).get("labels", []) or []
    status_values = chart_data.get("status", {}).get("values", []) or []
    reasons_labels = chart_data.get("reasons", {}).get("labels", []) or []
    reasons_values = chart_data.get("reasons", {}).get("values", []) or []
    decision_labels = chart_data.get("decision", {}).get("labels", []) or []
    decision_values = chart_data.get("decision", {}).get("values", []) or []
    fin_labels = chart_data.get("financials", {}).get("labels", []) or []
    fin_values = chart_data.get("financials", {}).get("values", []) or []
    out_labels = chart_data.get("outstanding_orgs", {}).get("labels", []) or []
    out_values = chart_data.get("outstanding_orgs", {}).get("values", []) or []

    status_top = status_labels[0] if status_labels else "Unknown"
    status_top_n = int(status_values[0]) if status_values else 0
    reason_top = reasons_labels[0] if reasons_labels else "Unspecified"
    reason_top_n = int(reasons_values[0]) if reasons_values else 0
    decision_top = decision_labels[0] if decision_labels else "Unclear/No Decision"
    decision_top_n = int(decision_values[0]) if decision_values else 0

    financial_map = {str(k): float(v or 0.0) for k, v in zip(fin_labels, fin_values)}
    total_penalty = financial_map.get("Total Penalty", 0.0)
    paid = financial_map.get("Penalty Paid", 0.0)
    granted = financial_map.get("Granted Waiver Amount", 0.0)
    balance = financial_map.get("Balance", 0.0)

    top_org = out_labels[0] if out_labels else "Unknown"
    top_org_amt = float(out_values[0]) if out_values else 0.0

    summary_txt = _ensure_min_words(
        f"For sheet {sheet_name}, the summary indicates rows={summary_map.get('Rows', '0')}, outstanding returns total "
        f"{summary_map.get('Outstanding Returns (sum)', '0')}, and recovery coverage {summary_map.get('Recovery Coverage', '0%')}. "
        f"The financial profile shows a combined penalty base of {summary_map.get('Total Penalty', '0')} against paid amount "
        f"{summary_map.get('Penalty Paid', '0')} and granted waiver {summary_map.get('Granted Waiver Amount', '0')}, leaving "
        f"balance {summary_map.get('Balance', '0')}. Operationally, this means the committee comment quality directly influences "
        f"financial closure quality. Where comments are clear, waiver classification is deterministic; where comments are vague, "
        f"records tend to remain partially paid and carry forward balance. This summary should be read as a control panel for collection "
        f"risk, waiver-governance consistency, and trend readiness before presenting month-end compliance outcomes."
    )
    status_txt = _ensure_min_words(
        f"The status distribution for {sheet_name} is led by '{status_top}' with {status_top_n} records. This pattern gives a direct "
        f"signal of closure quality because status is derived from penalty arithmetic and committee-comment outcomes. A concentration in "
        f"fully paid records usually indicates either stronger direct settlement or explicit waivers that reduce residual balances, while "
        f"partially paid concentrations usually point to unresolved liabilities or staged payment instructions. Use this chart to prioritize "
        f"follow-up queues: start with high-balance partially paid organizations, then review no-payment rows for enforcement escalation. "
        f"If status labels fluctuate heavily between periods, validate source consistency in total penalty, paid amount, and committee text "
        f"structure, because even small classification noise can distort management conclusions about recovery effectiveness."
    )
    decision_txt = _ensure_min_words(
        f"Committee-comment outcomes are currently dominated by '{decision_top}' with {decision_top_n} records. This chart is important "
        f"because it traces how textual decisions translate into granted waiver and final balance. Full-waiver language should collapse "
        f"balance to zero when unpaid exposure exists, to-pay amounts should convert to partial waivers, and rejected or unclear comments "
        f"should retain liabilities unless payment already closes the gap. If unclear decisions are frequent, management should standardize "
        f"comment syntax and decision wording to improve auditability. This distribution can also detect process drift: sudden shifts from "
        f"numeric to-pay instructions to vague narrative comments often reduce computational clarity and create reconciliation overhead "
        f"for operations, finance, and legal review."
    )
    reasons_txt = _ensure_min_words(
        f"The waiver-reasons chart highlights '{reason_top}' as the leading reason with {reason_top_n} records. This is a qualitative "
        f"risk lens showing why organizations seek relief and where policy communication may need reinforcement. Repeated reasons across "
        f"many records can indicate systemic friction such as filing bottlenecks, governance gaps, or recurring administrative constraints. "
        f"From an analytics perspective, reason frequency should be cross-read with status and balance outcomes: if high-frequency reasons "
        f"also show low closure rates, those cohorts deserve targeted intervention plans. Keep in mind that narrative reasons can vary in "
        f"spelling and detail, so periodic normalization and taxonomy cleanup will improve comparability and allow stronger trend analysis "
        f"without losing the original legal narrative in the source worksheet."
    )
    financial_txt = _ensure_min_words(
        f"The financial component chart summarizes core monetary movement: total penalty={total_penalty:,.2f}, penalty paid={paid:,.2f}, "
        f"granted waiver={granted:,.2f}, and remaining balance={balance:,.2f}. This panel represents the strongest quantitative checkpoint "
        f"for enforcement and revenue forecasting. When granted waiver rises without a corresponding reduction in balance, that usually signals "
        f"data quality issues or unparsed committee directives. When paid amount grows but balance remains high, collections may be progressing "
        f"against only a subset of high-exposure cases. The recommended operational use is to monitor the ratio of (paid + granted) to total "
        f"penalty and track whether balance is declining period over period. This interpretation supports both finance recovery planning and "
        f"governance assurance around waiver approvals."
    )
    outstanding_txt = _ensure_min_words(
        f"Outstanding returns concentration is led by '{top_org}' at {top_org_amt:,.2f}. This chart is a prioritization tool: it identifies "
        f"which organizations contribute most to unresolved exposure and should be reviewed first in enforcement or case-management workflows. "
        f"A steep concentration curve suggests that a relatively small set of organizations drives a large share of operational risk, making "
        f"targeted follow-up more efficient than broad generic reminders. Pair this chart with committee outcomes and status distribution to "
        f"separate high-risk unresolved accounts from cases that are already under clear payment or waiver pathways. If top organizations remain "
        f"unchanged across reporting periods, escalation strategy, timeline controls, and documentary verification should be strengthened to "
        f"improve closure velocity and reduce carry-over burden."
    )
    return {
        "summary": _paragraph_every_two_sentences(summary_txt),
        "status": _paragraph_every_two_sentences(status_txt),
        "decision": _paragraph_every_two_sentences(decision_txt),
        "reasons": _paragraph_every_two_sentences(reasons_txt),
        "financials": _paragraph_every_two_sentences(financial_txt),
        "outstanding": _paragraph_every_two_sentences(outstanding_txt),
    }


@app.get("/default-workbooks/<workbook_key>/sheet/<sheet_name>/export")
def default_workbook_sheet_export(workbook_key: str, sheet_name: str):
    key = str(workbook_key or "").strip().lower()
    if key != "waiver_summary_2025":
        abort(404, description="Default workbook not found.")
    if not WAIVER_SUMMARY_PATH.exists():
        abort(404, description="Workbook file not found.")
    xls = pd.ExcelFile(WAIVER_SUMMARY_PATH)
    if sheet_name not in xls.sheet_names:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for ws_name in xls.sheet_names:
            ws = pd.read_excel(WAIVER_SUMMARY_PATH, sheet_name=ws_name, dtype=str, keep_default_na=False).fillna("")
            ws = ws[
                [
                    c
                    for c in ws.columns
                    if not re.fullmatch(r"unnamed(?:[\s_:]*\d+)?", str(c).strip().lower())
                ]
            ].copy()
            ws.columns = [str(c).strip() for c in ws.columns]
            computed, _ = _prepare_waiver_sheet_computed_frame(ws, ws_name)
            computed.to_excel(writer, sheet_name=ws_name[:31] or "Sheet1", index=False)
    output.seek(0)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{Path(WAIVER_SUMMARY_PATH.name).stem}_edited_{stamp}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/default-workbooks/<workbook_key>/sheet/<sheet_name>/dashboard")
def default_workbook_sheet_dashboard(workbook_key: str, sheet_name: str):
    frame = _load_default_waiver_sheet_frame(workbook_key, sheet_name)
    computed, key_map = _prepare_waiver_sheet_computed_frame(frame, sheet_name)

    def col_series(key: str) -> pd.Series:
        col = key_map.get(key)
        if col and col in computed.columns:
            return computed[col]
        return pd.Series([""] * len(computed))

    org_series = col_series("org").astype(str).str.strip().replace("", "Unknown")
    reason_series = col_series("reason").astype(str).str.strip()
    comment_series = col_series("comment").astype(str).str.strip()
    status_series = col_series("status").astype(str).str.strip().replace("", "Unknown")

    total_penalty = col_series("total_penalty").map(parse_amount_value)
    penalty_paid = col_series("penalty_paid").map(parse_amount_value)
    balance_vals = col_series("balance").map(parse_amount_value)
    granted_vals = col_series("granted_waiver").map(parse_amount_value)
    outstanding_returns = col_series("outstanding_returns").map(parse_amount_value)

    rows = int(len(computed))
    sum_total_penalty = float(total_penalty.sum())
    sum_penalty_paid = float(penalty_paid.sum())
    sum_granted = float(granted_vals.sum())
    sum_balance = float(balance_vals.sum())
    sum_outstanding_returns = float(outstanding_returns.sum())
    recovery_pct = ((sum_penalty_paid + sum_granted) / sum_total_penalty * 100.0) if sum_total_penalty > 0 else 0.0

    status_counts = status_series.value_counts().head(10)
    reason_counts = reason_series[reason_series != ""].value_counts().head(10)
    if reason_counts.empty:
        reason_counts = pd.Series([rows], index=["Unspecified"]) if rows else pd.Series(dtype=int)

    decision_counter: Counter[str] = Counter()
    for txt in comment_series:
        parsed = _extract_amounts_from_comment(txt)
        if parsed.get("rejected"):
            decision_counter["Rejected"] += 1
        elif parsed.get("full_waiver"):
            decision_counter["Full Waiver"] += 1
        elif parsed.get("to_pay"):
            decision_counter["To Pay Amount"] += 1
        elif parsed.get("waived") or parsed.get("percent_waiver") is not None:
            decision_counter["Partial Waiver"] += 1
        else:
            decision_counter["Unclear/No Decision"] += 1

    top_outstanding = (
        pd.DataFrame({"org": org_series, "outstanding": outstanding_returns})
        .groupby("org", dropna=False)["outstanding"]
        .sum()
        .sort_values(ascending=False)
        .head(10)
    )

    chart_data = {
        "status": {"labels": status_counts.index.tolist(), "values": [int(v) for v in status_counts.tolist()]},
        "reasons": {"labels": reason_counts.index.tolist(), "values": [int(v) for v in reason_counts.tolist()]},
        "decision": {
            "labels": list(decision_counter.keys()),
            "values": [int(v) for v in decision_counter.values()],
        },
        "financials": {
            "labels": ["Total Penalty", "Penalty Paid", "Granted Waiver Amount", "Balance"],
            "values": [
                round(sum_total_penalty, 2),
                round(sum_penalty_paid, 2),
                round(sum_granted, 2),
                round(sum_balance, 2),
            ],
        },
        "outstanding_orgs": {
            "labels": top_outstanding.index.tolist(),
            "values": [round(float(v), 2) for v in top_outstanding.tolist()],
        },
    }

    summary_cards = [
        {"label": "Rows", "value": f"{rows:,}"},
        {"label": "Outstanding Returns (sum)", "value": f"{sum_outstanding_returns:,.2f}"},
        {"label": "Total Penalty", "value": f"{sum_total_penalty:,.2f}"},
        {"label": "Penalty Paid", "value": f"{sum_penalty_paid:,.2f}"},
        {"label": "Granted Waiver Amount", "value": f"{sum_granted:,.2f}"},
        {"label": "Balance", "value": f"{sum_balance:,.2f}"},
        {"label": "Recovery Coverage", "value": f"{recovery_pct:.2f}%"},
    ]
    interpretations = _build_waiver_dashboard_interpretations(sheet_name, summary_cards, chart_data)

    return render_template(
        "waiver_sheet_dashboard.html",
        dashboard_title=f"{WAIVER_SUMMARY_PATH.name} - {sheet_name}",
        back_url=url_for("default_workbook_sheet_view", workbook_key=workbook_key, sheet_name=sheet_name),
        export_url=url_for("default_workbook_sheet_export", workbook_key=workbook_key, sheet_name=sheet_name),
        summary_cards=summary_cards,
        chart_data=chart_data,
        key_map=key_map,
        interpretations=interpretations,
    )


@app.post("/annual-returns/refresh")
def refresh_annual_returns():
    try:
        mapping = build_csv_from_excel()
    except FileNotFoundError as exc:
        abort(500, description=str(exc))
    with get_db() as conn:
        recreate_table_from_csv(conn, mapping)
    return redirect(url_for("index", refreshed="1"))


@app.post("/uploaded/files/<int:file_id>/refresh")
def refresh_uploaded_file(file_id: int):
    meta = get_uploaded_file_meta(file_id)
    if meta is None:
        abort(404, description="Uploaded file not found.")
    workbook_path = UPLOAD_DIR / str(meta["stored_filename"])
    if not workbook_path.exists():
        abort(404, description=f"Uploaded workbook file not found on disk: {workbook_path.name}")

    with get_db() as conn:
        backup_count = backup_uploaded_file_state(conn, file_id)
        sheet_rows = conn.execute(
            "SELECT id, table_name FROM uploaded_sheets WHERE file_id = ? ORDER BY id ASC",
            (file_id,),
        ).fetchall()
        existing_sheet_ids = [int(row["id"]) for row in sheet_rows]
        for row in sheet_rows:
            conn.execute(f"DROP TABLE IF EXISTS {quote_identifier(row['table_name'])}")
        conn.execute("DELETE FROM uploaded_sheets WHERE file_id = ?", (file_id,))
        import_workbook_sheets_into_file(conn, workbook_path, file_id, existing_sheet_ids=existing_sheet_ids)
        conn.commit()
    enforce_waiver_balance_rule_for_existing_uploaded_sheets()
    clear_sidebar_uploaded_links_cache()

    return redirect(url_for("index", refreshed_uploaded=str(file_id), backup_count=str(backup_count)))


@app.get("/uploaded/files/<int:file_id>/backups")
def uploaded_file_backups(file_id: int):
    meta = get_uploaded_file_meta(file_id)
    if meta is None:
        abort(404, description="Uploaded file not found.")
    groups = get_backup_groups_for_file(file_id)
    return render_template(
        "backups.html",
        ai_context={"type": "annual", "sheet_name": SHEET_NAME},
        file_id=file_id,
        file_name=meta["original_filename"],
        backup_groups=groups,
    )


@app.post("/uploaded/files/<int:file_id>/backups/<snapshot_key>/restore")
def restore_uploaded_backup(file_id: int, snapshot_key: str):
    meta = get_uploaded_file_meta(file_id)
    if meta is None:
        abort(404, description="Uploaded file not found.")
    restored = restore_uploaded_file_snapshot(file_id, snapshot_key)
    if restored == 0:
        abort(404, description="Backup snapshot not found.")
    return redirect(url_for("uploaded_file_backups", file_id=file_id))


@app.route("/annual-returns/<sheet_name>")
def annual_returns_sheet(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    mapping = get_mapping()
    page, per_page = parse_paging(request.args)
    cursor_text = (request.args.get("cursor") or "").strip()
    direction = (request.args.get("direction") or "next").strip().lower()
    use_keyset = bool(cursor_text and page >= KEYSET_PAGINATION_MIN_PAGE)
    cursor_value: int | None = None
    if use_keyset:
        try:
            cursor_value = int(cursor_text)
        except Exception:
            use_keyset = False

    with get_db() as conn:
        total_rows = conn.execute("SELECT COUNT(*) AS total FROM annual_returns").fetchone()["total"]
        if use_keyset and cursor_value is not None and direction == "prev":
            db_rows = conn.execute(
                "SELECT * FROM annual_returns WHERE excel_row_number < ? ORDER BY excel_row_number DESC LIMIT ?",
                (cursor_value, per_page),
            ).fetchall()
            db_rows = list(reversed(db_rows))
        elif use_keyset and cursor_value is not None:
            db_rows = conn.execute(
                "SELECT * FROM annual_returns WHERE excel_row_number > ? ORDER BY excel_row_number LIMIT ?",
                (cursor_value, per_page),
            ).fetchall()
        else:
            offset = (page - 1) * per_page
            db_rows = conn.execute(
                "SELECT * FROM annual_returns ORDER BY excel_row_number LIMIT ? OFFSET ?",
                (per_page, offset),
            ).fetchall()

    total_pages = max(1, (total_rows + per_page - 1) // per_page)
    page = min(page, total_pages)
    select_config = get_select_config(mapping)
    date_field_keys = [
        item["db_key"]
        for item in mapping
        if is_date_field_label(item["label"])
    ]

    rows = []
    for db_row in db_rows:
        row = {"id": db_row["id"], "excel_row_number": db_row["excel_row_number"]}
        for item in mapping:
            key = item["db_key"]
            value = db_row[key] or ""
            row[key] = to_date_input_value(value) if key in date_field_keys else value
        rows.append(row)
    first_excel_row = None
    last_excel_row = None
    if rows:
        first_excel_row = int(rows[0]["excel_row_number"])
        last_excel_row = int(rows[-1]["excel_row_number"])
    can_prev = page > 1 and bool(rows)
    can_next = (page < total_pages) if not use_keyset else (len(rows) == per_page and bool(rows))

    return render_template(
        "sheet.html",
        ai_context={"type": "annual", "sheet_name": sheet_name},
        excel_name=EXCEL_PATH.name,
        sheet_name=sheet_name,
        dashboard_url=url_for(
            "annual_dashboard",
            sheet_name=sheet_name,
            dataset="annual_returns",
        ),
        columns=mapping,
        rows=rows,
        date_field_keys=date_field_keys,
        select_config=select_config,
        page=page,
        per_page=per_page,
        total_rows=total_rows,
        total_pages=total_pages,
        first_url=url_for("annual_returns_sheet", sheet_name=sheet_name, page=1, per_page=per_page)
        if page > 1
        else None,
        prev_url=url_for(
            "annual_returns_sheet",
            sheet_name=sheet_name,
            page=page - 1,
            per_page=per_page,
            cursor=first_excel_row,
            direction="prev",
        )
        if can_prev and first_excel_row is not None
        else None,
        next_url=url_for(
            "annual_returns_sheet",
            sheet_name=sheet_name,
            page=page + 1,
            per_page=per_page,
            cursor=last_excel_row,
            direction="next",
        )
        if can_next and last_excel_row is not None
        else None,
        last_url=url_for(
            "annual_returns_sheet", sheet_name=sheet_name, page=total_pages, per_page=per_page
        )
        if page < total_pages
        else None,
    )


@app.post("/annual-returns/<sheet_name>/row/<int:row_id>")
def update_row(sheet_name: str, row_id: int):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    mapping = get_mapping()
    with get_db() as conn:
        existing = conn.execute("SELECT * FROM annual_returns WHERE id = ?", (row_id,)).fetchone()
    if existing is None:
        abort(404, description=f"Row '{row_id}' not found.")

    updates = {}
    for item in mapping:
        key = item["db_key"]
        if key in request.form:
            updates[key] = request.form.get(key, "").strip()
        else:
            updates[key] = (existing[key] or "").strip()
    set_clause = ", ".join(f'"{key}" = ?' for key in updates)
    values = list(updates.values()) + [row_id]

    with get_db() as conn:
        conn.execute(f"UPDATE annual_returns SET {set_clause} WHERE id = ?", values)
        conn.commit()
    bump_data_version()

    page, per_page = parse_paging(request.form)
    return redirect(
        url_for("annual_returns_sheet", sheet_name=sheet_name, page=page, per_page=per_page) + f"#row-{row_id}"
    )


@app.post("/annual-returns/<sheet_name>/save-all")
def annual_save_all(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    payload = request.get_json(silent=True) or {}
    updates = payload.get("updates", [])
    if not isinstance(updates, list):
        abort(400, description="Invalid payload.")

    mapping = get_mapping()
    updated = bulk_update_rows("annual_returns", mapping, updates)
    return jsonify({"ok": True, "updated_rows": updated})


@app.post("/annual-returns/<sheet_name>/create-row")
def create_row(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    mapping = get_mapping()
    db_keys = [item["db_key"] for item in mapping]
    col_names = ", ".join(f'"{key}"' for key in db_keys)
    placeholders = ", ".join("?" for _ in db_keys)

    with get_db() as conn:
        next_excel_row = conn.execute(
            "SELECT COALESCE(MAX(excel_row_number), 1) + 1 AS next_row FROM annual_returns"
        ).fetchone()["next_row"]
        cursor = conn.execute(
            f'INSERT INTO annual_returns (excel_row_number, {col_names}) VALUES (?, {placeholders})',
            [next_excel_row] + ["" for _ in db_keys],
        )
        conn.commit()
        new_id = cursor.lastrowid
    bump_data_version()

    _, per_page = parse_paging(request.form)
    with get_db() as conn:
        total_rows = conn.execute("SELECT COUNT(*) AS total FROM annual_returns").fetchone()["total"]
    target_page = max(1, (total_rows + per_page - 1) // per_page)

    return redirect(
        url_for("annual_returns_sheet", sheet_name=sheet_name, page=target_page, per_page=per_page)
        + f"#row-{new_id}"
    )


@app.post("/annual-returns/<sheet_name>/row/<int:row_id>/delete")
def delete_row(sheet_name: str, row_id: int):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    with get_db() as conn:
        conn.execute("DELETE FROM annual_returns WHERE id = ?", (row_id,))
        conn.commit()
    bump_data_version()

    page, per_page = parse_paging(request.form)
    return redirect(url_for("annual_returns_sheet", sheet_name=sheet_name, page=page, per_page=per_page))


@app.get("/annual-returns/<sheet_name>/export")
def export_sheet(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")

    mapping = get_mapping()
    db_keys = [item["db_key"] for item in mapping]
    labels = [item["label"] for item in mapping]

    with get_db() as conn:
        rows = conn.execute(
            f'SELECT {", ".join(f"""\"{key}\"""" for key in db_keys)} FROM annual_returns ORDER BY excel_row_number'
        ).fetchall()

    records = [{labels[i]: row[db_keys[i]] for i in range(len(db_keys))} for row in rows]
    export_df = pd.DataFrame(records, columns=labels)

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_df.to_excel(writer, sheet_name=SHEET_NAME, index=False)
    output.seek(0)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Annual_Returns_DB_Export_{stamp}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/upload")
def upload_workbook():
    file = request.files.get("workbook")
    if file is None or not file.filename:
        abort(400, description="No file selected.")
    if not file.filename.lower().endswith(".xlsx"):
        abort(400, description="Only .xlsx files are supported.")

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    original_name = Path(file.filename).name
    safe_name = secure_filename(original_name) or "uploaded.xlsx"
    stored_name = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex}_{safe_name}"
    stored_path = UPLOAD_DIR / stored_name
    file.save(stored_path)
    file_id = create_uploaded_file_record(stored_path, original_name, import_status="processing")
    worker = threading.Thread(
        target=import_uploaded_workbook_async,
        args=(file_id, stored_path),
        daemon=True,
    )
    worker.start()
    return redirect(url_for("index", upload_started=str(file_id)))


@app.get("/uploaded/sheets/<int:sheet_id>")
def uploaded_sheet_view(sheet_id: int):
    sanitize_uploaded_sheet_remove_unnamed(sheet_id)
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")

    mapping = json.loads(meta["mapping_json"])
    validation_map = json.loads(meta["validation_json"])
    table_name = meta["table_name"]
    page, per_page = parse_paging(request.args)
    search_key = (request.args.get("search_key") or "").strip()
    search_query = (request.args.get("q") or "").strip()
    cursor_text = (request.args.get("cursor") or "").strip()
    direction = (request.args.get("direction") or "next").strip().lower()
    allowed_keys = {item["db_key"] for item in mapping}

    where_sql = ""
    where_params: list[str] = []
    if search_key in allowed_keys and search_query:
        where_sql = f" WHERE LOWER({quote_identifier(search_key)}) LIKE ?"
        where_params = [f"%{search_query.lower()}%"]

    use_keyset = bool(cursor_text and page >= KEYSET_PAGINATION_MIN_PAGE)
    cursor_value: int | None = None
    if use_keyset:
        try:
            cursor_value = int(cursor_text)
        except Exception:
            use_keyset = False

    with get_db() as conn:
        total_rows = conn.execute(
            f"SELECT COUNT(*) AS total FROM {quote_identifier(table_name)}{where_sql}",
            where_params,
        ).fetchone()["total"]
        if use_keyset and cursor_value is not None and direction == "prev":
            key_sql = (
                f"SELECT * FROM {quote_identifier(table_name)}{where_sql} "
                f"{' AND ' if where_sql else ' WHERE '}excel_row_number < ? "
                f"ORDER BY excel_row_number DESC LIMIT ?"
            )
            db_rows = conn.execute(key_sql, where_params + [cursor_value, per_page]).fetchall()
            db_rows = list(reversed(db_rows))
        elif use_keyset and cursor_value is not None:
            key_sql = (
                f"SELECT * FROM {quote_identifier(table_name)}{where_sql} "
                f"{' AND ' if where_sql else ' WHERE '}excel_row_number > ? "
                f"ORDER BY excel_row_number LIMIT ?"
            )
            db_rows = conn.execute(key_sql, where_params + [cursor_value, per_page]).fetchall()
        else:
            offset = (page - 1) * per_page
            db_rows = conn.execute(
                f"SELECT * FROM {quote_identifier(table_name)}{where_sql} "
                f"ORDER BY excel_row_number LIMIT ? OFFSET ?",
                where_params + [per_page, offset],
            ).fetchall()

    total_pages = max(1, (total_rows + per_page - 1) // per_page)
    page = min(page, total_pages)

    date_field_keys = date_keys_from_mapping(mapping)
    select_config = get_uploaded_select_config(table_name, mapping, validation_map)

    rows = []
    for db_row in db_rows:
        row = {"id": db_row["id"], "excel_row_number": db_row["excel_row_number"]}
        for item in mapping:
            key = item["db_key"]
            value = db_row[key] or ""
            row[key] = to_date_input_value(value) if key in date_field_keys else value
        rows.append(row)
    first_excel_row = None
    last_excel_row = None
    if rows:
        first_excel_row = int(rows[0]["excel_row_number"])
        last_excel_row = int(rows[-1]["excel_row_number"])
    can_prev = page > 1 and bool(rows)
    can_next = (page < total_pages) if not use_keyset else (len(rows) == per_page and bool(rows))

    return render_template(
        "sheet2.html",
        ai_context={"type": "uploaded", "sheet_id": sheet_id},
        sheet_id=sheet_id,
        file_id=meta["file_id"],
        file_name=meta["original_filename"],
        sheet_name=meta["sheet_name"],
        dashboard_url=url_for(
            "uploaded_dashboard",
            sheet_id=sheet_id,
            file_id=meta["file_id"],
            dataset=meta["table_name"],
        ),
        columns=mapping,
        rows=rows,
        date_field_keys=date_field_keys,
        select_config=select_config,
        page=page,
        per_page=per_page,
        search_key=search_key,
        search_query=search_query,
        total_rows=total_rows,
        total_pages=total_pages,
        first_url=url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=1,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
        if page > 1
        else None,
        prev_url=url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page - 1,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
            cursor=first_excel_row,
            direction="prev",
        )
        if can_prev and first_excel_row is not None
        else None,
        next_url=url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page + 1,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
            cursor=last_excel_row,
            direction="next",
        )
        if can_next and last_excel_row is not None
        else None,
        last_url=url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=total_pages,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
        if page < total_pages
        else None,
    )


@app.post("/uploaded/sheets/<int:sheet_id>/titles")
def uploaded_update_titles(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")

    mapping = json.loads(meta["mapping_json"])
    for item in mapping:
        key = item["db_key"]
        incoming = (request.form.get(f"label__{key}") or "").strip()
        if incoming:
            item["label"] = incoming

    new_sheet_name = (request.form.get("sheet_name") or "").strip() or str(meta["sheet_name"])

    with get_db() as conn:
        conn.execute(
            """
            UPDATE uploaded_sheets
            SET sheet_name = ?, mapping_json = ?
            WHERE id = ?
            """,
            (new_sheet_name, json.dumps(mapping), sheet_id),
        )
        conn.commit()
    bump_data_version()
    clear_sidebar_uploaded_links_cache()

    page, per_page = parse_paging(request.form)
    search_key = (request.form.get("search_key") or "").strip()
    search_query = (request.form.get("q") or "").strip()
    return redirect(
        url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
    )


def rebuild_uploaded_table_with_keys(conn: sqlite3.Connection, table_name: str, db_keys: list[str]) -> None:
    table_ident = quote_identifier(table_name)
    tmp_name = f"{table_name}_tmp_{uuid4().hex[:8]}"
    tmp_ident = quote_identifier(tmp_name)

    db_cols_sql = ", ".join(f"{quote_identifier(key)} TEXT" for key in db_keys)
    cols_fragment = f", {db_cols_sql}" if db_cols_sql else ""
    conn.execute(
        f"""
        CREATE TABLE {tmp_ident} (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            excel_row_number INTEGER NOT NULL UNIQUE
            {cols_fragment}
        )
        """
    )

    select_cols = ["id", "excel_row_number"] + db_keys
    select_sql = ", ".join(quote_identifier(col) for col in select_cols)
    conn.execute(
        f"INSERT INTO {tmp_ident} ({select_sql}) SELECT {select_sql} FROM {table_ident}"
    )
    conn.execute(f"DROP TABLE {table_ident}")
    conn.execute(f"ALTER TABLE {tmp_ident} RENAME TO {table_name}")
    conn.execute(
        f"CREATE INDEX IF NOT EXISTS {quote_identifier(safe_index_name('idx', table_name, 'excel_row'))} "
        f"ON {quote_identifier(table_name)}(excel_row_number)"
    )
    ensure_postgres_text_search_indexes(conn, table_name, db_keys)


def drop_uploaded_columns(conn: DbConnection, table_name: str, removed_keys: list[str], keep_keys: list[str]) -> None:
    if not removed_keys:
        return
    if getattr(conn, "backend", "") == "postgres":
        table_ident = quote_identifier(table_name)
        cols = {row["name"] for row in conn.execute(f"PRAGMA table_info({table_name})").fetchall()}
        dropped_any = False
        for key in removed_keys:
            if key not in cols:
                continue
            conn.execute(f"ALTER TABLE {table_ident} DROP COLUMN IF EXISTS {quote_identifier(key)}")
            dropped_any = True
        if dropped_any:
            conn.execute(
                f"CREATE INDEX IF NOT EXISTS {quote_identifier(safe_index_name('idx', table_name, 'excel_row'))} "
                f"ON {quote_identifier(table_name)}(excel_row_number)"
            )
            ensure_postgres_text_search_indexes(conn, table_name, keep_keys)
        return
    rebuild_uploaded_table_with_keys(conn, table_name, keep_keys)


def sanitize_existing_uploaded_sheets_remove_unnamed(force: bool = False) -> int:
    ensure_uploaded_registry()
    if not force and get_app_state_value(UNNAMED_SANITIZE_STATE_KEY) == "1":
        return 0
    changed = 0
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, table_name, mapping_json, validation_json FROM uploaded_sheets ORDER BY id ASC"
        ).fetchall()
        for row in rows:
            try:
                mapping = json.loads(row["mapping_json"])
            except Exception:
                continue
            if not isinstance(mapping, list):
                continue

            keep_mapping = []
            removed_keys: list[str] = []
            for item in mapping:
                if not isinstance(item, dict):
                    continue
                label = str(item.get("label", ""))
                key = str(item.get("db_key", ""))
                if not key:
                    continue
                if is_unnamed_label(label) or is_unnamed_label(key):
                    removed_keys.append(key)
                    continue
                keep_mapping.append({"label": label, "db_key": key})

            if not removed_keys:
                continue
            if not keep_mapping:
                continue

            table_name = str(row["table_name"])
            keep_keys = [item["db_key"] for item in keep_mapping]
            drop_uploaded_columns(conn, table_name, removed_keys, keep_keys)

            try:
                validation_map = json.loads(row["validation_json"])
            except Exception:
                validation_map = {}
            if not isinstance(validation_map, dict):
                validation_map = {}
            for removed in removed_keys:
                validation_map.pop(removed, None)

            row_count = conn.execute(
                f"SELECT COUNT(*) AS total FROM {quote_identifier(table_name)}"
            ).fetchone()["total"]
            conn.execute(
                """
                UPDATE uploaded_sheets
                SET mapping_json = ?, validation_json = ?, row_count = ?
                WHERE id = ?
                """,
                (json.dumps(keep_mapping), json.dumps(validation_map), int(row_count), int(row["id"])),
            )
            changed += 1

        if changed > 0:
            conn.commit()
    set_app_state_value(UNNAMED_SANITIZE_STATE_KEY, "1")
    if changed > 0:
        bump_data_version()
    return changed


def sanitize_uploaded_sheet_remove_unnamed(sheet_id: int) -> bool:
    changed = False
    with get_db() as conn:
        row = conn.execute(
            """
            SELECT s.id, s.table_name, s.mapping_json, s.validation_json, s.sheet_name, f.original_filename
            FROM uploaded_sheets s
            JOIN uploaded_files f ON f.id = s.file_id
            WHERE s.id = ?
            """,
            (sheet_id,),
        ).fetchone()
        if row is None:
            return False
        if not is_waiver_rule_target_sheet(str(row["original_filename"] or ""), str(row["sheet_name"] or "")):
            return False

        try:
            mapping = json.loads(row["mapping_json"])
        except Exception:
            return False
        if not isinstance(mapping, list):
            return False

        keep_mapping = []
        removed_keys: list[str] = []
        for item in mapping:
            if not isinstance(item, dict):
                continue
            label = str(item.get("label", ""))
            key = str(item.get("db_key", ""))
            if not key:
                continue
            if is_unnamed_label(label) or is_unnamed_label(key):
                removed_keys.append(key)
                continue
            keep_mapping.append({"label": label, "db_key": key})

        if not removed_keys or not keep_mapping:
            return False

        table_name = str(row["table_name"])
        keep_keys = [item["db_key"] for item in keep_mapping]
        drop_uploaded_columns(conn, table_name, removed_keys, keep_keys)

        try:
            validation_map = json.loads(row["validation_json"])
        except Exception:
            validation_map = {}
        if not isinstance(validation_map, dict):
            validation_map = {}
        for removed in removed_keys:
            validation_map.pop(removed, None)

        row_count = conn.execute(
            f"SELECT COUNT(*) AS total FROM {quote_identifier(table_name)}"
        ).fetchone()["total"]
        conn.execute(
            """
            UPDATE uploaded_sheets
            SET mapping_json = ?, validation_json = ?, row_count = ?
            WHERE id = ?
            """,
            (json.dumps(keep_mapping), json.dumps(validation_map), int(row_count), int(row["id"])),
        )
        conn.commit()
        changed = True

    if changed:
        bump_data_version()
    return changed


def ensure_waiver_granted_waiver_column(
    conn: DbConnection,
    sheet_id: int,
    table_name: str,
    mapping: list[dict[str, str]],
    validation_json: str,
) -> tuple[list[dict[str, str]], str | None, bool]:
    if not isinstance(mapping, list):
        return mapping, None, False
    keys = get_waiver_rule_keys(mapping)
    existing_key = str(keys.get("granted_waiver") or "").strip()
    if existing_key:
        return mapping, existing_key, False

    used = {str(item.get("db_key", "")).strip() for item in mapping if isinstance(item, dict)}
    used = {k for k in used if k}
    new_key = "granted_waiver_amount"
    n = 2
    while new_key in used:
        new_key = f"granted_waiver_amount_{n}"
        n += 1

    conn.execute(
        f"ALTER TABLE {quote_identifier(table_name)} ADD COLUMN {quote_identifier(new_key)} TEXT"
    )
    mapping = list(mapping) + [{"label": "Granted Waiver Amount", "db_key": new_key}]

    try:
        validation_map = json.loads(validation_json)
    except Exception:
        validation_map = {}
    if not isinstance(validation_map, dict):
        validation_map = {}
    validation_map[new_key] = {}
    conn.execute(
        """
        UPDATE uploaded_sheets
        SET mapping_json = ?, validation_json = ?
        WHERE id = ?
        """,
        (json.dumps(mapping), json.dumps(validation_map), int(sheet_id)),
    )
    return mapping, new_key, True


def enforce_waiver_balance_rule_for_existing_uploaded_sheets() -> int:
    ensure_uploaded_registry()
    updated_rows = 0
    mapping_changed = False
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT s.id, s.table_name, s.mapping_json, s.validation_json, s.sheet_name, f.original_filename
            FROM uploaded_sheets s
            JOIN uploaded_files f ON f.id = s.file_id
            ORDER BY s.id ASC
            """
        ).fetchall()
        for row in rows:
            file_name = str(row["original_filename"] or "")
            sheet_name = str(row["sheet_name"] or "")
            if not is_waiver_rule_target_sheet(file_name, sheet_name):
                continue
            try:
                mapping = json.loads(row["mapping_json"])
            except Exception:
                continue
            if not isinstance(mapping, list):
                continue
            table_name = str(row["table_name"])
            mapping, _, added_granted_col = ensure_waiver_granted_waiver_column(
                conn,
                int(row["id"]),
                table_name,
                mapping,
                str(row["validation_json"] or "{}"),
            )
            if added_granted_col:
                mapping_changed = True
            keys = get_waiver_rule_keys(mapping)
            requested_key = keys.get("requested")
            balance_key = keys.get("balance")
            total_penalty_key = keys.get("total_penalty")
            penalty_paid_key = keys.get("penalty_paid")
            granted_waiver_key = keys.get("granted_waiver")
            comment_key = keys.get("comment")
            comment_keys_raw = keys.get("comment_keys")
            comment_keys = [str(k) for k in (comment_keys_raw or []) if str(k).strip()]
            if not comment_keys and comment_key:
                comment_keys = [str(comment_key)]
            status_key = keys.get("status")
            if not balance_key:
                continue
            if not ((total_penalty_key and penalty_paid_key) or requested_key):
                continue

            select_cols = [quote_identifier("id"), quote_identifier(balance_key)]
            if requested_key:
                q = quote_identifier(requested_key)
                if q not in select_cols:
                    select_cols.append(q)
            if total_penalty_key:
                q = quote_identifier(total_penalty_key)
                if q not in select_cols:
                    select_cols.append(q)
            if penalty_paid_key:
                q = quote_identifier(penalty_paid_key)
                if q not in select_cols:
                    select_cols.append(q)
            if granted_waiver_key:
                q = quote_identifier(granted_waiver_key)
                if q not in select_cols:
                    select_cols.append(q)
            for c_key in comment_keys:
                q = quote_identifier(c_key)
                if q not in select_cols:
                    select_cols.append(q)
            if status_key:
                select_cols.append(quote_identifier(status_key))
            data_rows = conn.execute(
                f"SELECT {', '.join(select_cols)} FROM {quote_identifier(table_name)} ORDER BY id ASC"
            ).fetchall()
            updates: list[tuple[object, ...]] = []
            for r in data_rows:
                values = {balance_key: r[balance_key]}
                if requested_key:
                    values[requested_key] = r[requested_key]
                if total_penalty_key:
                    values[total_penalty_key] = r[total_penalty_key]
                if penalty_paid_key:
                    values[penalty_paid_key] = r[penalty_paid_key]
                if granted_waiver_key:
                    values[granted_waiver_key] = r[granted_waiver_key]
                for c_key in comment_keys:
                    values[c_key] = r[c_key]
                if status_key:
                    values[status_key] = r[status_key]
                before_balance = str(values.get(balance_key, "") or "").strip()
                before_granted = str(values.get(granted_waiver_key, "") or "").strip() if granted_waiver_key else ""
                before_status = str(values.get(status_key, "") or "").strip() if status_key else ""
                apply_waiver_balance_and_status_to_values(values, mapping, file_name, sheet_name)
                after_balance = str(values.get(balance_key, "") or "").strip()
                after_granted = str(values.get(granted_waiver_key, "") or "").strip() if granted_waiver_key else ""
                after_status = str(values.get(status_key, "") or "").strip() if status_key else ""
                row_changed = (after_balance != before_balance)
                if granted_waiver_key and after_granted != before_granted:
                    row_changed = True
                if status_key and after_status != before_status:
                    row_changed = True
                if row_changed:
                    if status_key and granted_waiver_key:
                        updates.append((after_balance, after_granted, after_status, int(r["id"])))
                    elif status_key:
                        updates.append((after_balance, after_status, int(r["id"])))
                    elif granted_waiver_key:
                        updates.append((after_balance, after_granted, int(r["id"])))
                    else:
                        updates.append((after_balance, int(r["id"])))

            if not updates:
                continue
            if status_key and granted_waiver_key:
                conn.executemany(
                    f"""
                    UPDATE {quote_identifier(table_name)}
                    SET {quote_identifier(balance_key)} = ?, {quote_identifier(granted_waiver_key)} = ?, {quote_identifier(status_key)} = ?
                    WHERE id = ?
                    """,
                    updates,
                )
            elif status_key:
                conn.executemany(
                    f"""
                    UPDATE {quote_identifier(table_name)}
                    SET {quote_identifier(balance_key)} = ?, {quote_identifier(status_key)} = ?
                    WHERE id = ?
                    """,
                    updates,
                )
            elif granted_waiver_key:
                conn.executemany(
                    f"""
                    UPDATE {quote_identifier(table_name)}
                    SET {quote_identifier(balance_key)} = ?, {quote_identifier(granted_waiver_key)} = ?
                    WHERE id = ?
                    """,
                    updates,
                )
            else:
                conn.executemany(
                    f"""
                    UPDATE {quote_identifier(table_name)}
                    SET {quote_identifier(balance_key)} = ?
                    WHERE id = ?
                    """,
                    updates,
                )
            updated_rows += len(updates)
        if updated_rows > 0 or mapping_changed:
            conn.commit()
    if updated_rows > 0 or mapping_changed:
        bump_data_version()
    return updated_rows


@app.post("/uploaded/sheets/<int:sheet_id>/columns/<db_key>/delete")
def uploaded_delete_column(sheet_id: int, db_key: str):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")

    mapping = json.loads(meta["mapping_json"])
    existing_keys = [item["db_key"] for item in mapping]
    if db_key not in existing_keys:
        abort(404, description="Column not found.")
    if len(existing_keys) <= 1:
        abort(400, description="Cannot delete the last remaining column.")

    new_mapping = [item for item in mapping if item["db_key"] != db_key]
    validation_map = json.loads(meta["validation_json"])
    if isinstance(validation_map, dict):
        validation_map.pop(db_key, None)
    else:
        validation_map = {}

    with get_db() as conn:
        rebuild_uploaded_table_with_keys(conn, str(meta["table_name"]), [item["db_key"] for item in new_mapping])
        conn.execute(
            """
            UPDATE uploaded_sheets
            SET mapping_json = ?, validation_json = ?
            WHERE id = ?
            """,
            (json.dumps(new_mapping), json.dumps(validation_map), sheet_id),
        )
        conn.commit()
    bump_data_version()
    clear_sidebar_uploaded_links_cache()

    page, per_page = parse_paging(request.form)
    search_key = (request.form.get("search_key") or "").strip()
    search_query = (request.form.get("q") or "").strip()
    if search_key == db_key:
        search_key = ""
        search_query = ""
    return redirect(
        url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
    )


@app.post("/uploaded/sheets/<int:sheet_id>/row/<int:row_id>")
def uploaded_update_row(sheet_id: int, row_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    mapping = json.loads(meta["mapping_json"])
    table_name = meta["table_name"]

    with get_db() as conn:
        existing = conn.execute(
            f"SELECT * FROM {quote_identifier(table_name)} WHERE id = ?",
            (row_id,),
        ).fetchone()
    if existing is None:
        abort(404, description="Row not found.")

    updates = {}
    for item in mapping:
        key = item["db_key"]
        updates[key] = request.form.get(key, (existing[key] or "")).strip()
    apply_waiver_balance_and_status_to_values(
        updates,
        mapping,
        str(meta["original_filename"]),
        str(meta["sheet_name"]),
    )

    set_clause = ", ".join(f"{quote_identifier(key)} = ?" for key in updates)
    values = list(updates.values()) + [row_id]
    with get_db() as conn:
        conn.execute(
            f"UPDATE {quote_identifier(table_name)} SET {set_clause} WHERE id = ?",
            values,
        )
        conn.commit()
    bump_data_version()

    page, per_page = parse_paging(request.form)
    search_key = (request.form.get("search_key") or "").strip()
    search_query = (request.form.get("q") or "").strip()
    return redirect(
        url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
        + f"#row-{row_id}"
    )


@app.post("/uploaded/sheets/<int:sheet_id>/save-all")
def uploaded_save_all(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")

    payload = request.get_json(silent=True) or {}
    updates = payload.get("updates", [])
    if not isinstance(updates, list):
        abort(400, description="Invalid payload.")

    mapping = json.loads(meta["mapping_json"])
    updates = apply_waiver_balance_rule_to_updates(
        updates,
        mapping,
        str(meta["original_filename"]),
        str(meta["sheet_name"]),
    )
    updated = bulk_update_rows(meta["table_name"], mapping, updates)
    return jsonify({"ok": True, "updated_rows": updated})


@app.post("/uploaded/sheets/<int:sheet_id>/create-row")
def uploaded_create_row(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    mapping = json.loads(meta["mapping_json"])
    table_name = meta["table_name"]
    db_keys = [item["db_key"] for item in mapping]

    with get_db() as conn:
        next_excel_row = conn.execute(
            f"SELECT COALESCE(MAX(excel_row_number), 1) + 1 AS next_row FROM {quote_identifier(table_name)}"
        ).fetchone()["next_row"]
        cursor = conn.execute(
            f"INSERT INTO {quote_identifier(table_name)} (excel_row_number, "
            f"{', '.join(quote_identifier(k) for k in db_keys)}) VALUES "
            f"(?, {', '.join('?' for _ in db_keys)})",
            [next_excel_row] + ["" for _ in db_keys],
        )
        conn.execute("UPDATE uploaded_sheets SET row_count = row_count + 1 WHERE id = ?", (sheet_id,))
        conn.commit()
        new_id = cursor.lastrowid
    bump_data_version()

    _, per_page = parse_paging(request.form)
    search_key = (request.form.get("search_key") or "").strip()
    search_query = (request.form.get("q") or "").strip()
    with get_db() as conn:
        total_rows = conn.execute(
            f"SELECT COUNT(*) AS total FROM {quote_identifier(table_name)}"
        ).fetchone()["total"]
    target_page = max(1, (total_rows + per_page - 1) // per_page)
    return redirect(
        url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=target_page,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
        + f"#row-{new_id}"
    )


@app.post("/uploaded/sheets/<int:sheet_id>/row/<int:row_id>/delete")
def uploaded_delete_row(sheet_id: int, row_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    table_name = meta["table_name"]

    with get_db() as conn:
        conn.execute(f"DELETE FROM {quote_identifier(table_name)} WHERE id = ?", (row_id,))
        conn.execute(
            "UPDATE uploaded_sheets SET row_count = CASE WHEN row_count > 0 THEN row_count - 1 ELSE 0 END WHERE id = ?",
            (sheet_id,),
        )
        conn.commit()
    bump_data_version()

    page, per_page = parse_paging(request.form)
    search_key = (request.form.get("search_key") or "").strip()
    search_query = (request.form.get("q") or "").strip()
    return redirect(
        url_for(
            "uploaded_sheet_view",
            sheet_id=sheet_id,
            page=page,
            per_page=per_page,
            search_key=search_key,
            q=search_query,
        )
    )


@app.get("/uploaded/sheets/<int:sheet_id>/export")
def uploaded_export_sheet(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    mapping = json.loads(meta["mapping_json"])
    table_name = meta["table_name"]
    db_keys = [item["db_key"] for item in mapping]
    labels = [item["label"] for item in mapping]

    with get_db() as conn:
        rows = conn.execute(
            f"SELECT {', '.join(quote_identifier(k) for k in db_keys)} "
            f"FROM {quote_identifier(table_name)} ORDER BY excel_row_number"
        ).fetchall()

    records = [{labels[i]: row[db_keys[i]] for i in range(len(db_keys))} for row in rows]
    export_df = pd.DataFrame(records, columns=labels)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_df.to_excel(writer, sheet_name=meta["sheet_name"], index=False)
    output.seek(0)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{Path(meta['original_filename']).stem}_{meta['sheet_name']}_{stamp}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/annual-returns/<sheet_name>/dashboard")
def annual_dashboard(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")
    dataset = (request.args.get("dataset") or "").strip()
    if dataset and dataset != "annual_returns":
        abort(400, description="Dataset mismatch for dashboard request.")
    mapping = get_mapping()
    thresholds = get_dashboard_thresholds("annual", sheet_name)
    payload = get_cached_dashboard_payload("annual_returns", mapping, sheet_name_hint=sheet_name, thresholds=thresholds)
    return render_template(
        "dashboard.html",
        ai_context={"type": "annual", "sheet_name": sheet_name},
        dashboard_title=f"{EXCEL_PATH.name} - {sheet_name}",
        back_url=url_for("annual_returns_sheet", sheet_name=sheet_name),
        chat_url=url_for("annual_dashboard_chat", sheet_name=sheet_name),
        summary_docx_url=url_for("annual_dashboard_summary_docx", sheet_name=sheet_name),
        payload=payload,
    )


@app.post("/annual-returns/<sheet_name>/dashboard/chat")
def annual_dashboard_chat(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")
    question = (request.get_json(silent=True) or {}).get("question", "").strip()
    if not question:
        abort(400, description="Question is required.")
    answer = answer_annual_question(sheet_name, question)
    return jsonify({"ok": True, "answer": answer})


@app.get("/annual-returns/<sheet_name>/dashboard/summary.docx")
def annual_dashboard_summary_docx(sheet_name: str):
    if sheet_name != SHEET_NAME:
        abort(404, description=f"Sheet '{sheet_name}' not found.")
    mapping = get_mapping()
    thresholds = get_dashboard_thresholds("annual", sheet_name)
    payload = get_cached_dashboard_payload("annual_returns", mapping, sheet_name_hint=sheet_name, thresholds=thresholds)
    summary_text = str(payload.get("submission_summary", "")).strip()
    output = summary_text_to_docx(summary_text, f"{sheet_name} Submission Summary")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{Path(EXCEL_PATH.name).stem}_{sheet_name}_submission_summary_{stamp}.docx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/uploaded/sheets/<int:sheet_id>/dashboard")
def uploaded_dashboard(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    file_id = (request.args.get("file_id") or "").strip()
    dataset = (request.args.get("dataset") or "").strip()
    if file_id and str(meta["file_id"]) != file_id:
        abort(400, description="File and sheet mismatch for dashboard request.")
    if dataset and dataset != meta["table_name"]:
        abort(400, description="Dataset mismatch for dashboard request.")
    mapping = json.loads(meta["mapping_json"])
    thresholds = get_dashboard_thresholds("uploaded", str(sheet_id))
    payload = get_cached_dashboard_payload(
        meta["table_name"],
        mapping,
        sheet_name_hint=meta["sheet_name"],
        snapshot_sheet_id=sheet_id,
        thresholds=thresholds,
    )
    return render_template(
        "dashboard.html",
        ai_context={"type": "uploaded", "sheet_id": sheet_id},
        dashboard_title=f"{meta['original_filename']} - {meta['sheet_name']}",
        back_url=url_for("uploaded_sheet_view", sheet_id=sheet_id),
        chat_url=url_for("uploaded_dashboard_chat", sheet_id=sheet_id),
        summary_docx_url=url_for("uploaded_dashboard_summary_docx", sheet_id=sheet_id),
        payload=payload,
    )


@app.post("/uploaded/sheets/<int:sheet_id>/dashboard/chat")
def uploaded_dashboard_chat(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    question = (request.get_json(silent=True) or {}).get("question", "").strip()
    if not question:
        abort(400, description="Question is required.")
    answer = answer_uploaded_question(sheet_id, question)
    return jsonify({"ok": True, "answer": answer})


@app.get("/uploaded/sheets/<int:sheet_id>/dashboard/summary.docx")
def uploaded_dashboard_summary_docx(sheet_id: int):
    meta = get_uploaded_sheet_meta(sheet_id)
    if meta is None:
        abort(404, description="Uploaded sheet not found.")
    mapping = json.loads(meta["mapping_json"])
    thresholds = get_dashboard_thresholds("uploaded", str(sheet_id))
    payload = get_cached_dashboard_payload(
        meta["table_name"],
        mapping,
        sheet_name_hint=meta["sheet_name"],
        snapshot_sheet_id=sheet_id,
        thresholds=thresholds,
    )
    summary_text = str(payload.get("submission_summary", "")).strip()
    output = summary_text_to_docx(summary_text, f"{meta['sheet_name']} Submission Summary")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{Path(meta['original_filename']).stem}_{meta['sheet_name']}_submission_summary_{stamp}.docx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/dashboard/drilldown")
def dashboard_drilldown():
    payload = request.get_json(silent=True) or {}
    context = payload.get("context", {})
    drill = payload.get("drilldown", {})
    label = str(payload.get("label", "")).strip()
    label_index = payload.get("label_index")
    if not isinstance(context, dict) or not isinstance(drill, dict):
        abort(400, description="Invalid drilldown payload.")

    ctx_type = str(context.get("type") or "")
    mapping: list[dict[str, str]]
    table_name: str
    view_url: str
    if ctx_type == "uploaded":
        try:
            sheet_id = int(context.get("sheet_id"))
        except Exception:
            abort(400, description="Invalid uploaded context.")
        meta = get_uploaded_sheet_meta(sheet_id)
        if meta is None:
            abort(404, description="Uploaded sheet not found.")
        mapping = json.loads(meta["mapping_json"])
        table_name = str(meta["table_name"])
        view_url = url_for("uploaded_sheet_view", sheet_id=sheet_id)
    else:
        sheet_name = str(context.get("sheet_name") or SHEET_NAME)
        if sheet_name != SHEET_NAME:
            abort(400, description="Invalid annual context.")
        mapping = get_mapping()
        table_name = "annual_returns"
        view_url = url_for("annual_returns_sheet", sheet_name=sheet_name)

    keys = [item["db_key"] for item in mapping]
    labels = {item["db_key"]: item["label"] for item in mapping}
    with get_db() as conn:
        rows = conn.execute(
            f"SELECT id, {', '.join(quote_identifier(k) for k in keys)} FROM {quote_identifier(table_name)}"
        ).fetchall()
    frame = pd.DataFrame([{k: row[k] for k in keys} | {"id": row["id"]} for row in rows], columns=["id"] + keys)
    if frame.empty:
        return jsonify({"ok": True, "total": 0, "columns": [], "rows": [], "view_url": view_url})

    d_type = str(drill.get("type") or "")
    col = str(drill.get("column") or "")
    if col not in keys:
        abort(400, description="Invalid drilldown column.")

    mask = pd.Series([True] * len(frame), index=frame.index)
    if d_type == "categorical":
        mask = frame[col].astype(str).str.strip().str.lower() == label.lower()
    elif d_type == "date_month":
        parsed = to_date_series(frame[col])
        mask = parsed.dt.to_period("M").astype(str) == label
    elif d_type == "numeric_bucket":
        numeric = pd.to_numeric(clean_numeric_text(frame[col]), errors="coerce")
        bucket_ranges = drill.get("bucket_ranges", [])
        bucket = None
        if isinstance(label_index, int) and isinstance(bucket_ranges, list) and 0 <= label_index < len(bucket_ranges):
            bucket = bucket_ranges[label_index]
        if bucket and isinstance(bucket, dict):
            left = float(bucket.get("left", float("-inf")))
            right = float(bucket.get("right", float("inf")))
            mask = numeric.ge(left) & numeric.le(right)
        else:
            abort(400, description="Invalid numeric bucket selection.")

    filtered = frame[mask].copy()
    display_cols = keys[:8]
    out_rows = []
    for _, row in filtered.head(50).iterrows():
        out_row = {"id": int(row["id"])}
        for k in display_cols:
            out_row[labels.get(k, k)] = str(row[k] or "")
        out_rows.append(out_row)
    return jsonify(
        {
            "ok": True,
            "total": int(len(filtered)),
            "columns": ["id"] + [labels.get(k, k) for k in display_cols],
            "rows": out_rows,
            "view_url": view_url,
        }
    )


@app.post("/dashboard/thresholds")
def dashboard_save_thresholds():
    payload = request.get_json(silent=True) or {}
    context = payload.get("context", {})
    config = payload.get("thresholds", {})
    if not isinstance(context, dict) or not isinstance(config, dict):
        abort(400, description="Invalid threshold payload.")

    ctx_type = str(context.get("type") or "").strip().lower()
    if ctx_type == "uploaded":
        try:
            sheet_id = int(context.get("sheet_id"))
        except Exception:
            abort(400, description="Invalid uploaded context.")
        meta = get_uploaded_sheet_meta(sheet_id)
        if meta is None:
            abort(404, description="Uploaded sheet not found.")
        saved = save_dashboard_thresholds("uploaded", str(sheet_id), config)
        return jsonify({"ok": True, "saved": saved})

    sheet_name = str(context.get("sheet_name") or SHEET_NAME)
    if sheet_name != SHEET_NAME:
        abort(400, description="Invalid annual context.")
    saved = save_dashboard_thresholds("annual", sheet_name, config)
    return jsonify({"ok": True, "saved": saved})


@app.post("/ai/chat")
def global_ai_chat():
    payload = request.get_json(silent=True) or {}
    question = str(payload.get("question", "")).strip()
    if not question:
        abort(400, description="Question is required.")

    context = payload.get("context", {})
    if not isinstance(context, dict):
        context = {}
    inferred = infer_ai_dataset_from_question(question)
    active = inferred if inferred else context

    ctx_type = str(active.get("type", "")).strip().lower()
    if ctx_type == "uploaded":
        try:
            sheet_id = int(active.get("sheet_id"))
        except Exception:
            return jsonify({"ok": True, "answer": "I could not resolve an uploaded sheet from your question."})
        return jsonify({"ok": True, "answer": answer_uploaded_question(sheet_id, question)})

    sheet_name = str(active.get("sheet_name") or SHEET_NAME)
    return jsonify({"ok": True, "answer": answer_annual_question(sheet_name, question)})


try:
    init_database()
    ensure_uploaded_registry()
    if RUN_STARTUP_MAINTENANCE:
        sanitize_existing_uploaded_sheets_remove_unnamed()
        enforce_waiver_balance_rule_for_existing_uploaded_sheets()
    if AUTO_SYNC_DEFAULT_WORKBOOKS:
        threading.Thread(target=sync_default_workbooks_from_static_files, daemon=True).start()
except Exception:
    app.logger.exception("Startup initialization failed.")

if __name__ == "__main__":
    init_database()
    ensure_uploaded_registry()
    if RUN_STARTUP_MAINTENANCE:
        sanitize_existing_uploaded_sheets_remove_unnamed()
        enforce_waiver_balance_rule_for_existing_uploaded_sheets()
    if AUTO_SYNC_DEFAULT_WORKBOOKS:
        sync_default_workbooks_from_static_files()
    app.run(debug=True)
